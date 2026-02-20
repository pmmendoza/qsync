"""Export survey content into a translation-review Word document (.docx)."""

from __future__ import annotations

import base64
import csv
import html as _html
import json
import os
import re
from urllib.parse import urlencode
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Mapping, Tuple

from .config import get_client_config, resolve_root, resolve_scoped_dir
from .excel_io import EMBEDDED_EMPTY_VALUE, build_embedded_data_rows
from .flow_traversal import (
    FlowTraversalHandlers,
    eval_boolean_expression as _eval_boolean_expression,
    eval_boolean_expression_with_unasked_selected_false as _eval_boolean_expression_with_unasked_selected_false,
    walk_flow,
)
from .markdown_codec import is_markdown_safe_html, should_treat_as_html
from .question_types import is_system_question_type
from .qualtrics_client import (
    load_cached_survey,
    fetch_survey_definition_live,
    refresh_survey_cache,
)
from .dimensions.translations_language_blocks import (
    get_base_language,
    list_enabled_languages,
    read_answer_display,
    read_choice_display,
    read_label_display,
    read_question_text,
)
from .translations_utils import normalize_language_code

EXPORT_DIRNAME = "export"
_EDF_TOKEN_RE = re.compile(r"\$\{e://Field/[^}]+\}")
_DANGEROUS_HTML_RE = re.compile(
    r"</?\s*(script|style|form|input|textarea|select|option|button|iframe|svg|canvas|img|video|audio)\b",
    re.IGNORECASE,
)
_TABLE_TAG_RE = re.compile(r"(?is)</?\s*table\b")


_DEFAULT_BASE_LANGUAGE = "EN"

_PUBLISHED_VERSION_CACHE: dict[str, dict[str, Any] | None] = {}


def _load_latest_published_version_from_push_log(
    survey_id: str,
) -> dict[str, Any] | None:
    """Best-effort lookup of the latest published version metadata from local logs.

    This avoids any extra network calls during export (important for tests and
    offline workflows). When the log file is missing or the survey hasn't been
    published from this workspace, returns None.
    """

    if survey_id in _PUBLISHED_VERSION_CACHE:
        return _PUBLISHED_VERSION_CACHE[survey_id]

    root = resolve_root(required=False) or Path.cwd()
    log_path = root / "logs" / "qualtrics_push.log"
    if not log_path.exists():
        _PUBLISHED_VERSION_CACHE[survey_id] = None
        return None

    latest_ts: datetime | None = None
    latest_meta: dict[str, Any] | None = None

    for raw in log_path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw.strip()
        if not line or '"action": "qsync.survey.publish.definition"' not in line:
            continue
        if f'"survey_id": "{survey_id}"' not in line:
            continue
        try:
            obj = json.loads(line)
        except Exception:
            continue
        if obj.get("action") != "qsync.survey.publish.definition":
            continue
        if obj.get("survey_id") != survey_id:
            continue
        ts_raw = str(obj.get("timestamp") or "")
        try:
            ts = datetime.fromisoformat(ts_raw.replace("Z", "+00:00"))
        except Exception:
            continue
        if latest_ts is None or ts > latest_ts:
            latest_ts = ts
            meta = obj.get("meta") or {}
            latest_meta = dict(meta) if isinstance(meta, dict) else {}

    _PUBLISHED_VERSION_CACHE[survey_id] = latest_meta
    return latest_meta


def _normalize_lang_code(value: str | None) -> str:
    return normalize_language_code(value or "")


def _normalize_label(value: str | None) -> str | None:
    if value is None:
        return None
    label = str(value).strip()
    return label or None


def _resolve_compare_labels(
    compare_labels: tuple[str, str] | None,
    *,
    base_language: str | None,
    target_language: str | None,
    fallback_base: str | None = None,
    fallback_target: str | None = None,
) -> tuple[str, str]:
    base_label = None
    target_label = None
    if compare_labels:
        base_label = _normalize_label(compare_labels[0])
        target_label = _normalize_label(compare_labels[1])
    if base_label is None:
        base_label = _normalize_label(fallback_base) or (
            _normalize_lang_code(base_language) or _DEFAULT_BASE_LANGUAGE
        )
    if target_label is None:
        target_label = _normalize_label(fallback_target) or (
            _normalize_lang_code(target_language) or "TARGET"
        )
    return base_label, target_label


def _survey_result(payload: dict) -> dict:
    if isinstance(payload.get("result"), dict):
        return payload["result"]
    return payload


def _survey_options(payload: dict) -> dict:
    return _survey_result(payload).get("SurveyOptions") or {}


def _parse_timestamp(value: object) -> datetime | None:
    if value is None:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    try:
        if raw.endswith("Z"):
            raw = raw[:-1] + "+00:00"
        return datetime.fromisoformat(raw)
    except ValueError:
        return None


def _preflight_cache_freshness(
    survey_id: str,
    *,
    interactive: bool,
    allow_prompt: bool = True,
) -> None:
    from .terminal_output import info, error
    from .survey_inventory import load_inventory_record
    from .interactive_menu import confirm, is_interactive

    try:
        cached = load_cached_survey(survey_id)
    except Exception as e:
        error(
            "[qsync:export]",
            f"WARNING: Unable to load cached survey definition for freshness check ({e}).",
        )
        return

    cached_payload = _survey_result(cached.payload or {})
    cached_last = (
        cached_payload.get("LastModified")
        or cached_payload.get("lastModified")
        or cached_payload.get("LastModifiedDate")
    )
    cached_dt = _parse_timestamp(cached_last)

    live_last = None
    try:
        live_payload = fetch_survey_definition_live(survey_id)
        live_result = _survey_result(live_payload or {})
        live_last = (
            live_result.get("LastModified")
            or live_result.get("lastModified")
            or live_result.get("LastModifiedDate")
        )
    except Exception:
        live_last = None

    if live_last is None:
        record = load_inventory_record(survey_id) or {}
        live_last = record.get("lastModified") or record.get("lastModifiedDate")

    live_dt = _parse_timestamp(live_last)

    if not cached_dt or not live_dt:
        info(
            "[qsync:export]",
            "WARNING: Could not compare cache freshness (missing timestamps).",
        )
        return

    if live_dt <= cached_dt:
        return

    message = (
        f"Cached survey definition appears stale (cache={cached_dt.isoformat()}, "
        f"live={live_dt.isoformat()}). Refresh cache now?"
    )
    if interactive and allow_prompt and is_interactive():
        if confirm(message=message, default=True):
            try:
                refresh_survey_cache(survey_id)
                info("[qsync:export]", "Refreshed cached survey definition from API.")
            except Exception as e:
                info(
                    "[qsync:export]",
                    f"WARNING: Failed to refresh cache automatically ({e}). Proceeding with stale cache.",
                )
        else:
            info(
                "[qsync:export]",
                "WARNING: Proceeding with stale cache (live survey is newer).",
            )
    else:
        info(
            "[qsync:export]",
            "WARNING: Live survey is newer than cached definition. Proceeding without refresh.",
        )


def _metadata_translations_for_language(payload: dict, language: str) -> dict[str, str]:
    lang = _normalize_lang_code(language)
    options = _survey_options(payload)
    meta = options.get("MetaDataTranslations")
    if not isinstance(meta, dict) or not lang:
        return {}
    entry = meta.get(lang) or meta.get(lang.lower()) or meta.get(lang.upper())
    if isinstance(entry, dict):
        return {str(k): str(v) for k, v in entry.items() if v is not None}
    return {}


def _language_present_in_cache(payload: dict, language: str) -> bool:
    lang = _normalize_lang_code(language)
    if not lang:
        return False
    enabled = list_enabled_languages(payload)
    if lang in enabled:
        return True
    options = _survey_options(payload)
    meta = options.get("MetaDataTranslations")
    if isinstance(meta, dict) and lang in meta:
        return True
    questions = _survey_result(payload).get("Questions") or {}
    for question in questions.values():
        lang_block = question.get("Language") or {}
        if isinstance(lang_block, dict) and lang in lang_block:
            return True
    return False


def _is_sbs_matrix_question(question: dict) -> bool:
    """Return True for SBS side-by-side matrix questions."""

    return (
        str(question.get("QuestionType") or "").strip() == "SBS"
        and str(question.get("Selector") or "").strip() == "SBSMatrix"
    )


def _ordered_numeric_string_ids(mapping: dict, order: object | None = None) -> list[str]:
    """Return ordered IDs from a mapping, honoring an optional order list."""

    if not isinstance(mapping, dict):
        return []

    ordered: list[str] = []
    if isinstance(order, list) and order:
        ordered.extend(str(item) for item in order if str(item) in mapping)
    for item_id in mapping.keys():
        item_id_s = str(item_id)
        if item_id_s not in ordered:
            ordered.append(item_id_s)
    return ordered


def _lookup_ordered_mapping_item(mapping: dict, item_id: str) -> Any | None:
    """Lookup a mapping item by string key and common numeric fallback."""

    if item_id in mapping:
        return mapping[item_id]
    if item_id.isdigit():
        as_int = int(item_id)
        if as_int in mapping:
            return mapping[as_int]
    return None


def _read_lang_block(question: dict, lang: str) -> dict:
    """Best-effort helper to read a language block from question payload."""

    language_blocks = question.get("Language")
    if not isinstance(language_blocks, dict):
        return {}
    lang_block = (
        language_blocks.get(lang)
        or language_blocks.get(lang.lower())
        or language_blocks.get(lang.upper())
    )
    if isinstance(lang_block, dict):
        return lang_block
    return {}


def build_translation_map_from_cache(
    payload: dict,
    *,
    language: str,
    base_language: str,
) -> dict[str, str]:
    result = _survey_result(payload)
    questions = result.get("Questions") or {}
    lang = _normalize_lang_code(language)
    base_lang = _normalize_lang_code(base_language)
    is_base = lang == base_lang

    translation_map: dict[str, str] = {}
    survey_options = result.get("SurveyOptions") or {}
    if not isinstance(survey_options, dict):
        survey_options = {}

    if is_base:
        survey_title = str(
            result.get("SurveyTitle") or survey_options.get("SurveyTitle") or ""
        )
        survey_description = str(
            result.get("SurveyDescription")
            or result.get("SurveyMetaDescription")
            or survey_options.get("SurveyMetaDescription")
            or survey_options.get("SurveyDescription")
            or ""
        )
        if survey_title:
            translation_map["SurveyTitle"] = survey_title
        if survey_description:
            translation_map["SurveyDescription"] = survey_description
    else:
        meta = _metadata_translations_for_language(payload, lang)
        if "SurveyTitle" in meta:
            translation_map["SurveyTitle"] = str(meta.get("SurveyTitle") or "")
        if "SurveyDescription" in meta:
            translation_map["SurveyDescription"] = str(
                meta.get("SurveyDescription") or ""
            )

    for qid, question in questions.items():
        qid_s = str(qid)
        if is_base:
            q_text = question.get("QuestionText")
        else:
            q_text = read_question_text(question, lang)
        translation_map[f"{qid_s}_QuestionText"] = _coerce_display_text(q_text)

        is_sbs_matrix = _is_sbs_matrix_question(question)

        choices = question.get("Choices") or {}
        if isinstance(choices, dict):
            for cid, choice in choices.items():
                if is_base:
                    display = (choice or {}).get("Display")
                else:
                    display = read_choice_display(question, lang, str(cid))
                translation_map[f"{qid_s}_Choice{cid}"] = _coerce_display_text(display)

        if is_sbs_matrix:
            additional = question.get("AdditionalQuestions") or {}
            if isinstance(additional, dict) and additional:
                lang_block = _read_lang_block(question, lang)
                add_lang = (
                    lang_block.get("AdditionalQuestions")
                    if isinstance(lang_block.get("AdditionalQuestions"), dict)
                    else {}
                )
                q_lang_answers = (
                    lang_block.get("Answers")
                    if isinstance(lang_block.get("Answers"), dict)
                    else {}
                )
                q_answers = question.get("Answers") or {}
                if not isinstance(q_answers, dict):
                    q_answers = {}
                for column_id in _ordered_numeric_string_ids(additional):
                    column = additional.get(column_id)
                    if not isinstance(column, dict):
                        continue
                    base_answers = column.get("Answers") or {}
                    if not isinstance(base_answers, dict):
                        base_answers = {}
                    if not base_answers:
                        base_answers = q_answers

                    col_lang_block: dict[str, Any] = {}
                    if is_base:
                        col_text = column.get("QuestionText")
                    else:
                        maybe_block = (
                            add_lang.get(str(column_id))
                            if isinstance(add_lang, dict)
                            else None
                        )
                        if isinstance(maybe_block, dict):
                            col_lang_block = maybe_block
                            col_text = maybe_block.get("QuestionText")
                        else:
                            col_text = None
                    translation_map[
                        f"{qid_s}#{column_id}_QuestionText"
                    ] = _coerce_display_text(col_text)

                    answers = base_answers
                    answer_order = column.get("AnswerOrder")
                    lang_answers = (
                        col_lang_block.get("Answers")
                        if isinstance(col_lang_block, dict)
                        else None
                    )
                    if not isinstance(lang_answers, dict):
                        lang_answers = {}
                    column_answer_order = answer_order if isinstance(answer_order, list) else None
                    if not answers:
                        answers = q_answers
                    for ans_id in _ordered_numeric_string_ids(answers, column_answer_order):
                        answer = _lookup_ordered_mapping_item(answers, ans_id)
                        if is_base:
                            ans_text = (answer or {}).get("Display")
                        else:
                            ans_lang = _lookup_ordered_mapping_item(lang_answers, ans_id)
                            if not isinstance(ans_lang, dict):
                                ans_lang = {}
                            if not ans_lang:
                                ans_lang = _lookup_ordered_mapping_item(q_lang_answers, ans_id)
                                if not isinstance(ans_lang, dict):
                                    ans_lang = {}
                            ans_text = (
                                ans_lang.get("Display") if isinstance(ans_lang, dict) else None
                            )
                            if ans_text is None:
                                ans_text = _coerce_display_text((answer or {}).get("Display"))
                                if ans_text == "":
                                    ans_text = None
                        translation_map[
                            f"{qid_s}#{column_id}_Answer{ans_id}"
                        ] = _coerce_display_text(ans_text)
            # SBSMatrix carries per-column answers in AdditionalQuestions.
            # Keep this branch intentionally isolated from top-level Answers.
        else:
            answers = question.get("Answers") or {}
            if isinstance(answers, dict):
                for aid, answer in answers.items():
                    if is_base:
                        display = (answer or {}).get("Display")
                    else:
                        display = read_answer_display(question, lang, str(aid))
                    translation_map[f"{qid_s}_Answer{aid}"] = _coerce_display_text(display)

        labels = question.get("Labels") or {}
        if isinstance(labels, dict):
            for lid, label in labels.items():
                if is_base:
                    display = (label or {}).get("Display")
                else:
                    display = read_label_display(question, lang, str(lid))
                translation_map[f"{qid_s}_Label{lid}"] = _coerce_display_text(display)

    return translation_map


@dataclass
class TranslationRenderPlan:
    """Precomputed translation coverage for rendering a language export."""

    survey_id: str
    base_language: str
    target_language: str
    expected_keys: list[str]
    base_empty_keys: set[str]
    missing_keys: list[str]
    empty_but_base_nonempty_keys: list[str]

    @property
    def total_expected(self) -> int:
        return len(self.expected_keys)

    @property
    def total_missing(self) -> int:
        return len(self.missing_keys)

    @property
    def total_empty_but_base_nonempty(self) -> int:
        return len(self.empty_but_base_nonempty_keys)

    @property
    def total_ok(self) -> int:
        # "OK" = has a non-empty translation OR is allowed-empty because base is empty.
        return max(
            0,
            self.total_expected
            - self.total_missing
            - self.total_empty_but_base_nonempty,
        )


@dataclass
class TranslationRenderContext:
    survey_id: str
    base_language: str
    target_language: str
    target_map: Mapping[str, str]
    base_map: Mapping[str, str] | None
    compare_to_base: bool
    plan: TranslationRenderPlan | None = None

    def key_for_question_text(self, qid: str) -> str:
        return f"{qid}_QuestionText"

    def key_for_choice(self, qid: str, choice_id: str) -> str:
        return f"{qid}_Choice{choice_id}"

    def key_for_answer(self, qid: str, answer_id: str) -> str:
        return f"{qid}_Answer{answer_id}"

    def key_for_sbs_column_question_text(self, qid: str, column_id: str) -> str:
        return f"{qid}#{column_id}_QuestionText"

    def key_for_sbs_column_answer(
        self, qid: str, column_id: str, answer_id: str
    ) -> str:
        return f"{qid}#{column_id}_Answer{answer_id}"

    def key_for_label(self, qid: str, label_id: str) -> str:
        return f"{qid}_Label{label_id}"


@dataclass
class ExportContent:
    """Prepared content for export (format-agnostic).

    This dataclass holds all prepared survey data for export rendering.
    It separates content preparation from format-specific rendering (DOCX/PDF),
    enabling code reuse and simpler testing.
    """

    survey_id: str
    survey_name: str
    survey_title: str | None
    survey_description: str | None
    version_number: int | None
    version_id: str | None
    version_description: str | None
    survey_payload: dict[str, Any]
    survey_link: str
    active_qids: set[str]
    translation_ctx: TranslationRenderContext | None
    render_plan: TranslationRenderPlan | None
    qid_to_js: dict[str, str]
    mermaid_code: str | None
    mermaid_path: Path | None
    mermaid_image_path: Path | None
    edf_overrides: dict[str, str] | None
    include_html_source: bool
    layout_heuristics: bool
    compare_to_base: bool
    render_language: str | None
    base_language: str
    output_path: Path
    include_js_strings: bool
    flow_trace: Callable[[str], None] | None
    compare_labels: tuple[str, str] | None = None
    compare_survey_id: str | None = None
    compare_survey_name: str | None = None
    compare_survey_link: str | None = None
    compare_survey_base_language: str | None = None


def _sanitize_filename(value: str) -> str:
    s = "".join(c if c.isalnum() or c in " -_." else "_" for c in str(value or ""))
    s = " ".join(s.split()).strip().strip(".")
    return s or "export"


def _resolve_output_docx_path(
    *,
    survey_id: str,
    survey_name: str,
    export_dir: Path,
    output_path: Path | None,
    smart_name: bool,
    render_language: str | None,
    compare_to_base: bool,
    base_language: str,
    format: str = "docx",
) -> Path:
    """Resolve and validate output path for export (DOCX or PDF).

    Args:
        survey_id: Survey ID
        survey_name: Survey name
        export_dir: Default export directory
        output_path: User-specified output path (optional)
        smart_name: Use survey name + timestamp
        render_language: Translation language
        compare_to_base: Bilingual mode
        base_language: Base language for default naming
        format: Output format ("docx" or "pdf")

    Returns:
        Resolved output path
    """
    ext = f".{format}"
    base = _normalize_lang_code(base_language) or _DEFAULT_BASE_LANGUAGE

    def default_name() -> str:
        lang = _normalize_lang_code(render_language)
        lang_tag = ""
        if lang:
            if compare_to_base and format == "pdf":
                # PDFs do not present a robust side-by-side language comparison in practice,
                # so avoid naming them as if they were bilingual exports.
                lang_tag = f"__{lang}"
            else:
                lang_tag = f"__{base}-{lang}" if compare_to_base else f"__{lang}"
        else:
            lang_tag = f"__{base}"
        # Always use slug__surveyid format (consistent with survey JSON cache naming)
        safe = _sanitize_filename(survey_name) if survey_name else survey_id
        if not smart_name:
            return f"{safe}__{survey_id}{lang_tag}{ext}"
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{safe}__{survey_id}{lang_tag}__{stamp}{ext}"

    if output_path is None:
        return export_dir / default_name()

    output_path = Path(output_path)
    if output_path.is_dir():
        return output_path / default_name()

    if not output_path.parent.exists():
        raise ValueError(f"Output directory does not exist: {output_path.parent}")

    if output_path.suffix == "":
        return output_path.with_suffix(ext)
    if output_path.suffix.lower() != ext:
        raise ValueError(f"Output path must be a {ext} file (got: {output_path})")
    return output_path


def _resolve_output_side_by_side_docx_path(
    *,
    survey_a_id: str,
    survey_a_name: str,
    survey_b_id: str,
    survey_b_name: str,
    export_dir: Path,
    output_path: Path | None,
    smart_name: bool,
) -> Path:
    ext = ".docx"

    def default_name() -> str:
        safe_a = _sanitize_filename(survey_a_name) if survey_a_name else survey_a_id
        safe_b = _sanitize_filename(survey_b_name) if survey_b_name else survey_b_id
        base = f"{safe_a}__{survey_a_id}__VS__{safe_b}__{survey_b_id}"
        if not smart_name:
            return f"{base}{ext}"
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{base}__{stamp}{ext}"

    if output_path is None:
        return export_dir / default_name()

    output_path = Path(output_path)
    if output_path.is_dir():
        return output_path / default_name()

    if not output_path.parent.exists():
        raise ValueError(f"Output directory does not exist: {output_path.parent}")

    if output_path.suffix == "":
        return output_path.with_suffix(ext)
    if output_path.suffix.lower() != ext:
        raise ValueError(f"Output path must be a {ext} file (got: {output_path})")
    return output_path


def _build_survey_link(
    survey_id: str, *, edf_overrides: dict[str, str] | None, language: str | None
) -> str | None:
    try:
        base_url, _ = get_client_config()
    except Exception:
        return None

    base = str(base_url or "").strip()
    if not base:
        return None
    if not base.startswith("http"):
        base = "https://" + base.lstrip("/")
    if "/API/v3" in base:
        base = base.split("/API/v3", 1)[0]
    base = base.rstrip("/")

    url = f"{base}/jfe/form/{survey_id}"
    params: list[tuple[str, str]] = []
    if language:
        params.append(("Q_Language", _normalize_lang_code(language)))
    if edf_overrides:
        params.extend([(str(k), str(v)) for k, v in sorted(edf_overrides.items())])
    if params:
        query = urlencode(params)
        if query:
            url = f"{url}?{query}"
    return url


def _collect_block_question_ids(
    result: Mapping[str, Any],
    *,
    include_blocks: set[str],
) -> set[str]:
    blocks = result.get("Blocks") or {}
    questions = result.get("Questions") or {}
    if not isinstance(blocks, Mapping) or not isinstance(questions, Mapping):
        return set()
    wanted_blocks = {str(item).strip() for item in include_blocks if str(item).strip()}
    if not wanted_blocks:
        return set()

    qids: set[str] = set()
    for block_id in wanted_blocks:
        block = blocks.get(block_id)
        if not isinstance(block, Mapping):
            continue
        for elem in block.get("BlockElements", []) or []:
            if not isinstance(elem, Mapping):
                continue
            if str(elem.get("Type") or "") != "Question":
                continue
            qid = str(elem.get("QuestionID") or "").strip()
            if qid and qid in questions:
                qids.add(qid)
    return qids


def _apply_active_qid_filters(
    *,
    result: Mapping[str, Any],
    active_qids: set[str],
    include_qids: set[str] | None,
    include_tags: set[str] | None,
    include_blocks: set[str] | None,
) -> set[str]:
    include_qids_set = {str(item).strip() for item in (include_qids or set()) if str(item).strip()}
    include_tags_set = {str(item).strip() for item in (include_tags or set()) if str(item).strip()}
    include_blocks_set = {str(item).strip() for item in (include_blocks or set()) if str(item).strip()}
    if not include_qids_set and not include_tags_set and not include_blocks_set:
        return active_qids

    questions = result.get("Questions") or {}
    if not isinstance(questions, Mapping):
        return set()

    allowed: set[str] = set()
    if include_qids_set:
        allowed.update({qid for qid in include_qids_set if qid in questions})
    if include_tags_set:
        for qid, payload in questions.items():
            if not isinstance(payload, Mapping):
                continue
            tag = str(payload.get("DataExportTag") or "").strip()
            if tag and tag in include_tags_set:
                allowed.add(str(qid))
    if include_blocks_set:
        allowed.update(
            _collect_block_question_ids(result, include_blocks=include_blocks_set)
        )

    return {qid for qid in active_qids if qid in allowed}


def _prepare_export_content(
    survey_id: str,
    survey_payload: dict,
    output_path: Path,
    *,
    mermaid_path: Path | None = None,
    mermaid_image_path: Path | None = None,
    edf_overrides: dict[str, str] | None = None,
    mapping_path: Path | None = None,
    include_html_source: bool = True,
    layout_heuristics: bool = False,
    render_language: str | None = None,
    compare_to_base: bool = False,
    include_qids: set[str] | None = None,
    include_tags: set[str] | None = None,
    include_blocks: set[str] | None = None,
    include_js_strings: bool = True,
    flow_trace: Callable[[str], None] | None = None,
    translation_ctx_override: TranslationRenderContext | None = None,
    render_plan_override: TranslationRenderPlan | None = None,
    compare_labels: tuple[str, str] | None = None,
    compare_survey_id: str | None = None,
    compare_survey_name: str | None = None,
    compare_survey_link: str | None = None,
    compare_survey_base_language: str | None = None,
) -> ExportContent:
    """Prepare format-agnostic export content from survey payload.

    This function extracts and prepares all survey data needed for export,
    including translation overlays, flow analysis, and metadata extraction.
    The returned ExportContent can be rendered to any format (DOCX, PDF).

    Args:
        survey_id: Survey ID
        survey_payload: Raw survey definition payload (from API or cache)
        output_path: Target output file path
        mermaid_path: Optional path for Mermaid flow diagram (.mmd)
        mermaid_image_path: Optional path for rendered Mermaid image (.png)
        edf_overrides: Optional embedded data field overrides for scenario filtering
        mapping_path: Optional path to QID→JS mapping CSV
        include_html_source: Whether to include raw HTML source blocks
        layout_heuristics: Whether to apply reviewer-friendly layout transforms
        render_language: Optional translation language code (e.g., "FR")
        compare_to_base: Whether to render bilingual (base + target) mode
        include_js_strings: Whether to extract and display user-visible JS strings
        translation_ctx_override: Optional override for the translation context
        render_plan_override: Optional override for the translation render plan
        compare_labels: Optional labels for side-by-side rendering (left/right)
        compare_survey_id: Optional secondary survey id (side-by-side exports)
        compare_survey_name: Optional secondary survey name
        compare_survey_link: Optional secondary survey link
        compare_survey_base_language: Optional secondary survey base language

    Returns:
        ExportContent instance with all prepared data
    """
    """Prepare all content for export (format-agnostic).

    This function extracts and prepares all survey data needed for rendering,
    independent of the output format (DOCX, PDF, etc.). It handles:
    - Translation context building
    - Active QID extraction
    - Survey link generation
    - Mermaid diagram generation
    - QID to JS mapping

    Returns an ExportContent dataclass with all prepared data.
    """
    result = survey_payload.get("result", {}) or {}
    questions = result.get("Questions", {}) or {}
    survey_name = str(result.get("SurveyName") or "").strip()

    base_language = get_base_language(survey_payload) or _DEFAULT_BASE_LANGUAGE
    base_map = build_translation_map_from_cache(
        survey_payload,
        language=base_language,
        base_language=base_language,
    )
    survey_title = (
        str(result.get("SurveyTitle") or base_map.get("SurveyTitle") or "").strip()
        or None
    )
    survey_description = (
        str(
            result.get("SurveyDescription") or base_map.get("SurveyDescription") or ""
        ).strip()
        or None
    )
    version_meta = _load_latest_published_version_from_push_log(survey_id)

    # Load QID to JS mapping
    qid_to_js: dict[str, str] = {}
    if mapping_path is not None:
        qid_to_js = _load_qid_js_mapping(mapping_path, survey_id=survey_id)

    # Extract active QIDs
    active_qids = _active_qids_in_flow(result)
    active_qids = _apply_active_qid_filters(
        result=result,
        active_qids=active_qids,
        include_qids=include_qids,
        include_tags=include_tags,
        include_blocks=include_blocks,
    )

    # Build translation context if language specified
    translation_ctx: TranslationRenderContext | None = translation_ctx_override
    render_plan: TranslationRenderPlan | None = render_plan_override
    lang = _normalize_lang_code(render_language)
    if translation_ctx is None and lang:
        if _normalize_lang_code(lang) != _normalize_lang_code(
            base_language
        ) and not _language_present_in_cache(survey_payload, lang):
            raise RuntimeError(
                f"Language {lang} not found in cached survey definition for {survey_id}. "
                "Run `qsync survey pull --survey-id ...` to refresh the cache or enable the language."
            )
        target_map = build_translation_map_from_cache(
            survey_payload,
            language=lang,
            base_language=base_language,
        )
        base_lang = base_language
        base_map_lang = base_map
        render_plan = _build_translation_render_plan(
            survey_id=survey_id,
            base_language=base_lang,
            target_language=lang,
            questions=questions,
            active_qids=active_qids,
            target_map=target_map,
            base_map=base_map_lang,
        )
        translation_ctx = TranslationRenderContext(
            survey_id=survey_id,
            base_language=base_lang,
            target_language=lang,
            target_map=target_map,
            base_map=base_map_lang,
            compare_to_base=bool(compare_to_base),
            plan=render_plan,
        )
    if translation_ctx is not None and render_plan is None:
        render_plan = translation_ctx.plan

    # Build survey link
    survey_link = _build_survey_link(
        survey_id, edf_overrides=edf_overrides, language=render_language
    )

    # Generate Mermaid diagram code
    mermaid_code: str | None = None
    if mermaid_path is not None:
        mermaid_code = build_mermaid_flow(
            survey_id=survey_id, flow=result.get("SurveyFlow") or {}
        )
        if mermaid_code:
            mermaid_path.parent.mkdir(parents=True, exist_ok=True)
            mermaid_path.write_text(mermaid_code, encoding="utf-8")

    return ExportContent(
        survey_id=survey_id,
        survey_name=survey_name,
        survey_title=survey_title,
        survey_description=survey_description,
        version_number=version_meta.get("version_number") if version_meta else None,
        version_id=version_meta.get("version_id") if version_meta else None,
        version_description=version_meta.get("description") if version_meta else None,
        survey_payload=survey_payload,
        survey_link=survey_link or "",
        active_qids=active_qids,
        translation_ctx=translation_ctx,
        render_plan=render_plan,
        qid_to_js=qid_to_js,
        mermaid_code=mermaid_code,
        mermaid_path=mermaid_path,
        mermaid_image_path=mermaid_image_path,
        edf_overrides=edf_overrides,
        include_html_source=include_html_source,
        layout_heuristics=layout_heuristics,
        compare_to_base=bool(compare_to_base),
        render_language=lang or None,
        base_language=base_language,
        output_path=output_path,
        include_js_strings=include_js_strings,
        flow_trace=flow_trace,
        compare_labels=compare_labels,
        compare_survey_id=compare_survey_id,
        compare_survey_name=compare_survey_name,
        compare_survey_link=compare_survey_link,
        compare_survey_base_language=compare_survey_base_language,
    )


def export_survey_to_word(
    survey_id: str,
    output_path: Path | None = None,
    *,
    edf_overrides: dict[str, str] | None = None,
    smart_name: bool = False,
    include_html_source: bool = True,
    layout_heuristics: bool = False,
    render_language: str | None = None,
    compare_to_base: bool = False,
    include_qids: set[str] | None = None,
    include_tags: set[str] | None = None,
    include_blocks: set[str] | None = None,
    refresh: bool = False,
    include_js_strings: bool = True,
    interactive: bool = True,
    skip_preflight: bool = False,
    flow_trace: Callable[[str], None] | None = None,
) -> Path:
    """Export a survey to a Word document for translation validation."""
    if refresh:
        cache, _ = refresh_survey_cache(survey_id)
    else:
        if not skip_preflight:
            _preflight_cache_freshness(survey_id, interactive=interactive)
        cache = load_cached_survey(survey_id)
    root = resolve_root(required=False) or Path.cwd()
    export_dir = resolve_scoped_dir(EXPORT_DIRNAME, root=root)
    export_dir.mkdir(parents=True, exist_ok=True)

    output_path = Path(output_path) if output_path else None
    # Use inventory slug (not Qualtrics SurveyName) for consistent filename format
    from .survey_inventory import load_inventory_record

    try:
        inventory_record = load_inventory_record(survey_id)
        survey_name = str(inventory_record.get("name") or "").strip()
    except Exception:
        survey_name = ""
    base_language = get_base_language(cache.payload) or _DEFAULT_BASE_LANGUAGE
    output_path = _resolve_output_docx_path(
        survey_id=survey_id,
        survey_name=survey_name,
        export_dir=export_dir,
        output_path=output_path,
        smart_name=smart_name,
        render_language=render_language,
        compare_to_base=compare_to_base,
        base_language=base_language,
    )

    mermaid_path = output_path.with_suffix(".flow.mmd")
    mermaid_png_path = output_path.with_name(output_path.stem + ".flow.png")
    # Optional: refresh cached survey definition before exporting (network).
    # This is equivalent to `qsync survey pull --survey-id ...` and should be
    # used intentionally (e.g. smoke verification).

    return export_survey_payload_to_word(
        survey_id,
        cache.payload,
        output_path,
        mermaid_path=mermaid_path,
        mermaid_image_path=mermaid_png_path,
        render_mermaid=True,
        edf_overrides=edf_overrides,
        mapping_path=(
            resolve_scoped_dir("survey_js", root=root) / "survey_qid_js_map.csv"
        ),
        include_html_source=include_html_source,
        layout_heuristics=layout_heuristics,
        render_language=render_language,
        compare_to_base=compare_to_base,
        include_qids=include_qids,
        include_tags=include_tags,
        include_blocks=include_blocks,
        include_js_strings=include_js_strings,
        flow_trace=flow_trace,
    )


def export_survey_to_pdf(
    survey_id: str,
    output_path: Path | None = None,
    *,
    edf_overrides: dict[str, str] | None = None,
    smart_name: bool = False,
    include_html_source: bool = True,
    layout_heuristics: bool = False,
    render_language: str | None = None,
    compare_to_base: bool = False,
    include_qids: set[str] | None = None,
    include_tags: set[str] | None = None,
    include_blocks: set[str] | None = None,
    refresh: bool = False,
    include_js_strings: bool = True,
    interactive: bool = True,
    skip_preflight: bool = False,
    flow_trace: Callable[[str], None] | None = None,
) -> Path:
    """Export a survey to a PDF document for translation validation.

    This is parallel to export_survey_to_word() but generates PDF output.
    PDF provides better HTML rendering fidelity than DOCX.

    Args:
        survey_id: Qualtrics survey ID
        output_path: Optional output path (defaults to export/<SurveyName>__<SurveyID>__<BASE>.pdf)
        edf_overrides: EDF scenario filter overrides
        smart_name: Use survey name + timestamp in filename
        include_html_source: Include HTML source in output (not applicable to PDF)
        layout_heuristics: Enable layout heuristics
        render_language: Target translation language (e.g., "FR", "NL")
        compare_to_base: Bilingual mode (show base + target side-by-side)
        refresh: Refresh cached survey definition from Qualtrics before exporting

    Returns:
        Path to generated PDF file
    """
    if refresh:
        cache, _ = refresh_survey_cache(survey_id)
    else:
        if not skip_preflight:
            _preflight_cache_freshness(survey_id, interactive=interactive)
        cache = load_cached_survey(survey_id)
    root = resolve_root(required=False) or Path.cwd()
    export_dir = resolve_scoped_dir(EXPORT_DIRNAME, root=root)
    export_dir.mkdir(parents=True, exist_ok=True)

    output_path = Path(output_path) if output_path else None
    # Use inventory slug (not Qualtrics SurveyName) for consistent filename format
    from .survey_inventory import load_inventory_record

    try:
        inventory_record = load_inventory_record(survey_id)
        survey_name = str(inventory_record.get("name") or "").strip()
    except Exception:
        survey_name = ""
    base_language = get_base_language(cache.payload) or _DEFAULT_BASE_LANGUAGE
    output_path = _resolve_output_docx_path(
        survey_id=survey_id,
        survey_name=survey_name,
        export_dir=export_dir,
        output_path=output_path,
        smart_name=smart_name,
        render_language=render_language,
        compare_to_base=compare_to_base,
        base_language=base_language,
        format="pdf",
    )

    mermaid_path = output_path.with_suffix(".flow.mmd")
    mermaid_png_path = output_path.with_name(output_path.stem + ".flow.png")

    # Cached survey definition already refreshed above when requested.

    return export_survey_payload_to_pdf(
        survey_id,
        cache.payload,
        output_path,
        mermaid_path=mermaid_path,
        mermaid_image_path=mermaid_png_path,
        edf_overrides=edf_overrides,
        mapping_path=(
            resolve_scoped_dir("survey_js", root=root) / "survey_qid_js_map.csv"
        ),
        include_html_source=False,  # PDF renders HTML natively, no source needed
        layout_heuristics=layout_heuristics,
        render_language=render_language,
        compare_to_base=compare_to_base,
        include_qids=include_qids,
        include_tags=include_tags,
        include_blocks=include_blocks,
        include_js_strings=include_js_strings,
        flow_trace=flow_trace,
    )


def export_survey_payload_to_pdf(
    survey_id: str,
    survey_payload: dict,
    output_path: Path,
    *,
    mermaid_path: Path | None = None,
    mermaid_image_path: Path | None = None,
    edf_overrides: dict[str, str] | None = None,
    mapping_path: Path | None = None,
    include_html_source: bool = True,
    layout_heuristics: bool = False,
    render_language: str | None = None,
    compare_to_base: bool = False,
    include_qids: set[str] | None = None,
    include_tags: set[str] | None = None,
    include_blocks: set[str] | None = None,
    include_js_strings: bool = True,
    flow_trace: Callable[[str], None] | None = None,
) -> Path:
    """Export a survey payload (already loaded) to PDF.

    This is a testing-friendly entry point that avoids any API/network calls.
    Parallel to export_survey_payload_to_word() but generates PDF output.

    Note: PDF exports render HTML natively, so HTML source is never included
    regardless of the include_html_source parameter (kept for API compatibility).

    Args:
        survey_id: Qualtrics survey ID
        survey_payload: Pre-loaded survey payload
        output_path: Output PDF file path
        mermaid_path: Optional path to save Mermaid source
        mermaid_image_path: Optional path to Mermaid PNG image
        edf_overrides: EDF scenario filter overrides
        mapping_path: Path to QID-to-JS mapping CSV
        include_html_source: Ignored for PDF (HTML is natively rendered, not shown as source)
        layout_heuristics: Enable layout heuristics
        render_language: Target translation language
        compare_to_base: Bilingual mode

    Returns:
        Path to generated PDF file
    """
    output_path = Path(output_path)

    # Prepare all export content (format-agnostic)
    # Note: PDF never includes HTML source (HTML is natively rendered)
    content = _prepare_export_content(
        survey_id=survey_id,
        survey_payload=survey_payload,
        output_path=output_path,
        mermaid_path=mermaid_path,
        mermaid_image_path=mermaid_image_path,
        edf_overrides=edf_overrides,
        mapping_path=mapping_path,
        include_html_source=False,  # PDF renders HTML natively, no source needed
        layout_heuristics=layout_heuristics,
        render_language=render_language,
        compare_to_base=compare_to_base,
        include_qids=include_qids,
        include_tags=include_tags,
        include_blocks=include_blocks,
        include_js_strings=include_js_strings,
        flow_trace=flow_trace,
    )

    # Render to PDF
    return _render_to_pdf(content)


def _render_to_docx(content: ExportContent, *, render_mermaid: bool = False) -> Path:
    """Render prepared export content to a DOCX file.

    Args:
        content: Prepared export content
        render_mermaid: Whether to render Mermaid diagram to PNG

    Returns:
        Path to the saved DOCX file
    """
    from docx import Document

    content.output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _apply_base_body_style(doc)
    _add_doc_header(
        doc,
        survey_id=content.survey_id,
        survey_name=content.survey_name,
        survey_title=content.survey_title,
        survey_description=content.survey_description,
        version_number=content.version_number,
        version_id=content.version_id,
        version_description=content.version_description,
        edf_overrides=content.edf_overrides,
        survey_link=content.survey_link,
        base_language=content.base_language,
        render_language=content.render_language,
        compare_to_base=content.compare_to_base,
        compare_labels=content.compare_labels,
        compare_survey_id=content.compare_survey_id,
        compare_survey_name=content.compare_survey_name,
        compare_survey_link=content.compare_survey_link,
        compare_survey_base_language=content.compare_survey_base_language,
    )

    # Extract data from content for rendering
    result = content.survey_payload.get("result", {}) or {}
    blocks = result.get("Blocks", {}) or {}
    questions = result.get("Questions", {}) or {}

    # Add translation summary if language specified
    if content.render_plan is not None:
        _add_translation_rendering_summary(
            doc,
            plan=content.render_plan,
            compare_to_base=content.compare_to_base,
            compare_labels=content.compare_labels,
        )

    _add_coverage_summary(doc, questions=questions, active_qids=content.active_qids)
    _add_question_type_legend(doc, questions=questions, active_qids=content.active_qids)

    if content.edf_overrides:
        _warn_about_unused_edf_overrides(
            doc,
            flow_obj=result.get("SurveyFlow") or {},
            edf_overrides=content.edf_overrides,
        )

    _add_mermaid_section_and_file(
        doc,
        survey_id=content.survey_id,
        result=result,
        mermaid_path=content.mermaid_path,
        mermaid_image_path=content.mermaid_image_path,
        render_mermaid=render_mermaid,
    )
    _add_survey_content_section(
        doc,
        result=result,
        blocks=blocks,
        questions=questions,
        qid_to_js=content.qid_to_js,
        active_qids=content.active_qids,
        edf_overrides=content.edf_overrides,
        include_html_source=content.include_html_source,
        layout_heuristics=content.layout_heuristics,
        base_language=content.base_language,
        render_language=content.render_language,
        compare_to_base=content.compare_to_base,
        translation_ctx=content.translation_ctx,
        include_js_strings=content.include_js_strings,
        flow_trace=content.flow_trace,
        compare_labels=content.compare_labels,
    )
    _add_external_translation_surfaces_section(
        doc,
        survey_id=content.survey_id,
        result=result,
        questions=questions,
        qid_to_js=content.qid_to_js,
    )

    doc.save(str(content.output_path))
    return content.output_path


def export_survey_payload_to_word(
    survey_id: str,
    survey_payload: dict,
    output_path: Path,
    *,
    mermaid_path: Path | None = None,
    mermaid_image_path: Path | None = None,
    render_mermaid: bool = False,
    edf_overrides: dict[str, str] | None = None,
    mapping_path: Path | None = None,
    include_html_source: bool = True,
    layout_heuristics: bool = False,
    render_language: str | None = None,
    compare_to_base: bool = False,
    include_qids: set[str] | None = None,
    include_tags: set[str] | None = None,
    include_blocks: set[str] | None = None,
    include_js_strings: bool = True,
    flow_trace: Callable[[str], None] | None = None,
) -> Path:
    """Export a survey payload (already loaded) to Word.

    This is a testing-friendly entry point that avoids any API/network calls.
    """
    output_path = Path(output_path)

    # Prepare all export content (format-agnostic)
    content = _prepare_export_content(
        survey_id=survey_id,
        survey_payload=survey_payload,
        output_path=output_path,
        mermaid_path=mermaid_path,
        mermaid_image_path=mermaid_image_path,
        edf_overrides=edf_overrides,
        mapping_path=mapping_path,
        include_html_source=include_html_source,
        layout_heuristics=layout_heuristics,
        render_language=render_language,
        compare_to_base=compare_to_base,
        include_qids=include_qids,
        include_tags=include_tags,
        include_blocks=include_blocks,
        include_js_strings=include_js_strings,
        flow_trace=flow_trace,
    )

    # Render to DOCX
    return _render_to_docx(content, render_mermaid=render_mermaid)


def export_surveys_side_by_side_docx(
    survey_a_id: str,
    survey_b_id: str,
    output_path: Path | None = None,
    *,
    label_a: str | None = None,
    label_b: str | None = None,
    smart_name: bool = False,
    refresh: bool = False,
    include_html_source: bool = True,
    layout_heuristics: bool = False,
    include_js_strings: bool = True,
    interactive: bool = True,
    flow_trace: Callable[[str], None] | None = None,
) -> Path:
    """Export two surveys side-by-side (Survey A vs Survey B) to a single DOCX."""

    if refresh:
        cache_a, _ = refresh_survey_cache(survey_a_id)
        cache_b, _ = refresh_survey_cache(survey_b_id)
    else:
        _preflight_cache_freshness(survey_a_id, interactive=interactive)
        _preflight_cache_freshness(survey_b_id, interactive=interactive)
        cache_a = load_cached_survey(survey_a_id)
        cache_b = load_cached_survey(survey_b_id)

    payload_a = cache_a.payload
    payload_b = cache_b.payload
    result_a = _survey_result(payload_a)
    result_b = _survey_result(payload_b)

    survey_name_a = str(result_a.get("SurveyName") or "").strip()
    survey_name_b = str(result_b.get("SurveyName") or "").strip()

    base_lang_a = (
        _normalize_lang_code(get_base_language(payload_a)) or _DEFAULT_BASE_LANGUAGE
    )
    base_lang_b = (
        _normalize_lang_code(get_base_language(payload_b)) or _DEFAULT_BASE_LANGUAGE
    )

    label_a_clean = _normalize_label(label_a)
    label_b_clean = _normalize_label(label_b)
    if label_a_clean or label_b_clean:
        compare_labels = (
            label_a_clean or base_lang_a,
            label_b_clean or base_lang_b,
        )
    else:
        if base_lang_a != base_lang_b:
            compare_labels = (base_lang_a, base_lang_b)
        else:
            compare_labels = ("Survey A", "Survey B")

    root = resolve_root(required=False) or Path.cwd()
    export_dir = resolve_scoped_dir(EXPORT_DIRNAME, root=root)
    export_dir.mkdir(parents=True, exist_ok=True)

    output_path = _resolve_output_side_by_side_docx_path(
        survey_a_id=survey_a_id,
        survey_a_name=survey_name_a,
        survey_b_id=survey_b_id,
        survey_b_name=survey_name_b,
        export_dir=export_dir,
        output_path=output_path,
        smart_name=smart_name,
    )

    base_map = build_translation_map_from_cache(
        payload_a,
        language=base_lang_a,
        base_language=base_lang_a,
    )
    target_map = build_translation_map_from_cache(
        payload_b,
        language=base_lang_b,
        base_language=base_lang_b,
    )

    active_qids = _active_qids_in_flow(result_a)
    questions = result_a.get("Questions", {}) or {}

    render_plan = _build_translation_render_plan(
        survey_id=survey_a_id,
        base_language=base_lang_a,
        target_language=base_lang_b,
        questions=questions,
        active_qids=active_qids,
        target_map=target_map,
        base_map=base_map,
    )
    translation_ctx = TranslationRenderContext(
        survey_id=survey_a_id,
        base_language=base_lang_a,
        target_language=base_lang_b,
        target_map=target_map,
        base_map=base_map,
        compare_to_base=True,
        plan=render_plan,
    )

    compare_link = _build_survey_link(
        survey_b_id, edf_overrides=None, language=base_lang_b
    )

    content = _prepare_export_content(
        survey_id=survey_a_id,
        survey_payload=payload_a,
        output_path=output_path,
        include_html_source=include_html_source,
        layout_heuristics=layout_heuristics,
        render_language=base_lang_b,
        compare_to_base=True,
        include_js_strings=include_js_strings,
        flow_trace=flow_trace,
        translation_ctx_override=translation_ctx,
        render_plan_override=render_plan,
        compare_labels=compare_labels,
        compare_survey_id=survey_b_id,
        compare_survey_name=survey_name_b,
        compare_survey_link=compare_link,
        compare_survey_base_language=base_lang_b,
    )

    return _render_to_docx(content)


# ----------------------------
# Document building
# ----------------------------


def _add_doc_header(
    doc,
    *,
    survey_id: str,
    survey_name: str,
    survey_title: str | None,
    survey_description: str | None,
    version_number: int | None,
    version_id: str | None,
    version_description: str | None,
    edf_overrides: dict[str, str] | None,
    survey_link: str | None,
    base_language: str,
    render_language: str | None,
    compare_to_base: bool,
    compare_labels: tuple[str, str] | None = None,
    compare_survey_id: str | None = None,
    compare_survey_name: str | None = None,
    compare_survey_link: str | None = None,
    compare_survey_base_language: str | None = None,
) -> None:
    doc.add_heading("SURVEY TRANSLATION EXPORT", level=0)

    table = doc.add_table(rows=1, cols=2)
    first_row_used = False

    def _add_row(
        label: str, value: str | None, *, hyperlink: str | None = None
    ) -> None:
        nonlocal first_row_used
        if not first_row_used:
            row = table.rows[0]
            first_row_used = True
        else:
            row = table.add_row()
        row.cells[0].text = label
        if hyperlink:
            p = row.cells[1].paragraphs[0]
            _append_hyperlink(
                p, url=hyperlink, text=hyperlink, bold=False, italic=False, size_pt=None
            )
            return
        row.cells[1].text = str(value or "").strip() or "—"

    _add_row("SurveyID:", survey_id)
    _add_row("Survey name:", survey_name)
    if survey_title:
        _add_row("Survey title:", survey_title)
    if survey_description:
        _add_row("Survey description:", survey_description)
    if version_number is not None or version_description:
        version_label = "Version:"
        version_value = None
        if version_number is not None and version_id:
            version_value = f"{version_number} (id={version_id})"
        elif version_number is not None:
            version_value = str(version_number)
        elif version_id:
            version_value = str(version_id)
        _add_row(version_label, version_value)
        if version_description:
            _add_row("Version description:", version_description)
    _add_row("Generated:", datetime.now().isoformat(timespec="seconds"))

    if compare_survey_id:
        _add_row("Compare SurveyID:", compare_survey_id)
    if compare_survey_name:
        _add_row("Compare survey name:", compare_survey_name)
    if compare_survey_base_language:
        _add_row("Compare base language:", compare_survey_base_language)
    if compare_survey_link:
        _add_row("Compare survey link:", None, hyperlink=compare_survey_link)

    base = _normalize_lang_code(base_language) or _DEFAULT_BASE_LANGUAGE
    target = _normalize_lang_code(render_language)
    if target:
        base_label, target_label = _resolve_compare_labels(
            compare_labels,
            base_language=base,
            target_language=target,
        )
        mode = f"{base_label}-{target_label}" if compare_to_base else target_label
        _add_row("Render language:", f"{mode} (base={base})")
    if edf_overrides:
        joined = ", ".join([f"{k}={v}" for k, v in sorted(edf_overrides.items())])
        _add_row("Scenario EDF filters:", joined)
    if survey_link:
        _add_row("Survey link:", None, hyperlink=survey_link)


def _collect_expected_translation_keys(
    questions: dict, active_qids: set[str]
) -> list[str]:
    keys: set[str] = set()
    for qid in sorted(active_qids or set(), key=str):
        q = questions.get(qid) or {}
        if is_system_question_type(q.get("QuestionType")):
            continue

        qid_s = str(qid)
        keys.add(f"{qid_s}_QuestionText")

        qtype = str(q.get("QuestionType") or "").strip()
        selector = str(q.get("Selector") or "").strip()
        is_sbs_matrix = qtype == "SBS" and selector == "SBSMatrix"

        choices = q.get("Choices") or {}
        if isinstance(choices, dict):
            for cid in choices.keys():
                keys.add(f"{qid_s}_Choice{cid}")

        if is_sbs_matrix:
            additional = q.get("AdditionalQuestions") or {}
            if isinstance(additional, dict):
                top_answers = q.get("Answers") or {}
                if not isinstance(top_answers, dict):
                    top_answers = {}
                for column_id, column in additional.items():
                    column_id_s = str(column_id)
                    keys.add(f"{qid_s}#{column_id_s}_QuestionText")
                    if not isinstance(column, dict):
                        continue
                    answers = column.get("Answers") or {}
                    if not isinstance(answers, dict):
                        answers = {}
                    if not answers:
                        answers = top_answers
                    answer_order = column.get("AnswerOrder")
                    for ans_id in _ordered_numeric_string_ids(answers, answer_order):
                        keys.add(f"{qid_s}#{column_id_s}_Answer{ans_id}")
        else:
            answers = q.get("Answers") or {}
            if isinstance(answers, dict):
                for aid in answers.keys():
                    keys.add(f"{qid_s}_Answer{aid}")

        labels = q.get("Labels") or {}
        if isinstance(labels, dict):
            for lid in labels.keys():
                keys.add(f"{qid_s}_Label{lid}")

    return sorted(keys)


def _build_translation_render_plan(
    *,
    survey_id: str,
    base_language: str,
    target_language: str,
    questions: dict,
    active_qids: set[str],
    target_map: Mapping[str, Any],
    base_map: Mapping[str, Any] | None,
) -> TranslationRenderPlan:
    expected = _collect_expected_translation_keys(questions, active_qids)

    base_empty: set[str] = set()
    if base_map is not None:
        for k, v in base_map.items():
            if not isinstance(v, str) or not v.strip():
                base_empty.add(str(k))

    missing: list[str] = []
    empty_nonempty_base: list[str] = []
    for key in expected:
        if key not in target_map:
            missing.append(key)
            continue
        value = target_map.get(key)
        if isinstance(value, str) and value.strip():
            continue
        if key in base_empty:
            continue
        empty_nonempty_base.append(key)

    return TranslationRenderPlan(
        survey_id=survey_id,
        base_language=_normalize_lang_code(base_language) or _DEFAULT_BASE_LANGUAGE,
        target_language=_normalize_lang_code(target_language),
        expected_keys=expected,
        base_empty_keys=base_empty,
        missing_keys=missing,
        empty_but_base_nonempty_keys=empty_nonempty_base,
    )


def _add_translation_rendering_summary(
    doc,
    *,
    plan: TranslationRenderPlan,
    compare_to_base: bool,
    compare_labels: tuple[str, str] | None = None,
    max_samples: int = 20,
) -> None:
    doc.add_heading("LANGUAGE RENDERING SUMMARY", level=1)

    base_label, target_label = _resolve_compare_labels(
        compare_labels,
        base_language=plan.base_language,
        target_language=plan.target_language,
    )
    mode = f"{base_label}-{target_label}" if compare_to_base else target_label
    doc.add_paragraph(f"Mode: {mode} (base={plan.base_language})")
    doc.add_paragraph(f"Expected keys rendered by export: {plan.total_expected}")
    doc.add_paragraph(f"OK (translated or allowed-empty): {plan.total_ok}")
    if plan.total_missing:
        doc.add_paragraph(f"Missing keys (not present in map): {plan.total_missing}")
    if plan.total_empty_but_base_nonempty:
        doc.add_paragraph(
            f"Empty keys (need translation): {plan.total_empty_but_base_nonempty}"
        )

    samples: list[str] = []
    for k in plan.missing_keys[: max_samples // 2]:
        samples.append(f"- MISSING: {k}")
    for k in plan.empty_but_base_nonempty_keys[: max_samples - len(samples)]:
        samples.append(f"- EMPTY: {k}")
    if samples:
        doc.add_paragraph("Sample issues (fix these first):")
        for line in samples:
            doc.add_paragraph(line)

    _add_doc_spacer_paragraph(doc, depth=0)


def _add_coverage_summary(doc, *, questions: dict, active_qids: set[str]) -> None:
    doc.add_heading("COVERAGE SUMMARY", level=1)
    total = len(questions or {})
    active = len(active_qids or set())
    excluded = max(0, total - active)
    doc.add_paragraph(f"Total questions in JSON: {total}")
    doc.add_paragraph(f"Active & exported (in-flow, non-Trash): {active}")
    doc.add_paragraph(f"Excluded (unplaced/Trash/other): {excluded}")


def _add_question_type_legend(doc, *, questions: dict, active_qids: set[str]) -> None:
    """Add a short legend for the question type abbreviations used in metadata lines."""

    used: dict[str, str] = {}
    counts: dict[str, int] = {}
    for qid in sorted(active_qids or set(), key=str):
        q = questions.get(qid) or {}
        qt = (q.get("QuestionType") or "").strip()
        sel = (q.get("Selector") or "").strip()
        sub = (q.get("SubSelector") or "").strip()
        abbrev, label = _question_type_abbrev_and_label(
            question_type=qt, selector=sel, subselector=sub
        )
        if abbrev and label:
            used.setdefault(abbrev, label)
            counts[abbrev] = counts.get(abbrev, 0) + 1

    if not used:
        return

    doc.add_heading("QUESTION TYPE LEGEND", level=1)
    table = doc.add_table(rows=1, cols=3)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _set_table_column_widths_in(table, [1.0, 4.5, 1.0])
    hdr = table.rows[0].cells
    hdr[0].text = "Code"
    hdr[1].text = "Meaning"
    hdr[2].text = "Count"
    try:
        for cell in hdr:
            for p in getattr(cell, "paragraphs", []) or []:
                for r in getattr(p, "runs", []) or []:
                    r.bold = True
    except Exception:
        pass

    for code in sorted(used.keys(), key=str):
        row = table.add_row().cells
        row[0].text = code
        try:
            for p in getattr(row[0], "paragraphs", []) or []:
                for r in getattr(p, "runs", []) or []:
                    _style_id_run(r)
        except Exception:
            pass
        row[1].text = used[code]
        row[2].text = str(counts.get(code, 0))

    _shrink_table_font(table, size_pt=9)
    _add_doc_spacer_paragraph(doc, depth=0)


def _add_mermaid_section_and_file(
    doc,
    *,
    survey_id: str,
    result: dict,
    mermaid_path: Path | None,
    mermaid_image_path: Path | None,
    render_mermaid: bool,
) -> None:
    doc.add_heading("FLOW DIAGRAM (Mermaid)", level=1)
    code = build_mermaid_flow(survey_id=survey_id, flow=result.get("SurveyFlow") or {})
    should_create_mmd = not _env_flag_disabled("QSYNC_MERMAID_RENDER")
    if mermaid_path is not None and should_create_mmd:
        mermaid_path = Path(mermaid_path)
        mermaid_path.parent.mkdir(parents=True, exist_ok=True)
        mermaid_path.write_text(code + "\n", encoding="utf-8")
        doc.add_paragraph(f"Mermaid source file: {mermaid_path}")

    should_render = render_mermaid and not _env_flag_disabled("QSYNC_MERMAID_RENDER")
    if should_render:
        if mermaid_image_path is None:
            mermaid_image_path = Path("mermaid.flow.png")
        mermaid_image_path = Path(mermaid_image_path)
        mermaid_image_path.parent.mkdir(parents=True, exist_ok=True)
        _render_mermaid_to_png(code, mermaid_image_path)
        try:
            from docx.shared import Inches

            doc.add_picture(str(mermaid_image_path), width=Inches(6.5))
        except Exception:
            doc.add_paragraph(f"(Rendered Mermaid image: {mermaid_image_path})")
    else:
        doc.add_paragraph(
            "(Mermaid rendering disabled/unavailable for this run; see the .flow.mmd file.)"
        )
    # Intentionally do not embed Mermaid source code in the Word document.


def _add_embedded_data_section(doc, *, survey_id: str, survey_payload: dict) -> None:
    doc.add_heading("EMBEDDED DATA (Summary)", level=1)
    doc.add_paragraph(
        "Note: Embedded Data is often set conditionally in SurveyFlow; this section is grouped by field to reduce noise."
    )

    rows = build_embedded_data_rows(survey_id, survey_payload)
    if not rows:
        doc.add_paragraph("(No embedded data found.)")
        return

    # Translator-friendly summary: group by field and avoid listing every conditional write.
    by_field: Dict[str, List[object]] = {}
    for r in rows:
        by_field.setdefault(r.field, []).append(r)

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Values (examples)"
    hdr[2].text = "Set in flow (#)"
    hdr[3].text = "JS writers (QIDs)"
    hdr[4].text = "Notes"

    for field in sorted(by_field.keys(), key=str):
        field_rows = by_field.get(field) or []
        values: List[str] = []
        types: set[str] = set()
        js_writers: set[str] = set()
        set_in_flow = 0
        for r in field_rows:
            types.add(getattr(r, "ed_type", "") or "")
            if getattr(r, "flow_order", 0) and getattr(r, "flow_order", 0) > 0:
                set_in_flow += 1
            v = getattr(r, "value", None)
            if v is not None:
                sv = str(v).strip()
                if sv and sv not in values:
                    values.append(sv)
            writers = getattr(r, "written_by_qids", "") or ""
            for qid in [x.strip() for x in writers.split(",") if x.strip()]:
                js_writers.add(qid)

        if not values:
            values_str = EMBEDDED_EMPTY_VALUE
        else:
            values_str = "; ".join(values)

        if js_writers:
            writers_sorted = sorted(js_writers, key=str)
            writers_str = ", ".join(writers_sorted)
        else:
            writers_str = ""

        note_bits: List[str] = []
        if types:
            note_bits.append("Types=" + ",".join(sorted([t for t in types if t])))
        if _looks_like_routing_field(field):
            note_bits.append("routing/label field (heuristic)")
        note = "; ".join(note_bits)

        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = values_str
        row_cells[2].text = str(set_in_flow)
        row_cells[3].text = writers_str
        row_cells[4].text = note
        _style_edf_field_cell(row_cells[0])

    _shrink_table_font(table, size_pt=8)
    _add_doc_spacer_paragraph(doc, depth=0)


def _add_survey_content_section(
    doc,
    *,
    result: dict,
    blocks: dict,
    questions: dict,
    qid_to_js: dict[str, str],
    active_qids: set[str],
    edf_overrides: dict[str, str] | None,
    include_html_source: bool,
    layout_heuristics: bool,
    base_language: str,
    render_language: str | None,
    compare_to_base: bool,
    translation_ctx: TranslationRenderContext | None,
    include_js_strings: bool,
    flow_trace: Callable[[str], None] | None,
    compare_labels: tuple[str, str] | None,
) -> None:
    doc.add_heading("SURVEY CONTENT (Flow Order)", level=1)
    flow = result.get("SurveyFlow") or {}
    flow_list = flow.get("Flow") or []
    if not isinstance(flow_list, list):
        doc.add_paragraph("(SurveyFlow missing or malformed.)")
        return
    if not active_qids:
        doc.add_paragraph(
            "(No active questions detected in SurveyFlow non-Trash blocks.)"
        )
        return

    asked_qids: set[str] | None = set() if edf_overrides else None

    _traverse_flow(
        doc,
        flow_list=flow_list,
        blocks=blocks,
        questions=questions,
        qid_to_js=qid_to_js,
        active_qids=active_qids,
        edf_overrides=edf_overrides,
        include_html_source=include_html_source,
        layout_heuristics=layout_heuristics,
        asked_qids=asked_qids,
        base_language=base_language,
        render_language=render_language,
        compare_to_base=compare_to_base,
        translation_ctx=translation_ctx,
        include_js_strings=include_js_strings,
        flow_trace=flow_trace,
        compare_labels=compare_labels,
        depth=0,
    )


def _add_external_translation_surfaces_section(
    doc,
    *,
    survey_id: str,
    result: dict,
    questions: dict,
    qid_to_js: dict[str, str],
) -> None:
    doc.add_heading("EXTERNAL TRANSLATION SURFACES", level=1)

    # 1) Questions with QuestionJS (these often contain user-visible strings)
    doc.add_heading("Questions with QuestionJS", level=2)
    qids_with_js = sorted(
        [qid for qid, q in questions.items() if (q.get("QuestionJS") or "").strip()]
    )
    if not qids_with_js:
        doc.add_paragraph("(No questions with QuestionJS found.)")
    else:
        qid_to_js = dict(qid_to_js or {})
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "QID"
        hdr[1].text = "DataExportTag"
        hdr[2].text = "survey_js source"
        for qid in qids_with_js:
            q = questions.get(qid) or {}
            tag = (q.get("DataExportTag") or "").strip()
            js_file = qid_to_js.get(qid) or ""
            row_cells = table.add_row().cells
            row_cells[0].text = qid
            row_cells[1].text = tag
            row_cells[2].text = f"survey_js/core/{js_file}" if js_file else ""
            # Make QIDs monospace (ID styling) without changing other formatting.
            try:
                for p in row_cells[0].paragraphs:
                    for r in p.runs:
                        _style_qid_run(r)
            except Exception:
                pass

        _add_doc_spacer_paragraph(doc, depth=0)
        doc.add_paragraph(
            "Note: translate user-visible strings in these JS files under survey_js/ (ground truth)."
        )

    # 2) EndSurvey message references (message text lives in the Qualtrics message library)
    doc.add_heading("EndSurvey message library references", level=2)
    refs = _extract_end_survey_message_refs(result.get("SurveyFlow") or {})
    if not refs:
        doc.add_paragraph("(No EndSurvey DisplayMessage references found.)")
        return
    for ref in refs:
        have_local = (
            _read_eos_message_from_disk(ref.library_id, ref.message_id) is not None
        )
        _add_system_note(
            doc,
            f"- FlowID {ref.flow_id or ''}: EOSMessageLibrary={ref.library_id}, EOSMessage={ref.message_id}"
            + (" (local message found)" if have_local else " (not pulled to disk)"),
            depth=0,
        )
    doc.add_paragraph(
        "Note: message text is not present in cached survey JSON. If you pulled EOS messages to disk "
        "(qsync eos pull), this export will embed the message content inline at EndSurvey nodes."
    )


# ----------------------------
# Flow traversal + formatting
# ----------------------------


def _traverse_flow(
    doc,
    *,
    flow_list: list,
    blocks: dict,
    questions: dict,
    qid_to_js: dict[str, str],
    active_qids: set[str],
    edf_overrides: dict[str, str] | None,
    include_html_source: bool,
    layout_heuristics: bool,
    asked_qids: set[str] | None,
    base_language: str,
    render_language: str | None,
    compare_to_base: bool,
    translation_ctx: TranslationRenderContext | None,
    include_js_strings: bool,
    flow_trace: Callable[[str], None] | None,
    compare_labels: tuple[str, str] | None,
    depth: int,
) -> None:
    """Render SurveyFlow using the shared flow_traversal helper."""

    def on_block(node: dict, depth_level: int) -> None:
        _add_block(
            doc,
            block_id=str(node["ID"]),
            blocks=blocks,
            questions=questions,
            qid_to_js=qid_to_js,
            active_qids=active_qids,
            include_html_source=include_html_source,
            edf_overrides=edf_overrides,
            asked_qids=asked_qids,
            layout_heuristics=layout_heuristics,
            base_language=base_language,
            render_language=render_language,
            compare_to_base=compare_to_base,
            translation_ctx=translation_ctx,
            include_js_strings=include_js_strings,
            depth=depth_level,
            flow_trace=flow_trace,
            compare_labels=compare_labels,
        )

    def on_group(node: dict, depth_level: int) -> None:
        _add_system_note(
            doc, f"GROUP: {node.get('Description') or ''}".strip(), depth=depth_level
        )

    def on_embedded_data(node: dict, depth_level: int) -> None:
        _render_embedded_data_node(
            doc, node=node, depth=depth_level, edf_overrides=edf_overrides
        )

    def on_web_service(node: dict, depth_level: int) -> None:
        _render_web_service_node(doc, node=node, depth=depth_level)

    def on_randomizer(node: dict, depth_level: int) -> None:
        _add_system_note(doc, _format_randomizer(node), depth=depth_level)

    def on_branch_open(node: dict, depth_level: int) -> None:
        cond = _format_logic_blob(
            node.get("BranchLogic"),
            questions=questions,
            translation_ctx=translation_ctx,
        )
        _add_logic_line(doc, f"BRANCH: IF {cond}".strip(), depth=depth_level)

    def on_branch_decision(
        node: dict, decision: bool, reason: str, depth_level: int
    ) -> None:
        if not flow_trace:
            return
        flow_id = str(node.get("FlowID") or "").strip()
        cond = _format_logic_blob(
            node.get("BranchLogic"),
            questions=questions,
            translation_ctx=translation_ctx,
        )
        taken = "THEN" if decision else "ELSE"
        label = f"FlowID={flow_id}" if flow_id else "FlowID=?"
        flow_trace(f"[branch:{reason}] {label} -> {taken} | {cond}")

    def on_branch_then(_node: dict, depth_level: int) -> None:
        _add_logic_line(doc, "THEN:", depth=depth_level)

    def on_branch_else(_node: dict, depth_level: int) -> None:
        _add_logic_line(doc, "ELSE:", depth=depth_level)

    def on_branch_end(_node: dict, depth_level: int) -> None:
        _add_logic_line(doc, "END BRANCH", depth=depth_level)

    def on_end_survey(node: dict, depth_level: int) -> None:
        opts = node.get("Options") or {}
        term = str(opts.get("SurveyTermination") or "").strip()
        flow_id = str(node.get("FlowID") or "").strip()
        lib_id = str(opts.get("EOSMessageLibrary") or "").strip()
        msg_id = str(opts.get("EOSMessage") or "").strip()

        label = f"END SURVEY: {term}" if term else "END SURVEY"
        details: list[str] = []
        if flow_id:
            details.append(f"FlowID={flow_id}")
        if lib_id:
            details.append(f"EOSMessageLibrary={lib_id}")
        if msg_id:
            details.append(f"EOSMessage={msg_id}")
        if details:
            label = f"{label} ({', '.join(details)})"

        _add_system_note(doc, label, depth=depth_level)
        if term == "DisplayMessage" and lib_id and msg_id:
            _render_eos_message_content(
                doc,
                library_id=lib_id,
                message_id=msg_id,
                flow_id=flow_id,
                depth=depth_level + 1,
                include_html_source=include_html_source,
                layout_heuristics=layout_heuristics,
                base_language=base_language,
                render_language=render_language,
                compare_to_base=compare_to_base,
            )

    def on_unknown(node: dict, depth_level: int) -> None:
        node_type = str(node.get("Type") or "").strip()
        if node_type:
            _add_system_note(doc, f"FLOW NODE: {node_type}", depth=depth_level)

    handlers = FlowTraversalHandlers(
        on_block=on_block,
        on_group=on_group,
        on_embedded_data=on_embedded_data,
        on_web_service=on_web_service,
        on_randomizer=on_randomizer,
        on_branch_decision=on_branch_decision,
        on_branch_open=on_branch_open,
        on_branch_then=on_branch_then,
        on_branch_else=on_branch_else,
        on_branch_end=on_branch_end,
        on_end_survey=on_end_survey,
        on_unknown=on_unknown,
    )

    walk_flow(
        flow_list=flow_list,
        handlers=handlers,
        edf_overrides=edf_overrides,
        asked_qids=asked_qids,
        depth=depth,
        eval_branch=_eval_boolean_expression,
        eval_branch_with_asked=_eval_boolean_expression_with_unasked_selected_false,
    )


def _add_block(
    doc,
    *,
    block_id: str,
    blocks: dict,
    questions: dict,
    qid_to_js: dict[str, str],
    active_qids: set[str],
    include_html_source: bool,
    edf_overrides: dict[str, str] | None,
    asked_qids: set[str] | None,
    layout_heuristics: bool,
    base_language: str,
    render_language: str | None,
    compare_to_base: bool,
    translation_ctx: TranslationRenderContext | None,
    include_js_strings: bool,
    flow_trace: Callable[[str], None] | None,
    compare_labels: tuple[str, str] | None,
    depth: int,
) -> None:
    block = blocks.get(block_id) or {}
    if (block.get("Type") or "").strip() == "Trash":
        return
    block_name = (block.get("Description") or "").strip()

    # Determine which questions actually render in this block (so we can omit
    # empty blocks in scenario exports, e.g. when all questions are hidden by
    # DisplayLogic under the provided EDF overrides).
    block_qids_in_order: list[str] = []
    for elem in block.get("BlockElements", []) or []:
        if (elem.get("Type") or "") != "Question":
            continue
        qid = elem.get("QuestionID")
        if not qid or qid not in questions:
            continue
        if qid not in active_qids:
            continue
        block_qids_in_order.append(str(qid))

    if not block_qids_in_order:
        return

    render_qids: list[str] = list(block_qids_in_order)
    if edf_overrides and asked_qids is not None:
        asked_sim = set(asked_qids)
        render_qids = []
        for qid in block_qids_in_order:
            visible = _eval_question_display_logic_visibility(
                questions.get(qid) or {},
                questions=questions,
                edf_overrides=edf_overrides,
                asked_qids=asked_sim,
            )
            if visible is False:
                if flow_trace:
                    flow_trace(
                        f"[display_logic] Block {block_id} hides QID {qid} (display logic false)"
                    )
                continue
            render_qids.append(qid)
            asked_sim.add(qid)
        if not render_qids:
            if flow_trace:
                flow_trace(
                    f"[block_drop] Block {block_id} dropped (all questions hidden by display logic)"
                )
            return

    # Block headers should have separation before and after (even without page breaks).
    _add_block_header_leading_spacer(doc, depth=depth)
    p = doc.add_paragraph()
    _set_indent(p, depth=depth)
    _shade_block_header_paragraph(p)

    def add_run(text: str, *, is_id: bool = False) -> None:
        r = p.add_run(text)
        r.bold = True
        try:
            from docx.shared import Pt

            r.font.size = Pt(14)
        except Exception:
            pass
        if is_id:
            _style_id_run(r)

    if block_name:
        add_run(f"BLOCK START: {block_name} (")
        add_run(block_id, is_id=True)
        add_run(")")
    else:
        add_run("BLOCK START: ")
        add_run(block_id, is_id=True)

    # Render block elements in order (questions and page breaks)
    for elem in block.get("BlockElements", []) or []:
        if not isinstance(elem, dict):
            continue
        elem_type = elem.get("Type") or ""

        if elem_type == "Question":
            qid = elem.get("QuestionID")
            if qid and qid in render_qids:
                _add_question(
                    doc,
                    qid=qid,
                    question=questions.get(qid) or {},
                    questions=questions,
                    qid_to_js=qid_to_js,
                    include_html_source=include_html_source,
                    layout_heuristics=layout_heuristics,
                    base_language=base_language,
                    render_language=render_language,
                    compare_to_base=compare_to_base,
                    translation_ctx=translation_ctx,
                    include_js_strings=include_js_strings,
                    depth=depth + 1,
                    compare_labels=compare_labels,
                )
                if asked_qids is not None:
                    asked_qids.add(str(qid))

        elif elem_type == "Page Break":
            _add_system_note(doc, "--- PAGE BREAK ---", depth=depth + 1)


def _add_question(
    doc,
    *,
    qid: str,
    question: dict,
    questions: dict,
    qid_to_js: dict[str, str],
    include_html_source: bool,
    layout_heuristics: bool,
    base_language: str,
    render_language: str | None,
    compare_to_base: bool,
    translation_ctx: TranslationRenderContext | None,
    include_js_strings: bool,
    compare_labels: tuple[str, str] | None,
    depth: int,
) -> None:
    tag = (question.get("DataExportTag") or "").strip()
    qtype = (question.get("QuestionType") or "").strip()
    selector = (question.get("Selector") or "").strip()
    subselector = (question.get("SubSelector") or "").strip()
    qt_abbrev, _qt_label = _question_type_abbrev_and_label(
        question_type=qtype, selector=selector, subselector=subselector
    )
    if question.get("Randomization"):
        qt_abbrev = (qt_abbrev + "+R").strip()

    has_js = bool((question.get("QuestionJS") or "").strip()) or bool(
        (qid_to_js or {}).get(qid)
    )

    marker = _question_validation_marker(question)
    lang_ctx = translation_ctx
    target_lang = lang_ctx.target_language if lang_ctx else ""
    base_label, target_label = _resolve_compare_labels(
        compare_labels,
        base_language=base_language,
        target_language=target_lang or render_language,
    )

    bilingual = bool(lang_ctx and compare_to_base)

    # Inside-table structure: one column, stacked rows (only when content exists):
    #   - Metadata (always)
    #   - Display logic (optional)
    #   - Question text (optional)
    #   - Statements (optional)
    #   - Answer options (optional)
    rows_to_render: list[tuple[str, object]] = []
    rows_to_render.append(
        ("meta", (qid, qt_abbrev, has_js, tag, marker)),
    )

    def _resolve_translation(key: str, base_value: str) -> str:
        if not lang_ctx:
            return base_value
        raw = lang_ctx.target_map.get(key)
        if isinstance(raw, str) and raw.strip():
            return raw
        # Allow empty in target if base translation map is also empty for this key.
        if isinstance(raw, str) and not raw.strip():
            if lang_ctx.base_map is not None:
                base_raw = lang_ctx.base_map.get(key)
                if not isinstance(base_raw, str) or not base_raw.strip():
                    return ""
        return base_value

    # System/technical questions are noise for translators; represent compactly.
    qtype_norm = qtype.strip().lower()
    if is_system_question_type(qtype):
        if qtype_norm == "timing":
            block_label = "Timing Block"
        elif qtype_norm in {"meta", "metainfo"}:
            block_label = "Meta Block"
        elif qtype_norm == "captcha":
            block_label = "Captcha Block"
        else:
            block_label = "System Block"
        rows_to_render.append(("compact", block_label))
        table = doc.add_table(rows=len(rows_to_render), cols=1)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for idx, (kind, payload) in enumerate(rows_to_render):
            cell = table.rows[idx].cells[0]
            if kind == "meta":
                qid_, qt_abbrev_, has_js_, tag_, marker_ = payload  # type: ignore[misc]
                _add_question_metadata_in_cell(
                    cell,
                    qid=str(qid_),
                    qt_abbrev=str(qt_abbrev_),
                    has_js=bool(has_js_),
                    export_tag=str(tag_),
                    validation_marker=str(marker_),
                )
            elif kind == "compact":
                # For system/technical blocks
                _add_system_note(cell, str(payload), depth=0)
        doc.add_paragraph("")  # spacing after the question
        return

    display_logic = question.get("DisplayLogic")
    if display_logic:
        dl = _format_logic_blob(
            display_logic,
            questions=questions,
            translation_ctx=translation_ctx,
        )
        rows_to_render.append(("logic", f"DISPLAY IF: {dl}".strip()))

    qtext_base = question.get("QuestionText") or ""
    qtext_key = lang_ctx.key_for_question_text(qid) if lang_ctx else ""
    qtext_target = (
        _resolve_translation(qtext_key, str(qtext_base))
        if qtext_key
        else str(qtext_base)
    )
    if bilingual:
        if _has_renderable_text(qtext_base) or _has_renderable_text(qtext_target):
            rows_to_render.append(
                ("text_bilingual", (qtext_base, qtext_target, target_lang))
            )
    else:
        if _has_renderable_text(qtext_target):
            rows_to_render.append(("text", qtext_target))

    # Precompute statements/answers so we can omit empty rows.
    # For options, track their origin so translation keys map correctly:
    # - non-Matrix Choices -> QID_ChoiceN
    # - non-Matrix Answers -> QID_AnswerN
    # - Matrix rows (subitems) -> QID_ChoiceN
    # - Matrix cols (options) -> QID_AnswerN
    # - SBSMatrix statements (Choices) -> QID_ChoiceN
    # - SBSMatrix per-column headers -> QID#<col>_QuestionText
    # - SBSMatrix per-column answers -> QID#<col>_AnswerN
    statement_items: list[tuple[str, str]] = []
    # (item_id, base_display, kind) where kind in {"choice","answer","label"}
    option_items: list[tuple[str, str, str]] = []
    label_items: list[tuple[str, str]] = []
    scale_info_lines: list[str] = []
    sbs_columns: list[tuple[object, ...]] = []

    qtype_code = (question.get("QuestionType") or "").strip()
    is_sbs_matrix = _is_sbs_matrix_question(question)
    slider_like = qtype_code in {"Slider", "CS"}
    if is_sbs_matrix:
        statement_items = list(
            _iter_ordered_displays(question, question.get("Choices") or {}, order_key="ChoiceOrder")
        )
    elif qtype_code == "Matrix":
        rows = question.get("Choices") or {}
        cols = question.get("Answers") or {}
        statement_items = list(
            _iter_ordered_displays(question, rows, order_key="ChoiceOrder")
        )
        option_items = [
            (aid, disp, "answer")
            for aid, disp in _iter_ordered_displays(
                question, cols, order_key="AnswerOrder"
            )
        ]
    elif slider_like:
        choices = question.get("Choices") or {}
        answers = question.get("Answers") or {}
        statement_items = list(
            _iter_ordered_displays(question, choices, order_key="ChoiceOrder")
        )
        for aid in sorted(
            answers.keys(), key=lambda x: int(x) if str(x).isdigit() else str(x)
        ):
            ans = answers.get(aid) or {}
            disp = _coerce_display_text(ans.get("Display"))
            if disp.strip():
                option_items.append((str(aid), disp, "answer"))

        labels = question.get("Labels") or {}
        if isinstance(labels, dict) and labels:
            for k in sorted(labels.keys(), key=str):
                lab = labels.get(k) or {}
                disp = _coerce_display_text(lab.get("Display"))
                if disp.strip():
                    label_items.append((str(k), _strip_html(disp)))

        cfg = question.get("Configuration") or {}
        if isinstance(cfg, dict):
            min_v = cfg.get("CSSliderMin")
            max_v = cfg.get("CSSliderMax")
            grid = cfg.get("GridLines")
            dec = cfg.get("NumDecimals")
            snap = cfg.get("SnapToGrid")
            show_val = cfg.get("ShowValue")
            if min_v is not None or max_v is not None:
                lo = "" if min_v is None else str(min_v)
                hi = "" if max_v is None else str(max_v)
                if lo and hi:
                    scale_info_lines.append(f"Range: {lo}–{hi}")
                else:
                    scale_info_lines.append(f"Range: {lo or '—'}–{hi or '—'}")
            if grid is not None:
                scale_info_lines.append(f"GridLines: {grid}")
            if dec is not None:
                scale_info_lines.append(f"Decimals: {dec}")
            if snap is not None:
                scale_info_lines.append(f"SnapToGrid: {snap}")
            if show_val is not None:
                scale_info_lines.append(f"ShowValue: {show_val}")
    else:
        choices = question.get("Choices") or {}
        answers = question.get("Answers") or {}
        for cid, disp in _iter_ordered_displays(
            question, choices, order_key="ChoiceOrder"
        ):
            option_items.append((str(cid), disp, "choice"))
        for aid in sorted(
            answers.keys(), key=lambda x: int(x) if str(x).isdigit() else str(x)
        ):
            ans = answers.get(aid) or {}
            disp = _coerce_display_text(ans.get("Display"))
            if disp.strip():
                option_items.append((str(aid), disp, "answer"))

    if not slider_like:
        labels = question.get("Labels") or {}
        if isinstance(labels, dict) and labels:
            for k in sorted(labels.keys(), key=str):
                lab = labels.get(k) or {}
                disp = _coerce_display_text(lab.get("Display"))
                if disp.strip():
                    option_items.append((f"label {k}", _strip_html(disp), "label"))

    if is_sbs_matrix:
        additional = question.get("AdditionalQuestions") or {}
        q_answers = question.get("Answers") or {}
        if not isinstance(q_answers, dict):
            q_answers = {}
        if isinstance(additional, dict):
            for col_id in _ordered_numeric_string_ids(additional):
                column = additional.get(col_id) or {}
                if not isinstance(column, dict):
                    continue
                col_base_text = _coerce_display_text(column.get("QuestionText"))
                col_base_answers = column.get("Answers") or {}
                if not isinstance(col_base_answers, dict):
                    col_base_answers = {}
                if not col_base_answers:
                    col_base_answers = q_answers
                col_key = (
                    lang_ctx.key_for_sbs_column_question_text(qid, str(col_id))
                    if lang_ctx
                    else ""
                )
                col_target_text = (
                    _resolve_translation(col_key, col_base_text) if col_key else col_base_text
                )

                answers = col_base_answers
                answer_order = column.get("AnswerOrder")
                translated_answers: list[tuple[str, str, str]] = []
                if isinstance(answers, dict):
                    for ans_id in _ordered_numeric_string_ids(answers, answer_order):
                        ans_data = _lookup_ordered_mapping_item(answers, str(ans_id)) or {}
                        if not isinstance(ans_data, dict):
                            continue
                        ans_base = _coerce_display_text(ans_data.get("Display"))
                        ans_target = ans_base
                        if lang_ctx:
                            ans_key = lang_ctx.key_for_sbs_column_answer(
                                qid, str(col_id), str(ans_id)
                            )
                            ans_target = _resolve_translation(ans_key, ans_base)
                            if compare_to_base:
                                translated_answers.append(
                                    (str(ans_id), ans_base, ans_target)
                                )
                            else:
                                translated_answers.append(
                                    (str(ans_id), ans_target)
                                )
                        else:
                            translated_answers.append((str(ans_id), ans_base))

                if compare_to_base:
                    sbs_columns.append(
                        (str(col_id), col_base_text, col_target_text, translated_answers)
                    )
                else:
                    sbs_columns.append((str(col_id), col_target_text, translated_answers))

    # Apply translation overlay (if requested).
    translated_statement_items: object = statement_items
    translated_option_items: object = [(aid, disp) for aid, disp, _kind in option_items]
    translated_label_items: object = label_items
    if lang_ctx:

        def translate_item(item_id: str, base_disp: str, *, kind: str) -> str:
            if kind == "label":
                label_id = (
                    str(item_id).split(" ", 1)[1]
                    if str(item_id).startswith("label ")
                    else str(item_id)
                )
                key = lang_ctx.key_for_label(qid, label_id)
            elif kind == "choice":
                key = lang_ctx.key_for_choice(qid, str(item_id))
            else:
                key = lang_ctx.key_for_answer(qid, str(item_id))
            return _resolve_translation(key, str(base_disp))

        if compare_to_base:
            if statement_items:
                translated_statement_items = [
                    (cid, disp, translate_item(str(cid), disp, kind="choice"))
                    for cid, disp in statement_items
                ]
            if option_items:
                translated_option_items = [
                    (aid, disp, translate_item(aid, disp, kind=kind))
                    for aid, disp, kind in option_items
                ]
            if label_items:
                translated_label_items = [
                    (lid, disp, translate_item(lid, disp, kind="label"))
                    for lid, disp in label_items
                ]
        else:
            if statement_items:
                translated_statement_items = [
                    (cid, translate_item(str(cid), disp, kind="choice"))
                    for cid, disp in statement_items
                ]
            if option_items:
                translated_option_items = [
                    (aid, translate_item(aid, disp, kind=kind))
                    for aid, disp, kind in option_items
                ]
            if label_items:
                translated_label_items = [
                    (lid, translate_item(lid, disp, kind="label"))
                    for lid, disp in label_items
                ]

    def _renderable(s: str) -> bool:
        return _has_renderable_text(_trim_html_edges(s))

    if statement_items:
        if bilingual:
            translated_statement_items = [
                (cid, base_disp, target_disp)
                for cid, base_disp, target_disp in translated_statement_items  # type: ignore[assignment]
                if _renderable(str(base_disp)) or _renderable(str(target_disp))
            ]
        else:
            translated_statement_items = [
                (cid, disp)
                for cid, disp in translated_statement_items  # type: ignore[assignment]
                if _renderable(str(disp))
            ]
        if bilingual:
            if translated_statement_items:
                rows_to_render.append(
                    ("statements_bilingual", translated_statement_items)
                )
        else:
            if translated_statement_items:
                rows_to_render.append(("statements", translated_statement_items))

    if label_items:
        if bilingual:
            translated_label_items = [
                (lid, base_disp, target_disp)
                for lid, base_disp, target_disp in translated_label_items  # type: ignore[assignment]
                if _renderable(str(base_disp)) or _renderable(str(target_disp))
            ]
        else:
            translated_label_items = [
                (lid, disp)
                for lid, disp in translated_label_items  # type: ignore[assignment]
                if _renderable(str(disp))
            ]
        if bilingual:
            if translated_label_items:
                rows_to_render.append(("labels_bilingual", translated_label_items))
        else:
            if translated_label_items:
                rows_to_render.append(("labels", translated_label_items))

    if is_sbs_matrix:
        filtered_sbs_columns: list[tuple[object, ...]] = []
        for col_payload in sbs_columns:
            if compare_to_base:
                col_id, col_base_text, col_target_text, answers_payload = col_payload  # type: ignore[misc]
                if answers_payload:
                    answers_payload = [
                        (aid, base_ans, target_ans)
                        for aid, base_ans, target_ans in answers_payload  # type: ignore[assignment]
                        if _renderable(str(base_ans)) or _renderable(str(target_ans))
                    ]
                if (
                    _renderable(str(col_base_text))
                    or _renderable(str(col_target_text))
                    or bool(answers_payload)
                ):
                    filtered_sbs_columns.append(
                        (str(col_id), str(col_base_text), str(col_target_text), answers_payload)
                    )
            else:
                col_id, col_target_text, column_answers = col_payload  # type: ignore[misc]
                column_answers = [
                    (aid, disp)
                    for aid, disp in column_answers  # type: ignore[assignment]
                    if _renderable(str(disp))
                ]
                if _renderable(str(col_target_text)) or column_answers:
                    filtered_sbs_columns.append(
                        (str(col_id), str(col_target_text), column_answers)
                    )

        if filtered_sbs_columns:
            if compare_to_base:
                rows_to_render.append(("sbs_columns2", filtered_sbs_columns))
            else:
                rows_to_render.append(("sbs_columns", filtered_sbs_columns))

    if option_items:
        if bilingual:
            rows_to_render.append(("answers_bilingual", translated_option_items))
        else:
            rows_to_render.append(("answers", translated_option_items))
    elif scale_info_lines:
        rows_to_render.append(("scale_info", scale_info_lines))

    # Add JS strings if question has QuestionJS and feature is enabled
    if has_js and include_js_strings:
        js_code = str(question.get("QuestionJS") or "")
        if js_code:
            js_lang = render_language or base_language
            js_strings = _extract_js_strings(js_code, target_language=js_lang)
            if js_strings:
                rows_to_render.append(("js_strings", js_strings))

    if bilingual:
        # Side-by-side bilingual rendering: base label in the left column, target in the right.
        # Metadata + (optional) display logic are shared (merged across columns).
        bilingual_rows: list[tuple[str, object]] = []
        bilingual_rows.append(("meta", (qid, qt_abbrev, has_js, tag, marker)))
        if display_logic:
            bilingual_rows.append(("logic", f"DISPLAY IF: {dl}".strip()))  # type: ignore[name-defined]
        if _has_renderable_text(qtext_base) or _has_renderable_text(qtext_target):
            bilingual_rows.append(("text2", (qtext_base, qtext_target)))
        if statement_items:
            if translated_statement_items:
                bilingual_rows.append(("statements2", translated_statement_items))
        if label_items:
            if translated_label_items:
                bilingual_rows.append(("labels2", translated_label_items))
        if is_sbs_matrix and filtered_sbs_columns:
            bilingual_rows.append(("sbs_columns2", filtered_sbs_columns))
        if option_items:
            bilingual_rows.append(("answers2", translated_option_items))
        elif scale_info_lines:
            bilingual_rows.append(("scale_info2", scale_info_lines))

        # Add JS strings row
        if has_js and include_js_strings:
            js_code = str(question.get("QuestionJS") or "")
            if js_code:
                js_lang = render_language or base_language
                js_strings = _extract_js_strings(js_code, target_language=js_lang)
                if js_strings:
                    bilingual_rows.append(("js_strings", js_strings))

        table = doc.add_table(rows=len(bilingual_rows), cols=2)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        try:
            _set_table_column_widths_in(table, [3.25, 3.25])
        except Exception:
            pass

        for idx, (kind, payload) in enumerate(bilingual_rows):
            left = table.rows[idx].cells[0]
            right = table.rows[idx].cells[1]
            if kind == "meta":
                merged = left.merge(right)
                qid_, qt_abbrev_, has_js_, tag_, marker_ = payload  # type: ignore[misc]
                _add_question_metadata_in_cell(
                    merged,
                    qid=str(qid_),
                    qt_abbrev=str(qt_abbrev_),
                    has_js=bool(has_js_),
                    export_tag=str(tag_),
                    validation_marker=str(marker_),
                )
            elif kind == "logic":
                merged = left.merge(right)
                _add_logic_line(merged, str(payload), depth=0)
            elif kind == "text2":
                base_text, target_text = payload  # type: ignore[misc]
                p_left = _container_add_paragraph(left)
                _style_table_label_paragraph(p_left)
                p_left.add_run(base_label).bold = True
                _add_rich_text_block(
                    left,
                    str(base_text),
                    depth=0,
                    include_html_source=include_html_source,
                    layout_heuristics=layout_heuristics,
                )
                p_right = _container_add_paragraph(right)
                _style_table_label_paragraph(p_right)
                p_right.add_run(target_label).bold = True
                _add_rich_text_block(
                    right,
                    str(target_text),
                    depth=0,
                    include_html_source=include_html_source,
                    layout_heuristics=layout_heuristics,
                )
            elif kind == "statements2":
                p = _container_add_paragraph(left)
                _style_table_label_paragraph(p)
                p.add_run("Statements").bold = True
                p2 = _container_add_paragraph(right)
                _style_table_label_paragraph(p2)
                p2.add_run("Statements").bold = True
                for cid, base_disp, target_disp in payload:  # type: ignore[assignment]
                    _add_choice_line(
                        left, prefix=f"[{cid}]", display=base_disp, depth=0
                    )
                    _add_choice_line(
                        right, prefix=f"[{cid}]", display=target_disp, depth=0
                    )
            elif kind == "sbs_columns2":
                p = _container_add_paragraph(left)
                _style_table_label_paragraph(p)
                p.add_run("Columns").bold = True
                p2 = _container_add_paragraph(right)
                _style_table_label_paragraph(p2)
                p2.add_run("Columns").bold = True
                for col_id, base_col_text, target_col_text, answers_payload in payload:  # type: ignore[assignment]
                    col_label = f"Column {str(col_id)}"
                    _add_choice_line(
                        left,
                        prefix=col_label,
                        display=base_col_text,
                        depth=0,
                    )
                    _add_choice_line(
                        right,
                        prefix=col_label,
                        display=target_col_text,
                        depth=0,
                    )
                    for aid, base_ans, target_ans in answers_payload:  # type: ignore[assignment]
                        _add_choice_line(
                            left,
                            prefix=f"[{col_id}.{aid}]",
                            display=base_ans,
                            depth=0,
                        )
                        _add_choice_line(
                            right,
                            prefix=f"[{col_id}.{aid}]",
                            display=target_ans,
                            depth=0,
                        )
            elif kind == "labels2":
                p = _container_add_paragraph(left)
                _style_table_label_paragraph(p)
                p.add_run("Labels").bold = True
                p2 = _container_add_paragraph(right)
                _style_table_label_paragraph(p2)
                p2.add_run("Labels").bold = True
                for lid, base_disp, target_disp in payload:  # type: ignore[assignment]
                    _add_choice_line(
                        left, prefix=f"[{lid}]", display=base_disp, depth=0
                    )
                    _add_choice_line(
                        right, prefix=f"[{lid}]", display=target_disp, depth=0
                    )
            elif kind == "scale_info2":
                merged = left.merge(right)
                p = _container_add_paragraph(merged)
                _style_table_label_paragraph(p)
                p.add_run("Scale").bold = True
                for line in payload:  # type: ignore[assignment]
                    _add_annotation(merged, str(line), depth=0)
            elif kind == "answers2":
                label = "Scale" if qtype_code == "Matrix" or slider_like else "Options"
                p = _container_add_paragraph(left)
                _style_table_label_paragraph(p)
                p.add_run(label).bold = True
                p2 = _container_add_paragraph(right)
                _style_table_label_paragraph(p2)
                p2.add_run(label).bold = True
                for aid, base_disp, target_disp in payload:  # type: ignore[assignment]
                    _add_choice_line(
                        left, prefix=f"[{aid}]", display=base_disp, depth=0
                    )
                    _add_choice_line(
                        right, prefix=f"[{aid}]", display=target_disp, depth=0
                    )

        doc.add_paragraph("")  # spacing after the question
        return

    table = doc.add_table(rows=len(rows_to_render), cols=1)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for idx, (kind, payload) in enumerate(rows_to_render):
        cell = table.rows[idx].cells[0]
        if kind == "meta":
            qid_, qt_abbrev_, has_js_, tag_, marker_ = payload  # type: ignore[misc]
            _add_question_metadata_in_cell(
                cell,
                qid=str(qid_),
                qt_abbrev=str(qt_abbrev_),
                has_js=bool(has_js_),
                export_tag=str(tag_),
                validation_marker=str(marker_),
            )
        elif kind == "logic":
            _add_logic_line(cell, str(payload), depth=0)
        elif kind == "text":
            _add_rich_text_block(
                cell,
                str(payload),
                depth=0,
                include_html_source=include_html_source,
                layout_heuristics=layout_heuristics,
            )
        elif kind == "text_bilingual":
            base_text, target_text, _lang_code = payload  # type: ignore[misc]
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run(base_label).bold = True
            _add_rich_text_block(
                cell,
                str(base_text),
                depth=0,
                include_html_source=include_html_source,
                layout_heuristics=layout_heuristics,
            )
            p2 = _container_add_paragraph(cell)
            _style_table_label_paragraph(p2)
            p2.add_run(target_label).bold = True
            _add_rich_text_block(
                cell,
                str(target_text),
                depth=0,
                include_html_source=include_html_source,
                layout_heuristics=layout_heuristics,
            )
        elif kind == "statements":
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run("Statements").bold = True
            for cid, disp in payload:  # type: ignore[assignment]
                _add_choice_line(cell, prefix=f"[{cid}]", display=disp, depth=0)
        elif kind == "statements_bilingual":
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run("Statements").bold = True
            for cid, base_disp, target_disp in payload:  # type: ignore[assignment]
                _add_choice_line(
                    cell, prefix=f"[{cid}] {base_label}:", display=base_disp, depth=0
                )
                _add_choice_line(
                    cell,
                    prefix=f"[{cid}] {target_label}:",
                    display=target_disp,
                    depth=0,
                )
        elif kind == "sbs_columns":
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run("Columns").bold = True
            for col_id, col_text, answers_payload in payload:  # type: ignore[assignment]
                _add_choice_line(
                    cell, prefix=f"Column {str(col_id)}", display=str(col_text), depth=0
                )
                for aid, disp in answers_payload:  # type: ignore[assignment]
                    _add_choice_line(
                        cell, prefix=f"[{col_id}.{aid}]", display=disp, depth=0
                    )
        elif kind == "labels":
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run("Labels").bold = True
            for lid, disp in payload:  # type: ignore[assignment]
                _add_choice_line(cell, prefix=f"[{lid}]", display=disp, depth=0)
        elif kind == "labels_bilingual":
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run("Labels").bold = True
            for lid, base_disp, target_disp in payload:  # type: ignore[assignment]
                _add_choice_line(
                    cell, prefix=f"[{lid}] {base_label}:", display=base_disp, depth=0
                )
                _add_choice_line(
                    cell,
                    prefix=f"[{lid}] {target_label}:",
                    display=target_disp,
                    depth=0,
                )
        elif kind == "answers":
            label = "Scale" if qtype_code == "Matrix" or slider_like else "Options"
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run(label).bold = True
            for aid, disp in payload:  # type: ignore[assignment]
                _add_choice_line(cell, prefix=f"[{aid}]", display=disp, depth=0)
        elif kind == "answers_bilingual":
            label = "Scale" if qtype_code == "Matrix" or slider_like else "Options"
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run(label).bold = True
            for aid, base_disp, target_disp in payload:  # type: ignore[assignment]
                _add_choice_line(
                    cell, prefix=f"[{aid}] {base_label}:", display=base_disp, depth=0
                )
                _add_choice_line(
                    cell,
                    prefix=f"[{aid}] {target_label}:",
                    display=target_disp,
                    depth=0,
                )
        elif kind == "scale_info":
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run("Scale").bold = True
            for line in payload:  # type: ignore[assignment]
                _add_annotation(cell, str(line), depth=0)
        elif kind == "js_strings":
            p = _container_add_paragraph(cell)
            _style_table_label_paragraph(p)
            p.add_run("JavaScript User-Visible Strings").bold = True
            for s in payload:  # type: ignore[union-attr]
                _add_system_note(cell, s, depth=0)

    doc.add_paragraph("")  # spacing after the question


def _iter_ordered_displays(
    question: dict, mapping: dict, *, order_key: str
) -> Iterable[Tuple[str, str]]:
    order = question.get(order_key)
    ordered_ids: List[str] = []
    if isinstance(order, list) and order:
        ordered_ids.extend([str(x) for x in order if str(x) in mapping])
    for k in mapping.keys():
        sk = str(k)
        if sk not in ordered_ids:
            ordered_ids.append(sk)
    for k in ordered_ids:
        item = mapping.get(k) or {}
        disp = _coerce_display_text(item.get("Display"))
        if disp.strip():
            yield k, disp


def _coerce_display_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, bool):
        return "True" if value else "False"
    if isinstance(value, (int, float)):
        return str(value)
    return str(value)


# ----------------------------
# Rendering helpers
# ----------------------------


def _add_heading(doc, text: str, *, level: int, depth: int) -> None:
    p = doc.add_heading(text, level=level)
    _set_indent(p, depth=depth)


def _add_annotation(container, text: str, *, depth: int) -> None:
    p = _container_add_paragraph(container)
    p.add_run(text)
    _set_indent(p, depth=depth)


def _style_system_run(run) -> None:
    """Style a run as a system note (monospace), without changing size/boldness."""

    try:
        run.font.name = "Courier New"
    except Exception:
        return


def _style_id_run(run) -> None:
    """Style an ID run (QID/BlockID) as monospace, preserving other formatting."""

    _style_system_run(run)


def _style_qid_run(run) -> None:
    """Style a QID run as monospace with slightly larger font for readability."""

    _style_system_run(run)
    try:
        from docx.shared import Pt

        run.font.size = Pt(11)
    except Exception:
        return


def _add_system_note(container, text: str, *, depth: int) -> None:
    p = _container_add_paragraph(container)
    _set_indent(p, depth=depth)
    run = p.add_run(text)
    _style_system_run(run)


def _add_choice_line(container, *, prefix: str, display: str, depth: int) -> None:
    p = _container_add_paragraph(container)
    _set_indent(p, depth=depth)
    run = p.add_run(f"{prefix} ")
    run.bold = True
    _add_rich_text_to_paragraph(p, display)


_LEADING_BREAKS_RE = re.compile(r"(?is)^\s*(?:<br\s*/?>\s*)+")
_TRAILING_BREAKS_RE = re.compile(r"(?is)(?:<br\s*/?>\s*)+\s*$")
_LEADING_EMPTY_BLOCK_RE = re.compile(
    r"(?is)^\s*<(p|div)[^>]*>\s*(?:&nbsp;|\s|<br\s*/?>)*\s*</\1>\s*"
)
_TRAILING_EMPTY_BLOCK_RE = re.compile(
    r"(?is)\s*<(p|div)[^>]*>\s*(?:&nbsp;|\s|<br\s*/?>)*\s*</\1>\s*$"
)


def _trim_html_edges(html_str: str) -> str:
    """Trim leading/trailing whitespace and empty paragraphs from HTML-ish strings."""

    s = str(html_str or "").strip()
    if not s:
        return ""

    # Trim common leading/trailing empty paragraphs and <br> runs.
    for _ in range(50):
        before = s
        s = _LEADING_BREAKS_RE.sub("", s)
        s = _TRAILING_BREAKS_RE.sub("", s)
        s = _LEADING_EMPTY_BLOCK_RE.sub("", s)
        s = _TRAILING_EMPTY_BLOCK_RE.sub("", s)
        s = s.strip()
        if s == before:
            break

    return s.strip()


def _trim_paragraph_edge_breaks(paragraph) -> None:
    """Trim leading/trailing break-only runs from a paragraph (best effort)."""

    try:
        from docx.oxml.ns import nsmap
    except Exception:
        return

    try:
        p_elm = paragraph._p  # type: ignore[attr-defined]
    except Exception:
        return

    ns = {
        "w": nsmap.get("w")
        or "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    }

    def is_break_run(r) -> bool:
        # Remove runs that contain only line breaks/tabs (no text).
        if r.find("w:t", ns) is not None:
            return False
        return (r.find("w:br", ns) is not None) or (r.find("w:tab", ns) is not None)

    runs = list(p_elm.findall("w:r", ns))
    while runs and is_break_run(runs[0]):
        p_elm.remove(runs[0])
        runs = list(p_elm.findall("w:r", ns))

    runs = list(p_elm.findall("w:r", ns))
    while runs and is_break_run(runs[-1]):
        p_elm.remove(runs[-1])
        runs = list(p_elm.findall("w:r", ns))


@dataclass
class _HtmlBlock:
    kind: str  # "p" | "heading" | "list_item"
    html: str
    heading_level: int = 0
    list_kind: str = ""  # "ul" | "ol"
    list_depth: int = 0


def _split_html_into_blocks(html_str: str) -> list[_HtmlBlock]:
    """Split HTML into block-level chunks so we can render real paragraphs/lists."""

    from html.parser import HTMLParser

    blocks: list[_HtmlBlock] = []
    buf: list[str] = []
    current_kind: str | None = None  # "p" | "heading" | "list_item"
    current_heading_level: int = 0
    current_list_kind: str = ""
    current_list_depth: int = 0
    list_stack: list[str] = []

    def flush() -> None:
        nonlocal buf, current_kind, current_heading_level, current_list_kind, current_list_depth
        s = _trim_html_edges("".join(buf))
        buf = []
        if not s or not current_kind:
            current_kind = None
            current_heading_level = 0
            current_list_kind = ""
            current_list_depth = 0
            return
        blocks.append(
            _HtmlBlock(
                kind=current_kind,
                html=s,
                heading_level=current_heading_level,
                list_kind=current_list_kind,
                list_depth=current_list_depth,
            )
        )
        current_kind = None
        current_heading_level = 0
        current_list_kind = ""
        current_list_depth = 0

    def start_block(
        kind: str, *, heading_level: int = 0, list_kind: str = "", list_depth: int = 0
    ) -> None:
        nonlocal current_kind, current_heading_level, current_list_kind, current_list_depth
        if current_kind is not None:
            flush()
        current_kind = kind
        current_heading_level = heading_level
        current_list_kind = list_kind
        current_list_depth = list_depth

    def attrs_str(attrs) -> str:
        # Keep only attributes relevant to our inline renderer.
        keep = {"style", "href", "target", "rel", "title"}
        parts: list[str] = []
        for k, v in attrs or []:
            if not k:
                continue
            lk = str(k).lower()
            if lk not in keep:
                continue
            if v is None:
                parts.append(f" {lk}")
            else:
                parts.append(f' {lk}="{_html.escape(str(v))}"')
        return "".join(parts)

    class Parser(HTMLParser):
        def handle_starttag(self, tag, attrs):
            t = str(tag or "").lower()
            if t in {"p", "div"}:
                # Common Qualtrics pattern: <li><p>...</p></li>. Inside list items,
                # treat <p>/<div> as structural noise (do not split the list item).
                if current_kind == "list_item":
                    return
                start_block("p")
                return
            if t in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                lvl = int(t[1]) if len(t) == 2 and t[1].isdigit() else 2
                if current_kind == "list_item":
                    # Don't split list items on headings; render inline.
                    return
                start_block("heading", heading_level=lvl)
                return
            if t in {"ul", "ol"}:
                flush()
                list_stack.append(t)
                return
            if t == "li":
                # Each list item becomes its own paragraph.
                flush()
                lk = list_stack[-1] if list_stack else "ul"
                depth = max(0, len(list_stack) - 1)
                start_block("list_item", list_kind=lk, list_depth=depth)
                return
            if t == "br":
                if current_kind is None:
                    start_block("p")
                buf.append("<br/>")
                return

            # Inline tags: preserve minimal markup.
            if current_kind is None:
                start_block("p")
            buf.append(f"<{t}{attrs_str(attrs)}>")

        def handle_endtag(self, tag):
            t = str(tag or "").lower()
            if t in {"p", "div"}:
                if current_kind == "list_item":
                    return
                flush()
                return
            if t in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                if current_kind == "list_item":
                    return
                flush()
                return
            if t == "li":
                flush()
                return
            if t in {"ul", "ol"}:
                flush()
                if list_stack:
                    list_stack.pop()
                return

            if current_kind is None:
                start_block("p")
            buf.append(f"</{t}>")

        def handle_data(self, data):
            s = _html.escape(_html.unescape(data or ""))
            if not s:
                return
            if current_kind is None:
                start_block("p")
            buf.append(s)

    try:
        parser = Parser()
        parser.feed(html_str or "")
        parser.close()
    except Exception:
        # Fallback: treat as a single paragraph.
        s = _trim_html_edges(html_str)
        return [_HtmlBlock(kind="p", html=s)] if s else []

    flush()
    # Drop empty blocks.
    return [b for b in blocks if b and b.html.strip()]


def _render_html_blocks_to_container(
    container, html_str: str, *, depth: int, layout_heuristics: bool = False
) -> None:
    """Render safe HTML into real paragraphs/lists/headings in the container."""

    html_str = _trim_html_edges(html_str)
    if not html_str:
        return

    blocks = _split_html_into_blocks(html_str)
    if not blocks:
        return

    def heading_style(level: int) -> str:
        # Keep these as real heading styles so translators can navigate via outline.
        return {
            1: "Heading 3",
            2: "Heading 4",
            3: "Heading 5",
            4: "Heading 6",
            5: "Heading 6",
            6: "Heading 6",
        }.get(int(level) if level else 4, "Heading 6")

    def list_style(kind: str) -> str:
        return "List Number" if (kind or "").lower() == "ol" else "List Bullet"

    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.kind == "list_item":
            # Group consecutive list items of the same list depth/kind.
            j = i + 1
            group = [b]
            while j < len(blocks):
                nb = blocks[j]
                if (
                    nb.kind == "list_item"
                    and nb.list_kind == b.list_kind
                    and nb.list_depth == b.list_depth
                ):
                    group.append(nb)
                    j += 1
                    continue
                break

            parsed: list[tuple[str, str]] = []
            if layout_heuristics:
                # Special case (layout heuristic): "[task]/[time]/[pay]/[reward]" lists -> render as a 2-col table.
                # This does not match Qualtrics' runtime HTML structure; keep it opt-in.
                label_re = re.compile(
                    r"^\[\s*(task|time|pay|reward)\s*\]\s*(.*)$", re.IGNORECASE
                )
                for it in group:
                    txt = _strip_html(it.html).strip()
                    m = label_re.match(txt)
                    if not m:
                        parsed = []
                        break
                    parsed.append((m.group(1).lower(), m.group(2).strip()))

            if parsed:
                table = _container_add_table(container, rows=len(parsed), cols=2)
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
                _set_table_column_widths_in(table, [1.2, 4.8])
                for r_idx, (lab, val) in enumerate(parsed):
                    c0, c1 = table.rows[r_idx].cells
                    p0 = _container_add_paragraph(c0)
                    run0 = p0.add_run(f"[{lab}]")
                    run0.bold = True
                    p1 = _container_add_paragraph(c1)
                    _add_text_with_edf_styling(p1, val)
                _set_indent_table(table, depth=depth)
            else:
                for it in group:
                    p = _container_add_paragraph(
                        container, style=list_style(it.list_kind)
                    )
                    _set_indent(p, depth=depth + it.list_depth)
                    _SafeHtmlToDocx(p, emit_block_breaks=False).feed(it.html)
                    _trim_paragraph_edge_breaks(p)

            i = j
            continue

        if b.kind == "heading":
            p = _container_add_paragraph(
                container, style=heading_style(b.heading_level)
            )
            _set_indent(p, depth=depth)
            # Headings from Qualtrics are frequently just plain text (no <h*> wrapper
            # in our block representation); force bold for the heading content.
            _SafeHtmlToDocx(p, emit_block_breaks=False).feed(
                f"<strong>{b.html}</strong>"
            )
            _trim_paragraph_edge_breaks(p)
            i += 1
            continue

        # Normal paragraph block
        p = _container_add_paragraph(container)
        _set_indent(p, depth=depth)
        _SafeHtmlToDocx(p, emit_block_breaks=False).feed(b.html)
        _trim_paragraph_edge_breaks(p)
        i += 1


def _add_rich_text_block(
    container,
    html_str: str,
    *,
    depth: int,
    include_html_source: bool = True,
    layout_heuristics: bool = False,
) -> None:
    html_str = _trim_html_edges(str(html_str or ""))
    if not html_str:
        return
    # Special case: render HTML tables as simplified Word tables (even when other checks
    # might classify the HTML as "safe").
    if _TABLE_TAG_RE.search(html_str) and should_treat_as_html(html_str):
        _render_html_with_tables(container, html_str, depth=depth)
        if include_html_source:
            _add_annotation(container, "HTML (source):", depth=depth)
            _add_code_block(
                container, _sanitize_html_source_for_doc(html_str), depth=depth
            )
        return
    if is_markdown_safe_html(html_str):
        _render_html_blocks_to_container(
            container, html_str, depth=depth, layout_heuristics=layout_heuristics
        )
        return
    if should_treat_as_html(html_str):
        if _can_render_html_to_docx(html_str):
            _render_html_blocks_to_container(
                container, html_str, depth=depth, layout_heuristics=layout_heuristics
            )
            return
        # Best-effort rendering for complex-but-not-dangerous HTML:
        # render what we can, and keep the full source below in monospace black.
        if not _DANGEROUS_HTML_RE.search(html_str):
            _render_html_blocks_to_container(
                container, html_str, depth=depth, layout_heuristics=layout_heuristics
            )
            if include_html_source:
                _add_annotation(container, "HTML (source):", depth=depth)
                _add_code_block(
                    container, _sanitize_html_source_for_doc(html_str), depth=depth
                )
            return
        # Dangerous HTML (interactive/media/table). Try to render what we can by
        # stripping/replacing unsafe tags, and always keep the full source below.
        sanitized = _sanitize_dangerous_html_for_best_effort(html_str)
        if sanitized.strip():
            _render_html_blocks_to_container(
                container, sanitized, depth=depth, layout_heuristics=layout_heuristics
            )
            if include_html_source:
                _add_annotation(
                    container,
                    "HTML (source):",
                    depth=depth,
                )
                _add_code_block(
                    container, _sanitize_html_source_for_doc(html_str), depth=depth
                )
            return
        _add_annotation(container, "RAW HTML:", depth=depth)
        _add_code_block(container, _sanitize_html_source_for_doc(html_str), depth=depth)
        return
    p = _container_add_paragraph(container)
    _set_indent(p, depth=depth)
    _add_text_with_edf_styling(p, _strip_html(html_str))


def _add_code_block(container, code: str, depth: int = 0) -> None:
    # python-docx doesn't have true code blocks; we approximate with monospaced runs.
    for line in str(code or "").splitlines() or [""]:
        p = _container_add_paragraph(container)
        _set_indent(p, depth=depth)
        run = p.add_run(line)
        run.font.name = "Courier New"
        try:
            from docx.shared import Pt

            run.font.size = Pt(8)
        except Exception:
            pass


def _set_indent(paragraph, *, depth: int) -> None:
    if depth <= 0:
        return
    try:
        from docx.shared import Inches
    except Exception:
        return
    paragraph.paragraph_format.left_indent = Inches(0.25 * depth)


_TABLE_RE = re.compile(r"(?is)<table\b.*?</table>")
_CANVAS_RE = re.compile(r"(?is)<canvas\b[^>]*>.*?</canvas>")
_IMG_RE = re.compile(r"(?is)<img\b[^>]*>")


def _sanitize_dangerous_html_for_best_effort(html_str: str) -> str:
    """Best-effort sanitization for HTML we don't want to render verbatim.

    Goal: preserve readable content while removing interactive/media elements.
    We still include the full HTML source in the export document.
    """

    s = str(html_str or "")
    if not s.strip():
        return ""

    # Drop scripts/styles entirely (no meaningful visible content for translators).
    s = re.sub(r"(?is)<script\b.*?</script>", "", s)
    s = re.sub(r"(?is)<style\b.*?</style>", "", s)

    # Replace interactive/media elements with placeholders.
    s = _CANVAS_RE.sub("<p><code>[Interactive chart omitted in export]</code></p>", s)

    # Replace buttons with an explicit placeholder, keeping their visible label.
    def button_repl(m: re.Match) -> str:
        label = _strip_html(m.group(1) or "").strip() or "button"
        # Use a bracketed placeholder so it reads clearly in Word exports.
        return f'<code>[button "{_html.escape(label)}" omitted]</code>'

    s = re.sub(r"(?is)<button\b[^>]*>(.*?)</button>", button_repl, s)

    def img_repl(m: re.Match) -> str:
        tag = m.group(0) or ""
        alt = None
        src = None
        m_alt = re.search(r"\balt\s*=\s*\"([^\"]*)\"", tag, flags=re.IGNORECASE)
        if m_alt:
            alt = m_alt.group(1).strip()
        m_src = re.search(r"\bsrc\s*=\s*\"([^\"]*)\"", tag, flags=re.IGNORECASE)
        if m_src:
            src = m_src.group(1).strip()
        label = alt or (src if src else "image")
        label = _strip_html(label) if label else "image"
        return f"<p><code>[Image omitted: {label}]</code></p>"

    s = _IMG_RE.sub(img_repl, s)

    # Unwrap form controls/buttons while preserving their inner text (if any).
    s = re.sub(r"(?is)</?\s*(form|input|textarea|select|option)\b[^>]*>", "", s)
    s = re.sub(r"(?is)</?\s*(iframe|svg|video|audio)\b[^>]*>", "", s)

    # Remove common JS handler/data attrs so remaining tags are simpler.
    s = re.sub(r"\s+data-[a-zA-Z0-9_:-]+\s*=\s*\"[^\"]*\"", "", s)
    s = re.sub(r"\s+on[a-zA-Z0-9_:-]+\s*=\s*\"[^\"]*\"", "", s)

    return s.strip()


def _sanitize_html_source_for_doc(html_str: str) -> str:
    """Sanitize HTML shown in `HTML (source):` blocks to avoid noisy open-mechanics."""

    s = str(html_str or "")
    if not s.strip():
        return ""

    # Drop event handler attrs (e.g., onclick JS that controls how links open).
    s = re.sub(r"\s+on[a-zA-Z0-9_:-]+\s*=\s*\"[^\"]*\"", "", s)
    # Remove data/aria attrs (not relevant to translators, but very noisy).
    s = re.sub(r"\s+(data|aria)-[a-zA-Z0-9_:-]+\s*=\s*\"[^\"]*\"", "", s)

    # Normalize known app deep-links to a stable web URL.
    s = re.sub(
        r"(?i)href\s*=\s*\"\s*bsky:[^\"]*\"",
        'href="https://bsky.app/"',
        s,
    )

    # Replace buttons with a readable placeholder, keeping their visible label.
    def button_src_repl(m: re.Match) -> str:
        label = _strip_html(m.group(1) or "").strip() or "button"
        return f'[button "{label}" omitted]'

    s = re.sub(r"(?is)<button\b[^>]*>(.*?)</button>", button_src_repl, s)

    return s.strip()


def _render_html_with_tables(container, html_str: str, *, depth: int) -> None:
    """Render HTML that contains <table> blocks into paragraphs + simplified Word tables."""

    s = str(html_str or "")
    if not s.strip():
        return

    pos = 0
    for m in _TABLE_RE.finditer(s):
        before = s[pos : m.start()]
        if _has_renderable_text(before):
            p = _container_add_paragraph(container)
            _set_indent(p, depth=depth)
            _add_rich_text_to_paragraph(p, before)

        _render_html_table_to_docx(container, m.group(0) or "", depth=depth)
        pos = m.end()

    tail = s[pos:]
    if _has_renderable_text(tail):
        p = _container_add_paragraph(container)
        _set_indent(p, depth=depth)
        _add_rich_text_to_paragraph(p, tail)


def _render_html_table_to_docx(container, table_html: str, *, depth: int) -> None:
    """Convert an HTML <table> into a simplified Word table (no buttons, no JS)."""

    from html.parser import HTMLParser

    class TableParser(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._in_td = False
            self._in_tr = False
            self._cell_parts: list[str] = []
            self._row: list[str] = []
            self.rows: list[list[str]] = []

        def handle_starttag(self, tag, attrs):
            t = str(tag or "").lower()
            if t == "tr":
                self._in_tr = True
                self._row = []
                return
            if t in {"td", "th"}:
                self._in_td = True
                self._cell_parts = []
                return
            if not self._in_td:
                return
            if t == "br":
                self._cell_parts.append("<br/>")
                return
            # Keep nested tags (with minimal attrs) so we can render bold/italics/colors.
            attrs_str = "".join(
                f' {k}="{_html.escape(str(v))}"' if v is not None else f" {k}"
                for k, v in (attrs or [])
                if k
            )
            self._cell_parts.append(f"<{t}{attrs_str}>")

        def handle_endtag(self, tag):
            t = str(tag or "").lower()
            if t in {"td", "th"} and self._in_td:
                self._row.append("".join(self._cell_parts).strip())
                self._in_td = False
                return
            if t == "tr" and self._in_tr:
                if any(_strip_html(c).strip() for c in self._row):
                    self.rows.append(self._row)
                self._row = []
                self._in_tr = False
                return
            if self._in_td:
                self._cell_parts.append(f"</{t}>")

        def handle_data(self, data):
            if not self._in_td:
                return
            self._cell_parts.append(_html.escape(_html.unescape(data or "")))

    parser = TableParser()
    try:
        parser.feed(table_html or "")
        parser.close()
    except Exception:
        p = _container_add_paragraph(container)
        _set_indent(p, depth=depth)
        _add_text_with_edf_styling(p, _strip_html(table_html))
        return

    rows = parser.rows
    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    if max_cols <= 0:
        return

    # Drop columns that are entirely empty.
    col_text: list[list[str]] = [[] for _ in range(max_cols)]
    for r in rows:
        for i in range(max_cols):
            col_text[i].append(_strip_html(r[i] if i < len(r) else "").strip())
    keep_cols = [i for i in range(max_cols) if any(v for v in col_text[i])]

    # Heuristic: drop a right-most column that is mostly copy-buttons.
    if keep_cols:
        last_i = keep_cols[-1]
        vals = [v for v in col_text[last_i] if v]
        if vals:
            copyish = sum(
                1 for v in vals if v.lower() in {"copy"} or "button" in v.lower()
            )
            if copyish / len(vals) >= 0.6:
                keep_cols = keep_cols[:-1]

    if not keep_cols:
        # Fallback: show stripped content.
        p = _container_add_paragraph(container)
        _set_indent(p, depth=depth)
        _add_text_with_edf_styling(p, _strip_html(table_html))
        return

    table = _container_add_table(container, rows=len(rows), cols=len(keep_cols))
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for r_idx, row in enumerate(rows):
        for out_c, src_c in enumerate(keep_cols):
            cell = table.rows[r_idx].cells[out_c]
            html_cell = row[src_c] if src_c < len(row) else ""
            html_cell = _sanitize_dangerous_html_for_best_effort(html_cell)
            if not _has_renderable_text(html_cell):
                continue
            p = _container_add_paragraph(cell)
            _add_rich_text_to_paragraph(p, html_cell)

    _set_indent_table(table, depth=depth)


def _container_add_table(container, *, rows: int, cols: int):
    return container.add_table(rows=rows, cols=cols)


def _set_indent_table(table, *, depth: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in getattr(cell, "paragraphs", []) or []:
                _set_indent(p, depth=depth)


class _SafeHtmlToDocx:
    """Very small HTML -> docx renderer for our safe subset."""

    def __init__(self, paragraph, *, emit_block_breaks: bool = True) -> None:
        from html.parser import HTMLParser

        class Parser(HTMLParser):
            def __init__(self, outer: "_SafeHtmlToDocx") -> None:
                super().__init__()
                self.outer = outer

            def handle_starttag(self, tag, attrs):
                self.outer._start(tag.lower(), attrs)

            def handle_endtag(self, tag):
                self.outer._end(tag.lower())

            def handle_startendtag(self, tag, attrs):
                self.handle_starttag(tag, attrs)

            def handle_data(self, data):
                self.outer._text(data)

        self._parser = Parser(self)
        self.p = paragraph
        self.bold = False
        self.italic = False
        self.underline = False
        self.superscript = False
        self.subscript = False
        self.code = False
        self.size_pt: int | None = None
        self._heading_stack: list[tuple[bool, int | None]] = []
        self._list_stack: list[tuple[str, int]] = []  # (ul|ol, counter)
        self._in_li = False
        self._style_stack: list[
            tuple[
                bool,
                bool,
                bool,
                bool,
                tuple[int, int, int] | None,
                int | None,
            ]
        ] = []
        self._href: str | None = None
        self._href_stack: list[str | None] = []
        self.color_rgb: tuple[int, int, int] | None = None
        self._has_visible_content = False
        self._emit_block_breaks = bool(emit_block_breaks)

    def feed(self, html_str: str) -> None:
        self._parser.feed(html_str or "")
        self._parser.close()

    def _maybe_break(self) -> None:
        # Avoid leading/trailing blank lines produced by repeated <br>/<p>/<div> wrappers.
        if not self._has_visible_content:
            return
        self.p.add_run().add_break()

    def _start(self, tag: str, attrs) -> None:
        if tag == "br":
            self._maybe_break()
        elif tag in {"p", "div"}:
            # Avoid breaking within list items where <li><p>...</p></li> is common;
            # inserting a break there would separate the bullet marker from the text.
            if self._emit_block_breaks and (not self._in_li):
                self._maybe_break()
        elif tag == "details":
            if self._emit_block_breaks:
                self._maybe_break()
        elif tag == "summary":
            if self._emit_block_breaks:
                self._maybe_break()
            self.p.add_run("▸ ")
            self.bold = True
            self._has_visible_content = True
        elif tag in {"ul", "ol"}:
            # Nested lists are rendered inline with simple bullet/number prefixes.
            if self._emit_block_breaks:
                self._maybe_break()
            if tag == "ul":
                self._list_stack.append(("ul", 0))
            else:
                self._list_stack.append(("ol", 0))
        elif tag == "li":
            self._in_li = True
            if self._emit_block_breaks:
                self._maybe_break()
            depth = max(0, len(self._list_stack) - 1)
            indent = ("\u00a0\u00a0" * depth) if depth else ""
            prefix = "•"
            if self._list_stack and self._list_stack[-1][0] == "ol":
                t, n = self._list_stack.pop()
                n += 1
                self._list_stack.append((t, n))
                prefix = f"{n}."
            _add_text_with_edf_styling(
                self.p,
                f"{indent}{prefix} ",
                bold=self.bold,
                italic=self.italic,
                underline=self.underline,
                rgb=self.color_rgb,
            )
            self._has_visible_content = True
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            if self._emit_block_breaks:
                self._maybe_break()
            self._heading_stack.append((self.bold, self.size_pt))
            self.bold = True
            # Approximate Qualtrics UI hierarchy (base text is 10pt).
            self.size_pt = {
                "h1": 14,
                "h2": 13,
                "h3": 12,
                "h4": 11,
                "h5": 11,
                "h6": 11,
            }.get(tag, self.size_pt)
        elif tag in {"strong", "b"}:
            self.bold = True
        elif tag in {"em", "i"}:
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag == "sup":
            self.superscript = True
        elif tag == "sub":
            self.subscript = True
        elif tag == "code":
            self.code = True
        elif tag == "a":
            self._href_stack.append(self._href)
            self._href = _href_normalize(_extract_href(attrs))
        elif tag in {"span", "font"}:
            # Track style via a small CSS subset.
            self._style_stack.append(
                (
                    self.bold,
                    self.italic,
                    self.underline,
                    self.code,
                    self.color_rgb,
                    self.size_pt,
                )
            )
            new_color = _extract_color_rgb(attrs)
            if new_color is not None:
                self.color_rgb = new_color
            style_str = None
            for k, v in attrs or []:
                if str(k).lower() == "style" and v:
                    style_str = str(v)
                    break
            css = _parse_inline_style(style_str)
            if _css_font_weight_is_bold(css.get("font-weight", "")):
                self.bold = True
            if _css_font_style_is_italic(css.get("font-style", "")):
                self.italic = True
            if _css_text_decoration_is_underline(css.get("text-decoration", "")):
                self.underline = True
            if _css_font_family_is_monospace(css.get("font-family", "")):
                self.code = True
            sz = _css_font_size_pt(css.get("font-size", ""))
            if sz is not None:
                self.size_pt = sz
        # Lists are handled upstream (we do not attempt to create list styles here).

    def _end(self, tag: str) -> None:
        if tag in {"strong", "b"}:
            self.bold = False
        elif tag in {"em", "i"}:
            self.italic = False
        elif tag == "u":
            self.underline = False
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            if self._heading_stack:
                self.bold, self.size_pt = self._heading_stack.pop()
            else:
                self.bold = False
                self.size_pt = None
            if self._emit_block_breaks:
                self._maybe_break()
        elif tag == "summary":
            self.bold = False
            if self._emit_block_breaks:
                self._maybe_break()
        elif tag == "details":
            if self._emit_block_breaks:
                self._maybe_break()
        elif tag == "li":
            self._in_li = False
        elif tag in {"ul", "ol"}:
            if self._list_stack:
                self._list_stack.pop()
            if self._emit_block_breaks:
                self._maybe_break()
        elif tag == "sup":
            self.superscript = False
        elif tag == "sub":
            self.subscript = False
        elif tag == "code":
            self.code = False
        elif tag == "a":
            # Don't inline hrefs for normal links (we keep the link embedded/clickable).
            # Only surface href when it contains embedded data fields.
            if self._href and _EDF_TOKEN_RE.search(self._href):
                self.p.add_run(" (")
                _add_text_with_edf_styling(self.p, self._href)
                self.p.add_run(")")
            if self._href_stack:
                self._href = self._href_stack.pop()
            else:
                self._href = None
        elif tag in {"span", "font"}:
            if self._style_stack:
                (
                    self.bold,
                    self.italic,
                    self.underline,
                    self.code,
                    self.color_rgb,
                    self.size_pt,
                ) = self._style_stack.pop()

    def _text(self, data: str) -> None:
        if data is None:
            return
        text = _html.unescape(data or "")
        if not text:
            return
        # Normalize common whitespace noise from pretty-printed HTML.
        if text.strip() == "":
            return
        text = text.replace("\r", " ").replace("\n", " ").replace("\t", " ")
        text = re.sub(r"[ ]{2,}", " ", text)
        if not self._has_visible_content:
            # Avoid leading spaces at the very start of a block.
            text = text.lstrip()
            if not text:
                return
        # In-link text: embed the link, do not inline-open mechanics.
        if self._href and not _EDF_TOKEN_RE.search(self._href):
            url = self._href
            if url.lower().startswith(("http://", "https://")):
                _append_hyperlink(
                    self.p,
                    url=url,
                    text=text,
                    bold=self.bold,
                    italic=self.italic,
                    size_pt=self.size_pt,
                )
                self._has_visible_content = True
                return
            # Non-http schemes: keep as styled text only.
            n0 = len(self.p.runs)
            _add_text_with_edf_styling(
                self.p,
                text,
                bold=self.bold,
                italic=self.italic,
                underline=True,
                rgb=(0, 0, 200),
            )
            new_runs = self.p.runs[n0:]
            self._has_visible_content = True
        else:
            n0 = len(self.p.runs)
            _add_text_with_edf_styling(
                self.p,
                text,
                bold=self.bold,
                italic=self.italic,
                underline=self.underline,
                rgb=self.color_rgb,
            )
            new_runs = self.p.runs[n0:]
            self._has_visible_content = True

        # Apply run-level superscript/subscript/monospace/font-size to newly added runs.
        for r in new_runs:
            try:
                if self.superscript:
                    r.font.superscript = True
                if self.subscript:
                    r.font.subscript = True
                if self.code:
                    r.font.name = "Courier New"
                if self.size_pt is not None:
                    from docx.shared import Pt

                    r.font.size = Pt(self.size_pt)
            except Exception:
                continue


def _add_rich_text_to_paragraph(paragraph, html_str: str) -> None:
    html_str = _trim_html_edges(str(html_str or ""))
    if not html_str:
        return

    # Quick list handling: if it's a safe HTML list, render each <li> as a bullet paragraph.
    # For simplicity, we only handle lists when the whole content starts with <ul>/<ol>.
    lower = html_str.strip().lower()
    if lower.startswith("<ul") or lower.startswith("<ol"):
        items = _extract_list_items(html_str)
        if items:
            container = paragraph._parent  # type: ignore[attr-defined]
            # Avoid leaving a blank placeholder paragraph by using it for the first bullet.
            first = items[0]
            try:
                paragraph.style = "List Bullet"
            except Exception:
                pass
            _SafeHtmlToDocx(paragraph).feed(first)
            _trim_paragraph_edge_breaks(paragraph)
            for item in items[1:]:
                p = container.add_paragraph(style="List Bullet")  # type: ignore[call-arg]
                _set_indent(p, depth=_guess_paragraph_depth(paragraph))
                _SafeHtmlToDocx(p).feed(item)
                _trim_paragraph_edge_breaks(p)
            return

    _SafeHtmlToDocx(paragraph).feed(html_str)
    _trim_paragraph_edge_breaks(paragraph)


def _guess_paragraph_depth(paragraph) -> int:
    # Best-effort: infer depth from left indent. Not critical; fallback to 0.
    try:
        left = paragraph.paragraph_format.left_indent
        if left is None:
            return 0
        # Inches are EMUs in python-docx; we avoid converting precisely.
        return 0
    except Exception:
        return 0


_LI_RE = re.compile(r"<li[^>]*>(.*?)</li>", re.IGNORECASE | re.DOTALL)


def _extract_list_items(html_str: str) -> List[str]:
    return [
        m.group(1).strip()
        for m in _LI_RE.finditer(html_str or "")
        if m.group(1).strip()
    ]


_TAG_RE = re.compile(r"<[^>]+>")


def _strip_html(value: str) -> str:
    text = _TAG_RE.sub("", str(value or ""))
    return " ".join(_html.unescape(text).split())


_LOGIC_MARK_RE = re.compile(r"\[\[(Q|A|OP|EDF)\]\](.*?)\[\[/\1\]\]", re.DOTALL)

# Regex patterns for extracting user-visible strings from JavaScript
_JS_STRING_PATTERNS = [
    # Single-quoted strings
    re.compile(r"'([^'\\\n]|\\.|\\\n){10,}'", re.DOTALL),
    # Double-quoted strings
    re.compile(r'"([^"\\\n]|\\.|\\\n){10,}"', re.DOTALL),
    # Template literals (backticks)
    re.compile(r"`([^`\\]|\\.|\$\{[^}]*\}){10,}`", re.DOTALL),
]


def _strip_logic_markers(text: str) -> str:
    return _LOGIC_MARK_RE.sub(lambda m: m.group(2), str(text or ""))


_JS_IDENTIFIER_RE = re.compile(r"[A-Za-z_$][A-Za-z0-9_$]*")


def _parse_js_string_literal(text: str, start: int) -> tuple[str, int] | None:
    if start >= len(text):
        return None
    quote = text[start]
    if quote not in {"'", '"', "`"}:
        return None
    i = start + 1
    escape = False
    chars: list[str] = []
    while i < len(text):
        ch = text[i]
        if escape:
            chars.append(ch)
            escape = False
        elif ch == "\\":
            escape = True
        elif ch == quote:
            return ("".join(chars), i + 1)
        else:
            chars.append(ch)
        i += 1
    return None


def _extract_balanced_block(
    text: str, start: int, open_char: str = "{", close_char: str = "}"
) -> tuple[str, int] | None:
    if start >= len(text) or text[start] != open_char:
        return None
    depth = 0
    i = start
    string: str | None = None
    escape = False
    while i < len(text):
        ch = text[i]
        if string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == string:
                string = None
        else:
            if ch in {"'", '"', "`"}:
                string = ch
            elif ch == open_char:
                depth += 1
            elif ch == close_char:
                depth -= 1
                if depth == 0:
                    return (text[start : i + 1], i + 1)
        i += 1
    return None


def _skip_js_value(text: str, start: int) -> int:
    i = start
    depth_brace = 0
    depth_bracket = 0
    depth_paren = 0
    string: str | None = None
    escape = False
    while i < len(text):
        ch = text[i]
        if string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == string:
                string = None
        else:
            if ch in {"'", '"', "`"}:
                string = ch
            elif ch == "{":
                depth_brace += 1
            elif ch == "}":
                if depth_brace > 0:
                    depth_brace -= 1
            elif ch == "[":
                depth_bracket += 1
            elif ch == "]":
                if depth_bracket > 0:
                    depth_bracket -= 1
            elif ch == "(":
                depth_paren += 1
            elif ch == ")":
                if depth_paren > 0:
                    depth_paren -= 1
            elif (
                ch == ","
                and depth_brace == 0
                and depth_bracket == 0
                and depth_paren == 0
            ):
                return i + 1
        i += 1
    return i


def _find_copy_object(js_code: str) -> str | None:
    if not js_code:
        return None
    for match in re.finditer(r"\bCOPY\b", js_code):
        i = match.end()
        while i < len(js_code) and js_code[i].isspace():
            i += 1
        if i >= len(js_code) or js_code[i] != "=":
            continue
        i += 1
        while i < len(js_code) and js_code[i].isspace():
            i += 1
        if i >= len(js_code) or js_code[i] != "{":
            continue
        block = _extract_balanced_block(js_code, i, "{", "}")
        if block:
            return block[0]
    return None


def _parse_js_object_blocks(obj_block: str) -> dict[str, str]:
    body = (obj_block or "").strip()
    if body.startswith("{") and body.endswith("}"):
        body = body[1:-1]
    i = 0
    blocks: dict[str, str] = {}
    while i < len(body):
        while i < len(body) and body[i] in " \t\r\n,":
            i += 1
        if i >= len(body):
            break
        key = ""
        if body[i] in {"'", '"', "`"}:
            parsed = _parse_js_string_literal(body, i)
            if not parsed:
                break
            key, i = parsed
        else:
            m = _JS_IDENTIFIER_RE.match(body, i)
            if not m:
                break
            key = m.group(0)
            i = m.end()
        while i < len(body) and body[i].isspace():
            i += 1
        if i >= len(body) or body[i] != ":":
            i = _skip_js_value(body, i)
            continue
        i += 1
        while i < len(body) and body[i].isspace():
            i += 1
        if i < len(body) and body[i] == "{":
            block = _extract_balanced_block(body, i, "{", "}")
            if block:
                blocks[str(key).upper()] = block[0]
                i = block[1]
                continue
        i = _skip_js_value(body, i)
    return blocks


def _extract_js_string_literals(
    obj_block: str, min_length: int, max_results: int
) -> list[str]:
    block = (obj_block or "").strip()
    if block.startswith("{") and block.endswith("}"):
        block = block[1:-1]
    values: list[str] = []
    seen: set[str] = set()
    pattern = re.compile(r':\s*([\'"`])((?:\\.|(?!\1).)*)\1', re.DOTALL)
    for match in pattern.finditer(block):
        value = match.group(2)
        value = value.replace("\\n", "\n").replace("\\t", "\t")
        value = value.replace("\\'", "'").replace('\\"', '"').replace("\\\\", "\\")
        trimmed = value.strip()
        if len(trimmed) < min_length:
            continue
        if trimmed in seen:
            continue
        seen.add(trimmed)
        values.append(trimmed)
        if len(values) >= max_results:
            break
    return values


def _extract_js_strings(
    js_code: str,
    min_length: int = 10,
    max_results: int = 30,
    target_language: str | None = None,
) -> list[str]:
    """Extract user-visible strings from JavaScript code.

    Prioritizes extraction from COPY objects (language dictionaries), then falls back
    to heuristic filtering for inline strings that might be user-visible.

    Args:
        js_code: JavaScript code string
        min_length: Minimum string length to consider (default 10 chars)
        max_results: Maximum number of strings to return
        target_language: Target language code (e.g., "NL", "FR"). If provided, extracts
                        from that language's COPY object; falls back to EN if not found.

    Returns:
        List of extracted string literals that might be user-visible
    """
    if not js_code:
        return []

    extracted: list[str] = []

    # Normalize target language
    target_lang = _normalize_lang_code(target_language) if target_language else None

    # First pass: Extract from COPY objects (preferred pattern)
    copy_obj = _find_copy_object(js_code)
    if copy_obj:
        lang_blocks = _parse_js_object_blocks(copy_obj)
        if target_lang:
            for key in (target_lang, "EN"):
                block = lang_blocks.get(key)
                if block:
                    extracted = _extract_js_string_literals(
                        block, min_length, max_results
                    )
                    if extracted:
                        return extracted[:max_results]
        else:
            if "EN" in lang_blocks:
                extracted = _extract_js_string_literals(
                    lang_blocks["EN"], min_length, max_results
                )
                if extracted:
                    return extracted[:max_results]
            for block in lang_blocks.values():
                extracted = _extract_js_string_literals(block, min_length, max_results)
                if extracted:
                    return extracted[:max_results]

    # Second pass: Heuristic extraction for inline strings
    # Only if no COPY object was found
    for pattern in _JS_STRING_PATTERNS:
        for match in pattern.finditer(js_code):
            raw = match.group(0)
            # Remove quotes/backticks from start/end
            if raw.startswith("'") or raw.startswith('"') or raw.startswith("`"):
                raw = raw[1:-1]

            # Unescape common JavaScript escapes
            raw = raw.replace("\\n", "\n")
            raw = raw.replace("\\t", "\t")
            raw = raw.replace("\\'", "'")
            raw = raw.replace('\\"', '"')
            raw = raw.replace("\\\\", "\\")

            trimmed = raw.strip()

            # Skip if too short
            if len(trimmed) < min_length:
                continue

            # Skip strings that are clearly technical/code
            lower = trimmed.lower()

            # Skip strings with JavaScript operators or syntax
            if any(
                op in trimmed
                for op in ["||", "&&", "===", "!==", "++", "--", "=>", "==", "!="]
            ):
                continue

            # Skip CSS selectors, jQuery selectors, and attribute selectors
            if trimmed.startswith(("#", ".", "[", "<", "{")) or trimmed.endswith(
                ("}", ">", "]", ",", ";")
            ):
                continue

            # Skip strings that look like attribute selectors or jQuery expressions
            if "[type=" in lower or "input[" in lower or "select[" in lower:
                continue

            # HTML fragments with tags or template parts (string concatenation artifacts)
            if any(x in trimmed for x in ['" +', '+ "', "' +", "+ '"]):
                continue

            # String concatenation artifacts: starts or ends with + operator
            # Examples: '+ debugId +', '+ String(i + 1).padStart(2,', etc.
            if trimmed.startswith("+ ") or trimmed.endswith(" +"):
                continue

            # Debug/logging messages with technical prefixes and colons
            # Examples: "fetch: error", "blur/change: trigger check", "startSignup: trying"
            # Also catches camelCase function names: "scheduleIdleCheck:", "resolveHandleDetails:"
            if re.match(r"^[a-z][a-zA-Z0-9_/]*:\s", trimmed):
                continue

            # Debug messages with variable assignments or technical patterns
            # Examples: "bs_ok=1", "verifying=true", ", choiceChecked=", "normalized="
            if re.search(r"[a-z_][a-z0-9_]*\s*=", trimmed):
                continue

            # Debug log fragments starting with punctuation
            # Examples: ", surname:", ", using fallback", "- falling back"
            if re.match(r'^[,\-"\s]+[a-z]', trimmed):
                continue

            # Strings ending with just punctuation (debug fragments)
            # Examples: "). Try again shortly."
            if re.match(r"^\W+\)", trimmed):
                continue

            # CSS/jQuery selectors (comma-separated class/element selectors)
            # Examples: "li, .ChoiceStructure, .QuestionAnswers, .QuestionBody"
            if re.search(r",\s*\.[A-Z]", trimmed) or re.match(
                r"^[a-z]+,\s*\.", trimmed
            ):
                continue

            # Technical strings ending with " (incomplete/fragment)
            # Examples: "Gender embedded data: \"", "Generated candidate \""
            if trimmed.endswith('"'):
                continue

            # Parenthetical debug/status indicators
            # Examples: "(bs_ok reset to 0)", "(No response)", "(Waiting for response…)"
            if trimmed.startswith("(") and trimmed.endswith(")"):
                continue

            # HTML tags or attributes
            if re.match(r"<[a-z][^>]*>", trimmed, re.IGNORECASE) or "_blank" in trimmed:
                continue

            # DOM/jQuery method names
            if any(
                word in lower
                for word in [
                    "innerhtml",
                    "innertext",
                    "textcontent",
                    "appendchild",
                    "insertbefore",
                    "insertafter",
                    "removechild",
                    "classlist",
                ]
            ):
                continue

            # CSS properties (key:value patterns)
            if ":" in trimmed and any(
                prop in lower
                for prop in [
                    "style=",
                    "display:",
                    "position:",
                    "margin:",
                    "padding:",
                    "width:",
                    "height:",
                    "color:",
                    "background:",
                    "border:",
                    "font-",
                    "text-",
                    "align:",
                    "flex",
                    "grid",
                    "z-index:",
                    "opacity:",
                    "cursor:",
                    "overflow:",
                    "float:",
                    "clear:",
                ]
            ):
                continue

            # URLs and technical paths
            if any(
                x in lower
                for x in [
                    "http://",
                    "https://",
                    "www.",
                    ".com/",
                    ".org/",
                    ".net/",
                    "api/",
                    "/api",
                    "localhost",
                    ".php",
                    ".json",
                    ".js",
                    ".css",
                    "function(",
                    "return ",
                    "var ",
                    "const ",
                    "let ",
                    "if (",
                    "typeof ",
                    "undefined",
                    "null",
                    ".length",
                    ".push(",
                    "console.",
                    "jquery",
                    "qualtrics.",
                    ".addoneready",
                    ".setembeddeddata",
                    ".getembeddeddata",
                    "queryselector",
                    "getelementby",
                    "addeventlistener",
                    "keydown",
                    "keyup",
                    "keypress",
                    "mousedown",
                    "mouseup",
                    "click ",
                    "blur ",
                    "focus ",
                    "change ",
                    "input ",
                    "submit ",
                    "touchstart",
                    "touchend",
                    "touchmove",
                ]
            ):
                continue

            # Must have some alphabetic characters and spaces (natural language indicator)
            if not any(c.isalpha() for c in trimmed):
                continue

            # Prefer strings with spaces (more likely to be messages)
            # or strings that start with capital letter and contain sentence-like patterns
            has_space = " " in trimmed
            starts_capital = trimmed and trimmed[0].isupper()
            has_punctuation = any(c in trimmed for c in ".!?,:;")

            if not (has_space or (starts_capital and has_punctuation)):
                continue

            extracted.append(trimmed)
            if len(extracted) >= max_results:
                break

        if len(extracted) >= max_results:
            break

    # Return sorted list for consistent output
    return sorted(set(extracted))[:max_results]


def _format_logic_blob(
    logic: object,
    *,
    questions: dict | None = None,
    translation_ctx: TranslationRenderContext | None = None,
) -> str:
    """Convert a Qualtrics logic object (BranchLogic/DisplayLogic) into readable text."""

    pretty = _format_logic_structured(
        logic,
        questions=questions,
        translation_ctx=translation_ctx,
    )
    if pretty:
        return pretty

    descs: List[str] = []

    def walk(obj: object) -> None:
        if obj is None:
            return
        if isinstance(obj, dict):
            d = obj.get("Description")
            if isinstance(d, str) and d.strip():
                descs.append(_strip_html(d))
            for v in obj.values():
                walk(v)
        elif isinstance(obj, list):
            for v in obj:
                walk(v)

    walk(logic)
    # Prefer unique, stable order.
    cleaned = [d for d in descs if d]
    if not cleaned:
        return "(logic present; no Description field)"
    # Collapse duplicates while preserving order.
    out: List[str] = []
    seen = set()
    for d in cleaned:
        if d in seen:
            continue
        seen.add(d)
        out.append(d)
    return " ".join(out)


_OP_MAP = {
    "EqualTo": "==",
    "NotEqualTo": "!=",
    "GreaterThan": ">",
    "GreaterThanOrEqualTo": ">=",
    "LessThan": "<",
    "LessThanOrEqualTo": "<=",
    "Contains": "contains",
    "DoesNotContain": "does not contain",
    "Selected": "is selected",
    "NotSelected": "is not selected",
}


def _format_logic_structured(
    logic: object,
    *,
    questions: dict | None = None,
    translation_ctx: TranslationRenderContext | None = None,
) -> str:
    """Best-effort structured formatter for the common BooleanExpression shape."""

    if not isinstance(logic, dict):
        return ""
    if (logic.get("Type") or "") != "BooleanExpression":
        return ""
    if_block = logic.get("0")
    if not isinstance(if_block, dict) or (if_block.get("Type") or "") != "If":
        return ""

    exprs: list[dict] = []
    conj: str | None = None
    for k, v in if_block.items():
        if not str(k).isdigit() or not isinstance(v, dict):
            continue
        if (v.get("Type") or "") != "Expression":
            continue
        exprs.append(v)
        conj = conj or (v.get("Conjuction") or v.get("Conjunction") or None)

    if not exprs:
        return ""

    parts: list[str] = []
    for e in exprs:
        s = _format_expression_structured(
            e,
            questions=questions,
            translation_ctx=translation_ctx,
        )
        if s:
            parts.append(s)

    if not parts:
        return ""

    conj_norm = str(conj or "And").strip().lower()
    joiner = " OR " if conj_norm == "or" else " AND "
    if len(parts) == 1:
        return parts[0]
    return "(" + joiner.join(parts) + ")"


def _parse_choice_id_from_locator(locator: str) -> str:
    s = str(locator or "").strip()
    if not s:
        return ""
    # Common Qualtrics locators: q://QID50/SelectableChoice/2
    m = re.search(r"/SelectableChoice/([^/]+)\s*$", s)
    if m:
        return str(m.group(1)).strip()
    m = re.search(r"/Choice/([^/]+)\s*$", s)
    if m:
        return str(m.group(1)).strip()
    return ""


def _format_expression_structured(
    expr: dict,
    *,
    questions: dict | None = None,
    translation_ctx: TranslationRenderContext | None = None,
) -> str:
    logic_type = (expr.get("LogicType") or "").strip()
    op = (expr.get("Operator") or "").strip()
    op_str = _OP_MAP.get(op, op or "?")

    if logic_type == "EmbeddedField":
        left = (expr.get("LeftOperand") or "").strip()
        right = str(expr.get("RightOperand") or "").strip()
        if not left:
            return ""
        if op in {"Contains", "DoesNotContain"}:
            return f'EDF:[[EDF]]{left}[[/EDF]] [[OP]]{op_str}[[/OP]] "{right}"'
        return f'EDF:[[EDF]]{left}[[/EDF]] [[OP]]{op_str}[[/OP]] "{right}"'

    qid = (expr.get("QuestionID") or expr.get("QuestionIDFromLocator") or "").strip()
    if questions and qid and isinstance(questions.get(qid), dict):
        q = questions.get(qid) or {}
        q_text_base = q.get("QuestionText") or q.get("QuestionDescription") or qid
        q_text = _strip_html(q_text_base)
        if translation_ctx is not None and not translation_ctx.compare_to_base:
            key = translation_ctx.key_for_question_text(qid)
            raw = translation_ctx.target_map.get(key)
            if isinstance(raw, str) and raw.strip():
                q_text = _strip_html(raw)

        if op in {"Selected", "NotSelected"}:
            loc = (
                expr.get("ChoiceLocator")
                or expr.get("LeftOperand")
                or expr.get("RightOperand")
                or ""
            )
            choice_id = _parse_choice_id_from_locator(str(loc))
            choice_text = ""
            if choice_id:
                choices = q.get("Choices") or {}
                answers = q.get("Answers") or {}
                if isinstance(choices, dict):
                    choice_obj = choices.get(choice_id) or {}
                    choice_text = _strip_html(choice_obj.get("Display") or "")
                if not choice_text and isinstance(answers, dict):
                    answer_obj = answers.get(choice_id) or {}
                    choice_text = _strip_html(answer_obj.get("Display") or "")
                if translation_ctx is not None and not translation_ctx.compare_to_base:
                    qtype = str(q.get("QuestionType") or "").strip()
                    kind = "choice"
                    if qtype == "Matrix":
                        if isinstance(answers, dict) and str(choice_id) in answers:
                            kind = "answer"
                        elif isinstance(choices, dict) and str(choice_id) in choices:
                            kind = "choice"
                    if kind == "answer":
                        key = translation_ctx.key_for_answer(qid, str(choice_id))
                    else:
                        key = translation_ctx.key_for_choice(qid, str(choice_id))
                    raw = translation_ctx.target_map.get(key)
                    if isinstance(raw, str) and raw.strip():
                        choice_text = _strip_html(raw)
            if not choice_text:
                # Fallback: Description contains the most human labels.
                desc = expr.get("Description")
                if isinstance(desc, str) and desc.strip():
                    desc_s = _strip_html(desc).strip()
                    # Prefer splitting "Question: Option Is Selected" when present.
                    if ":" in desc_s and desc_s.lower().endswith(op_str.lower()):
                        lhs, rhs = desc_s.rsplit(":", 1)
                        rhs = rhs.strip()
                        rhs = re.sub(
                            rf"(?i)\\s*{re.escape(op_str)}\\s*$", "", rhs
                        ).strip()
                        if rhs:
                            choice_text = rhs
                    else:
                        # Try to capture the last token group before the operator.
                        m = re.search(
                            rf"(?i)\\bif\\b\\s+.*?\\s+(.+?)\\s+{re.escape(op_str)}\\s*$",
                            desc_s.strip(),
                        )
                        if m:
                            choice_text = m.group(1).strip()
            if choice_text:
                return (
                    f'[[Q]]{qid}:"{q_text}"[[/Q]] '
                    f'[[A]]"{choice_text}"[[/A]] [[OP]]{op_str}[[/OP]]'
                )
            return f'[[Q]]{qid}:"{q_text}"[[/Q]] [[OP]]{op_str}[[/OP]]'

    # For question logic, the Description is typically the most readable when we can't map it.
    desc = expr.get("Description")
    if isinstance(desc, str) and desc.strip():
        d = _strip_html(desc).strip()
        if qid:
            return f"{qid}: {d}"
        return d

    left = (expr.get("LeftOperand") or "").strip()
    right = str(expr.get("RightOperand") or "").strip()
    pieces = [p for p in [qid or left, op_str, right] if p]
    return " ".join(pieces)


# ----------------------------
# Active/in-flow detection
# ----------------------------


def _collect_block_ids_from_flow(flow_obj: dict) -> List[str]:
    """Collect block IDs referenced in SurveyFlow (including nested flows)."""

    ordered: List[str] = []

    def walk(node: object) -> None:
        if node is None:
            return
        if isinstance(node, list):
            for child in node:
                walk(child)
            return
        if not isinstance(node, dict):
            return

        node_type = node.get("Type")
        if node_type in {"Block", "Standard"} and node.get("ID"):
            bid = str(node["ID"])
            if bid not in ordered:
                ordered.append(bid)

        # Common nesting patterns
        for key in ("Flow", "Then", "Else", "ElseFlow"):
            if key in node:
                walk(node.get(key))

    if isinstance(flow_obj, dict):
        walk(flow_obj.get("Flow"))
    return ordered


def _active_qids_in_flow(result: dict) -> set[str]:
    """Return QIDs placed in SurveyFlow blocks excluding Trash blocks."""

    blocks = result.get("Blocks", {}) or {}
    questions = result.get("Questions", {}) or {}
    flow = result.get("SurveyFlow") or result.get("Flow") or {}
    block_ids = _collect_block_ids_from_flow(flow) if isinstance(flow, dict) else []

    trash_blocks = {
        bid for bid, b in blocks.items() if (b.get("Type") or "").strip() == "Trash"
    }

    qids: set[str] = set()
    for bid in block_ids:
        if bid in trash_blocks:
            continue
        block = blocks.get(bid) or {}
        for elem in block.get("BlockElements", []) or []:
            if (elem.get("Type") or "") != "Question":
                continue
            qid = elem.get("QuestionID")
            if qid and qid in questions:
                qids.add(qid)
    return qids


def active_qids_in_flow(payload: Mapping[str, Any]) -> set[str]:
    """Return QIDs placed in SurveyFlow blocks excluding Trash blocks.

    This is a best-effort helper for tooling that should ignore translations for
    unplaced or Trash questions (e.g. cross-account copy validation).

    Args:
        payload: A `survey-definitions/{survey_id}` payload (or its `result` dict).

    Returns:
        Set of QIDs that are in SurveyFlow non-Trash blocks. Returns an empty set
        when the flow cannot be interpreted.
    """

    if not isinstance(payload, Mapping):
        return set()
    result = _survey_result(payload)
    if not isinstance(result, dict):
        return set()
    return _active_qids_in_flow(result)


def expected_translation_keys_for_qids(
    payload: Mapping[str, Any], *, qids: set[str]
) -> list[str]:
    """Return the translation keys (QuestionText/Choice/Answer/Label) for QIDs.

    This mirrors the key-generation logic used by `build_translation_map_from_cache`,
    but allows callers to scope validation/coverage to a specific QID set (e.g.
    in-flow questions only).
    """

    if not qids or not isinstance(payload, Mapping):
        return []
    result = _survey_result(payload)
    if not isinstance(result, dict):
        return []
    questions = result.get("Questions") or {}
    if not isinstance(questions, dict):
        return []
    return _collect_expected_translation_keys(questions, qids)


# ----------------------------
# Mermaid flow builder
# ----------------------------


def build_mermaid_flow(*, survey_id: str, flow: dict) -> str:
    """Build a Mermaid flowchart representation of SurveyFlow.

    MVP: include node types + block IDs + branch conditions (best-effort).
    """

    lines: List[str] = ["flowchart TD", f"%% SurveyID: {survey_id}"]
    if not isinstance(flow, dict):
        return "\n".join(lines + ["%% No SurveyFlow available"])
    flow_list = flow.get("Flow")
    if not isinstance(flow_list, list):
        return "\n".join(lines + ["%% SurveyFlow.Flow missing or malformed"])

    counter = 0
    node_lines: List[str] = []
    edge_lines: List[str] = []

    def new_id() -> str:
        nonlocal counter
        counter += 1
        return f"n{counter}"

    def label_for(node: dict) -> str:
        t = str(node.get("Type") or "").strip() or "Node"
        if t in {"Block", "Standard"} and node.get("ID"):
            return f"Block {node.get('ID')}"
        if t == "Branch":
            cond = _strip_logic_markers(_format_logic_blob(node.get("BranchLogic")))
            return f"IF {cond}"
        if t == "BlockRandomizer":
            return "Randomizer"
        if t == "EmbeddedData":
            return "EmbeddedData"
        if t == "EndSurvey":
            return "EndSurvey"
        if t == "Group":
            desc = (node.get("Description") or "").strip()
            return f"Group {desc}" if desc else "Group"
        if t == "WebService":
            return "WebService"
        return t

    def walk_list(nodes: list) -> Tuple[str | None, str | None]:
        """Return (first_id, last_id) for this sequence."""

        first: str | None = None
        last: str | None = None

        i = 0
        while i < len(nodes):
            node = nodes[i]
            i += 1
            if not isinstance(node, dict):
                continue
            nid = new_id()
            if first is None:
                first = nid
            lbl = label_for(node).replace('"', "'")
            shape_open, shape_close = ("[", "]")
            if str(node.get("Type") or "") == "Branch":
                shape_open, shape_close = ("{", "}")
            node_lines.append(f'{nid}{shape_open}"{lbl}"{shape_close}')

            if last is not None:
                edge_lines.append(f"{last} --> {nid}")

            if str(node.get("Type") or "") == "Branch":
                then_nodes = (
                    node.get("Flow") if isinstance(node.get("Flow"), list) else []
                )
                else_nodes = (
                    node.get("ElseFlow")
                    if isinstance(node.get("ElseFlow"), list)
                    else []
                )
                if isinstance(node.get("Then"), list):
                    then_nodes = node.get("Then")
                if isinstance(node.get("Else"), list):
                    else_nodes = node.get("Else")

                then_first, then_last = walk_list(list(then_nodes))
                else_first, else_last = walk_list(list(else_nodes))

                if then_first:
                    edge_lines.append(f"{nid} -- then --> {then_first}")
                if else_first:
                    edge_lines.append(f"{nid} -- else --> {else_first}")

                # Join back to next sibling if any exists.
                # We don't know the next node's ID yet; create an explicit join node.
                join_id = new_id()
                node_lines.append(f'{join_id}["Join"]')
                if then_last:
                    edge_lines.append(f"{then_last} --> {join_id}")
                else:
                    edge_lines.append(f"{nid} --> {join_id}")
                if else_last:
                    edge_lines.append(f"{else_last} --> {join_id}")
                else:
                    edge_lines.append(f"{nid} --> {join_id}")
                # The join becomes the last node for sequential linking.
                last = join_id
                continue

            sub = node.get("Flow")
            if isinstance(sub, list) and str(node.get("Type") or "") in {
                "Group",
                "BlockRandomizer",
            }:
                sub_first, sub_last = walk_list(list(sub))
                if sub_first:
                    edge_lines.append(f"{nid} --> {sub_first}")
                if sub_last:
                    last = sub_last
                else:
                    last = nid
                continue

            last = nid

        return first, last

    walk_list(flow_list)

    lines.extend(node_lines)
    lines.extend(edge_lines)
    return "\n".join(lines)


# ----------------------------
# QuestionJS mapping + EndSurvey refs
# ----------------------------


def _load_qid_js_mapping(mapping_path: Path, *, survey_id: str) -> Dict[str, str]:
    mapping_path = Path(mapping_path)
    if not mapping_path.exists():
        return {}

    with mapping_path.open(newline="", encoding="utf-8") as f:
        reader = csv.reader(f)
        header = next(reader, [])

        survey_col_idx: int | None = None
        for idx, field in enumerate(header):
            if field == "js_file":
                continue
            prefix = field.split("-", 1)[0]
            if prefix == survey_id:
                survey_col_idx = idx
                break
        if survey_col_idx is None:
            return {}

        out: Dict[str, str] = {}
        for row in reader:
            if not row or len(row) <= survey_col_idx:
                continue
            js_file = row[0].strip()
            qids_cell = (row[survey_col_idx] or "").strip()
            if not js_file or not qids_cell:
                continue
            for qid in [q.strip() for q in qids_cell.split(";") if q.strip()]:
                out[qid] = js_file
        return out


def _read_eos_message_from_disk(library_id: str, message_id: str) -> dict | None:
    """Read an EOS (EndSurvey DisplayMessage) library message from contents/, if present."""

    try:
        from .eos_messages import read_library_message_from_disk
    except Exception:
        return None

    try:
        return read_library_message_from_disk(library_id, message_id)
    except Exception:
        return None


def _apply_base_body_style(doc) -> None:
    """Set translator-friendly defaults for body text (headings remain Word defaults)."""

    try:
        from docx.shared import Pt

        for style in doc.styles:
            # Apply Arial broadly (including headings); sizes are preserved by style.
            try:
                style.font.name = "Arial"
            except Exception:
                continue

        normal = doc.styles["Normal"]
        normal.font.size = Pt(10)

        # Make the defaults explicit as well (docDefaults), so paragraphs/runs that
        # don't carry an explicit style still render consistently in Word.
        _set_doc_defaults_font(doc, font_name="Arial", size_pt=10)
        _set_doc_defaults_paragraph_spacing(doc, before_twips=0, after_twips=200)
    except Exception:
        return


def _set_doc_defaults_font(doc, *, font_name: str, size_pt: int) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    try:
        styles_elm = doc.styles.element  # type: ignore[attr-defined]
        doc_defaults = styles_elm.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_elm.insert(0, doc_defaults)

        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is None:
            rpr_default = OxmlElement("w:rPrDefault")
            doc_defaults.append(rpr_default)

        rpr = rpr_default.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rpr_default.append(rpr)

        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), font_name)

        sz = rpr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            rpr.append(sz)
        sz.set(qn("w:val"), str(int(size_pt) * 2))

        szcs = rpr.find(qn("w:szCs"))
        if szcs is None:
            szcs = OxmlElement("w:szCs")
            rpr.append(szcs)
        szcs.set(qn("w:val"), str(int(size_pt) * 2))
    except Exception:
        return


def _set_doc_defaults_paragraph_spacing(
    doc, *, before_twips: int | None, after_twips: int | None
) -> None:
    """Set docDefaults paragraph spacing (best effort)."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    try:
        styles_elm = doc.styles.element  # type: ignore[attr-defined]
        doc_defaults = styles_elm.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_elm.insert(0, doc_defaults)

        ppr_default = doc_defaults.find(qn("w:pPrDefault"))
        if ppr_default is None:
            ppr_default = OxmlElement("w:pPrDefault")
            doc_defaults.append(ppr_default)

        ppr = ppr_default.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            ppr_default.append(ppr)

        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)

        if before_twips is not None:
            spacing.set(qn("w:before"), str(int(before_twips)))
        if after_twips is not None:
            spacing.set(qn("w:after"), str(int(after_twips)))
    except Exception:
        return


# ----------------------------
# PDF Rendering
# ----------------------------


def _build_pdf_css() -> str:
    """Build CSS stylesheet for PDF export.

    Returns a complete CSS string with styles for all PDF sections.
    This includes page setup, typography, headings, tables, questions,
    logic blocks, and translation context styling.
    """
    return """
/* Page Setup */
@page {
  size: A4;
  margin: 2cm 1.5cm;
}

/* Typography */
body {
  font-family: 'Helvetica', 'Arial', sans-serif;
  font-size: 10pt;
  line-height: 1.4;
  color: #000;
}

/* Headings */
h1 {
  font-size: 18pt;
  color: #1a1a1a;
  margin-top: 1em;
  margin-bottom: 0.5em;
  page-break-after: avoid;
}

h2 {
  font-size: 14pt;
  color: #333;
  border-bottom: 2px solid #666;
  padding-bottom: 0.3em;
  margin-top: 1em;
  margin-bottom: 0.5em;
  page-break-after: avoid;
}

h3 {
  font-size: 12pt;
  color: #444;
  font-family: 'Courier New', monospace;
  margin-top: 0.8em;
  margin-bottom: 0.4em;
  page-break-after: avoid;
}

h4, h5, h6 {
  font-size: 11pt;
  color: #555;
  margin-top: 0.6em;
  margin-bottom: 0.3em;
}

/* Sections */
section {
  margin-bottom: 2em;
}

.export-header {
  border-bottom: 3px solid #000;
  padding-bottom: 1em;
  margin-bottom: 1.5em;
}

.translation-summary {
  background-color: #f0f8ff;
  padding: 1em;
  border-left: 4px solid #4a90e2;
  margin-bottom: 1em;
}

/* Tables */
table {
  width: 100%;
  border-collapse: collapse;
  margin: 0.5em 0;
  font-size: 9pt;
}

table th {
  background-color: #e0e0e0;
  font-weight: bold;
  text-align: left;
  padding: 0.4em;
  border: 1px solid #999;
}

table td {
  padding: 0.3em;
  border: 1px solid #ccc;
  vertical-align: top;
}

table.metadata th {
  width: 30%;
}

/* Questions */
.question {
  margin-bottom: 1.5em;
  padding: 0.8em;
  border: 1px solid #ddd;
  background: #fafafa;
  page-break-inside: avoid;
}

.question-header {
  font-family: 'Courier New', monospace;
  font-size: 11pt;
  font-weight: bold;
  margin-bottom: 0.5em;
}

.question-text {
  margin: 0.5em 0;
}

.choice-label {
  margin-top: 1em;
  margin-bottom: 0.3em;
  font-weight: bold;
}

.choices, .answers {
  list-style-type: disc;
  margin-left: 1.5em;
  margin-top: 0.5em;
}

.choices li, .answers li {
  margin-bottom: 0.3em;
}

/* JavaScript Strings */
.js-strings {
  margin-top: 1em;
  margin-bottom: 0.3em;
  font-weight: bold;
  color: #6a4c93;
}

.js-strings-list {
  list-style-type: square;
  margin-left: 1.5em;
  margin-top: 0.5em;
  color: #333;
}

.js-strings-list li {
  margin-bottom: 0.3em;
  font-family: 'Courier New', monospace;
  font-size: 0.9em;
}

/* Logic/Metadata */
.logic {
  background: #fff3cd;
  padding: 0.5em;
  margin: 0.5em 0;
  font-family: monospace;
  font-size: 9pt;
  border-left: 3px solid #ffc107;
}

.annotation {
  color: #666;
  font-style: italic;
  font-size: 9pt;
  margin: 0.3em 0;
}

.system-note {
  background: #e7f3ff;
  border-left: 3px solid #007bff;
  padding: 0.5em;
  margin: 0.5em 0;
  font-size: 9pt;
  page-break-inside: avoid;
}

/* Block containers */
.block {
  margin-bottom: 1.5em;
  page-break-inside: avoid;
}

.block-header {
  font-size: 12pt;
  font-weight: bold;
  color: #444;
  background: #f0f0f0;
  border-left: 4px solid #666;
  padding: 0.5em;
  margin-bottom: 0.8em;
  page-break-after: avoid;
}

/* Translation Context */
.translation-key {
  font-family: 'Courier New', monospace;
  color: #0066cc;
  font-size: 8pt;
}

.translation-bilingual {
  display: flex;
  gap: 1em;
}

.translation-base {
  flex: 1;
  padding: 0.5em;
  background: #f5f5f5;
}

.translation-target {
  flex: 1;
  padding: 0.5em;
  background: #f0fff0;
}

.translation-missing {
  color: #cc0000;
  font-style: italic;
}

.lang-label {
  font-weight: bold;
  font-size: 9pt;
  color: #666;
}

/* Code blocks */
pre, code {
  background: #f5f5f5;
  border: 1px solid #ccc;
  padding: 0.3em 0.5em;
  font-family: 'Courier New', monospace;
  font-size: 8pt;
  overflow-wrap: break-word;
  white-space: pre-wrap;
}

/* Utilities */
.page-break {
  page-break-before: always;
}

.no-break {
  page-break-inside: avoid;
}

/* QID and ID styling */
.qid {
  font-family: 'Courier New', monospace;
  font-weight: bold;
  color: #333;
}

/* EDF token styling */
.edf-token {
  color: #0066cc;
  font-family: 'Courier New', monospace;
  background: #e6f2ff;
  padding: 0.1em 0.3em;
  border-radius: 2px;
}

/* Block containers */
.block {
  margin-bottom: 2em;
  page-break-inside: avoid;
}

.block-header {
  font-size: 13pt;
  font-weight: bold;
  color: #444;
  border-bottom: 1px solid #999;
  padding-bottom: 0.3em;
  margin-bottom: 0.8em;
}

/* External surfaces */
.external-surface {
  background: #fffacd;
  border-left: 4px solid #ffd700;
  padding: 0.5em;
  margin: 0.5em 0;
}

/* Mermaid diagram */
.flow-diagram img {
  max-width: 100%;
  height: auto;
  margin: 1em 0;
}
""".strip()


def _escape_html(text: str) -> str:
    """Escape HTML special characters for safe embedding in HTML."""
    return _html.escape(str(text or ""))


def _sanitize_html_for_weasyprint(html: str) -> str:
    """Sanitize HTML content to work around WeasyPrint bugs.

    WeasyPrint 68 has a bug where it cannot handle calc() functions in table widths.
    This function removes calc() from inline styles on tables.

    See: https://github.com/Kozea/WeasyPrint/issues/2246
    """
    import re

    # Replace calc() in table width styles with 100%
    # Pattern matches: style="...width:calc(...)..." or style="...width: calc(...)..."
    html = re.sub(
        r'(<table[^>]*style="[^"]*width\s*:\s*)calc\([^)]+\)',
        r"\g<1>100%",
        html,
        flags=re.IGNORECASE,
    )

    return html


def _render_header_html(content: ExportContent) -> str:
    """Render header section HTML."""
    title_row = ""
    if content.survey_title:
        title_row = f"""
        <tr>
            <th>Survey title:</th>
            <td>{_escape_html(content.survey_title)}</td>
        </tr>
        """

    description_row = ""
    if content.survey_description:
        description_row = f"""
        <tr>
            <th>Survey description:</th>
            <td>{_escape_html(content.survey_description)}</td>
        </tr>
        """

    version_row = ""
    version_desc_row = ""
    if (
        content.version_number is not None
        or content.version_id
        or content.version_description
    ):
        version_value = ""
        if content.version_number is not None and content.version_id:
            version_value = f"{content.version_number} (id={content.version_id})"
        elif content.version_number is not None:
            version_value = str(content.version_number)
        elif content.version_id:
            version_value = str(content.version_id)
        version_row = f"""
        <tr>
            <th>Version:</th>
            <td>{_escape_html(version_value)}</td>
        </tr>
        """
        if content.version_description:
            version_desc_row = f"""
            <tr>
                <th>Version description:</th>
                <td>{_escape_html(content.version_description)}</td>
            </tr>
            """

    lang_info = ""
    if content.render_language:
        mode = (
            f"{content.base_language}-{content.render_language}"
            if content.compare_to_base
            else content.render_language
        )
        lang_info = f"""
        <tr>
            <th>Render language:</th>
            <td>{_escape_html(mode)} (base={_escape_html(content.base_language)})</td>
        </tr>
        """

    edf_info = ""
    if content.edf_overrides:
        edf_list = ", ".join(
            f"{k}={v}" for k, v in sorted(content.edf_overrides.items())
        )
        edf_info = f"""
        <tr>
            <th>Scenario EDF filters:</th>
            <td>{_escape_html(edf_list)}</td>
        </tr>
        """

    survey_link_row = ""
    if content.survey_link:
        survey_link_row = f"""
        <tr>
            <th>Survey link:</th>
            <td><a href="{_escape_html(content.survey_link)}" target="_blank">{_escape_html(content.survey_link)}</a></td>
        </tr>
        """

    return f"""
    <div class="export-header">
        <h1>SURVEY TRANSLATION EXPORT</h1>
        <table class="metadata">
            <tr>
                <th>SurveyID:</th>
                <td>{_escape_html(content.survey_id)}</td>
            </tr>
            <tr>
                <th>Survey name:</th>
                <td>{_escape_html(content.survey_name)}</td>
            </tr>
            {title_row}
            {description_row}
            {version_row}
            {version_desc_row}
            <tr>
                <th>Generated:</th>
                <td>{datetime.now().isoformat(timespec='seconds')}</td>
            </tr>
            {lang_info}
            {edf_info}
            {survey_link_row}
        </table>
    </div>
    """


def _render_translation_summary_html(content: ExportContent) -> str:
    """Render translation summary section HTML."""
    if content.render_plan is None:
        return ""

    plan = content.render_plan
    mode = (
        f"{plan.base_language}-{plan.target_language}"
        if content.compare_to_base
        else plan.target_language
    )

    issues_html = ""
    if plan.total_missing > 0 or plan.total_empty_but_base_nonempty > 0:
        issues = []
        if plan.total_missing > 0:
            issues.append(
                f"<li>Missing keys (not present in map): {plan.total_missing}</li>"
            )
        if plan.total_empty_but_base_nonempty > 0:
            issues.append(
                f"<li>Empty translations (base has content): {plan.total_empty_but_base_nonempty}</li>"
            )

        sample_issues = []
        for key in (plan.missing_keys + plan.empty_but_base_nonempty_keys)[:5]:
            sample_issues.append(f"<li><code>{_escape_html(key)}</code></li>")

        sample_html = ""
        if sample_issues:
            sample_html = f"""
            <p><strong>Sample issues (fix these first):</strong></p>
            <ul>
                {''.join(sample_issues)}
            </ul>
            """

        issues_html = f"""
        <ul>
            {''.join(issues)}
        </ul>
        {sample_html}
        """

    return f"""
    <section class="translation-summary">
        <h1>LANGUAGE RENDERING SUMMARY</h1>
        <p><strong>Mode:</strong> {_escape_html(mode)} (base={_escape_html(plan.base_language)})</p>
        <p><strong>Expected keys rendered by export:</strong> {plan.total_expected}</p>
        <p><strong>OK (translated or allowed-empty):</strong> {plan.total_ok}</p>
        {issues_html}
    </section>
    """


def _render_coverage_summary_html(content: ExportContent) -> str:
    """Render coverage summary section HTML."""
    result = content.survey_payload.get("result", {}) or {}
    questions = result.get("Questions", {}) or {}

    total = len(questions)
    active = len(content.active_qids)
    excluded = total - active

    return f"""
    <section>
        <h1>COVERAGE SUMMARY</h1>
        <p><strong>Total questions in JSON:</strong> {total}</p>
        <p><strong>Active & exported (in-flow, non-Trash):</strong> {active}</p>
        <p><strong>Excluded (unplaced/Trash/other):</strong> {excluded}</p>
    </section>
    """


def _render_type_legend_html(content: ExportContent) -> str:
    """Render question type legend section HTML."""
    result = content.survey_payload.get("result", {}) or {}
    questions = result.get("Questions", {}) or {}

    used: dict[str, str] = {}
    counts: dict[str, int] = {}
    for qid in sorted(content.active_qids or set(), key=str):
        q = questions.get(qid) or {}
        if not isinstance(q, dict):
            continue
        qt = str(q.get("QuestionType") or "").strip()
        sel = str(q.get("Selector") or "").strip()
        sub = str(q.get("SubSelector") or "").strip()
        abbrev, label = _question_type_abbrev_and_label(
            question_type=qt, selector=sel, subselector=sub
        )
        if not abbrev or not label:
            continue
        used.setdefault(abbrev, label)
        counts[abbrev] = counts.get(abbrev, 0) + 1

    if not used:
        return ""

    rows = "".join(
        f"<tr><td><code>{_escape_html(abbr)}</code></td><td>{_escape_html(used[abbr])}</td><td>{counts.get(abbr, 0)}</td></tr>"
        for abbr in sorted(used.keys(), key=str)
    )

    return f"""
    <section>
        <h1>QUESTION TYPE LEGEND</h1>
        <table>
            <tr><th>Abbreviation</th><th>Description</th><th>Count</th></tr>
            {rows}
        </table>
    </section>
    """


# ----------------------------
# HTML Flow Traversal (mirrors DOCX _traverse_flow)
# ----------------------------


def _traverse_flow_html(
    *,
    flow_list: list,
    blocks: dict,
    questions: dict,
    content: ExportContent,
    asked_qids: set[str] | None,
    depth: int,
    flow_trace: Callable[[str], None] | None,
) -> str:
    """Traverse SurveyFlow and render to HTML (centralized flow helper)."""
    html_parts: list[str] = []

    def on_block(node: dict, depth_level: int) -> None:
        block_html = _render_block_html(
            block_id=str(node["ID"]),
            blocks=blocks,
            questions=questions,
            content=content,
            asked_qids=asked_qids,
            depth=depth_level,
            flow_trace=flow_trace,
        )
        if block_html:
            html_parts.append(block_html)

    def on_group(node: dict, depth_level: int) -> None:
        desc = (node.get("Description") or "").strip()
        html_parts.append(_render_system_note_html(f"GROUP: {desc}", depth_level))

    def on_embedded_data(node: dict, depth_level: int) -> None:
        html_parts.append(
            _render_embedded_data_html(node, content.edf_overrides, depth_level)
        )

    def on_web_service(node: dict, depth_level: int) -> None:
        html_parts.append(_render_web_service_html(node, depth_level))

    def on_randomizer(node: dict, depth_level: int) -> None:
        html_parts.append(_render_randomizer_html(node, depth_level))

    def on_branch_open(node: dict, depth_level: int) -> None:
        cond = _format_logic_blob(
            node.get("BranchLogic"),
            questions=questions,
            translation_ctx=content.translation_ctx,
        )
        html_parts.append(_render_logic_line_html(f"BRANCH: IF {cond}", depth_level))

    def on_branch_decision(
        node: dict, decision: bool, reason: str, depth_level: int
    ) -> None:
        if not flow_trace:
            return
        flow_id = str(node.get("FlowID") or "").strip()
        cond = _format_logic_blob(
            node.get("BranchLogic"),
            questions=questions,
            translation_ctx=content.translation_ctx,
        )
        taken = "THEN" if decision else "ELSE"
        label = f"FlowID={flow_id}" if flow_id else "FlowID=?"
        flow_trace(f"[branch:{reason}] {label} -> {taken} | {cond}")

    def on_branch_then(_node: dict, depth_level: int) -> None:
        html_parts.append(_render_logic_line_html("THEN:", depth_level))

    def on_branch_else(_node: dict, depth_level: int) -> None:
        html_parts.append(_render_logic_line_html("ELSE:", depth_level))

    def on_branch_end(_node: dict, depth_level: int) -> None:
        html_parts.append(_render_logic_line_html("END BRANCH", depth_level))

    def on_end_survey(node: dict, depth_level: int) -> None:
        html_parts.append(_render_end_survey_html(node, content, depth_level))

    def on_unknown(node: dict, depth_level: int) -> None:
        node_type = str(node.get("Type") or "").strip()
        if node_type:
            html_parts.append(
                _render_system_note_html(f"FLOW NODE: {node_type}", depth_level)
            )

    handlers = FlowTraversalHandlers(
        on_block=on_block,
        on_group=on_group,
        on_embedded_data=on_embedded_data,
        on_web_service=on_web_service,
        on_randomizer=on_randomizer,
        on_branch_decision=on_branch_decision,
        on_branch_open=on_branch_open,
        on_branch_then=on_branch_then,
        on_branch_else=on_branch_else,
        on_branch_end=on_branch_end,
        on_end_survey=on_end_survey,
        on_unknown=on_unknown,
    )

    walk_flow(
        flow_list=flow_list,
        handlers=handlers,
        edf_overrides=content.edf_overrides,
        asked_qids=asked_qids,
        depth=depth,
        eval_branch=_eval_boolean_expression,
        eval_branch_with_asked=_eval_boolean_expression_with_unasked_selected_false,
    )

    return "".join(html_parts)


def _render_block_html(
    *,
    block_id: str,
    blocks: dict,
    questions: dict,
    content: ExportContent,
    asked_qids: set[str] | None,
    depth: int,
    flow_trace: Callable[[str], None] | None,
) -> str:
    """Render a single block with its questions in HTML."""
    block = blocks.get(block_id)
    if not isinstance(block, dict):
        return _render_system_note_html(
            f"BLOCK {block_id} (not found in Blocks)", depth
        )

    block_desc = (block.get("Description") or "").strip() or "(no description)"
    elements = block.get("BlockElements") or []
    if not isinstance(elements, list):
        elements = []

    # Collect QIDs from block elements
    block_qids_in_order: list[str] = []
    for elem in elements:
        if not isinstance(elem, dict):
            continue
        if elem.get("Type") == "Question":
            qid = elem.get("QuestionID")
            if qid and qid in content.active_qids:
                block_qids_in_order.append(str(qid))

    # Filter questions based on display logic when EDF overrides are provided
    render_qids: list[str] = list(block_qids_in_order)
    if content.edf_overrides and asked_qids is not None:
        asked_sim = set(asked_qids)
        render_qids = []
        for qid in block_qids_in_order:
            visible = _eval_question_display_logic_visibility(
                questions.get(qid) or {},
                questions=questions,
                edf_overrides=content.edf_overrides,
                asked_qids=asked_sim,
            )
            if visible is False:
                if flow_trace:
                    flow_trace(
                        f"[display_logic] Block {block_id} hides QID {qid} (display logic false)"
                    )
                continue
            render_qids.append(qid)
            asked_sim.add(qid)

    # If no questions to render after filtering, skip the block
    if not render_qids:
        if flow_trace and content.edf_overrides:
            flow_trace(
                f"[block_drop] Block {block_id} dropped (all questions hidden by display logic)"
            )
        return ""

    # Block header
    indent = f'style="margin-left:{depth * 20}px"'
    html = f'<div class="block" {indent}>\n'
    html += f'<div class="block-header">BLOCK START: {_escape_html(block_desc)} ({_escape_html(block_id)})</div>\n'

    # Render questions and page breaks in this block
    for elem in elements:
        if not isinstance(elem, dict):
            continue
        if elem.get("Type") == "Question":
            qid = elem.get("QuestionID")
            if qid and qid in render_qids:
                q_html = _render_question_html_full(qid, questions, content, depth)
                if q_html:
                    html += q_html
                    if asked_qids is not None:
                        asked_qids.add(qid)
        elif elem.get("Type") == "Page Break":
            html += f'<div class="system-note" style="margin-left:{(depth + 1) * 20}px">--- PAGE BREAK ---</div>\n'

    html += "</div>\n"
    return html


def _render_question_html_full(
    qid: str, questions: dict, content: ExportContent, depth: int
) -> str:
    """Render a complete question with all details (for flow traversal)."""
    q = questions.get(qid)
    if not isinstance(q, dict):
        return ""

    qtype = str(q.get("QuestionType") or "").strip()
    export_tag = str(q.get("DataExportTag") or "").strip() or qid

    # Get question type abbreviation
    qt_abbrev = _question_type_abbrev(q)

    # Add randomization marker (consistent with Word export)
    if q.get("Randomization"):
        qt_abbrev = (qt_abbrev + "+R").strip()

    # System/technical questions are noise for translators; represent them compactly.
    qtype_norm = qtype.lower()
    if is_system_question_type(qtype):
        if qtype_norm == "timing":
            block_label = "Timing Block"
        elif qtype_norm in {"meta", "metainfo"}:
            block_label = "Meta Block"
        elif qtype_norm == "captcha":
            block_label = "Captcha Block"
        else:
            block_label = "System Block"
        indent = f'style="margin-left:{(depth + 1) * 20}px"'
        html = f'<div class="question" {indent}>\n'
        html += '<h3 class="question-header">'
        html += f'<span class="qid">[{_escape_html(qid)}]</span>'
        html += f'<span class="qid">[{_escape_html(qt_abbrev)}]</span>'
        html += f" {_escape_html(export_tag)}"
        html += "</h3>\n"
        html += f'<div class="annotation">{_escape_html(block_label)}</div>\n'
        html += "</div>\n"
        return html

    # Question text (with translation if applicable)
    question_text = str(q.get("QuestionText") or "").strip()
    if content.translation_ctx:
        key = content.translation_ctx.key_for_question_text(qid)
        translated = content.translation_ctx.target_map.get(key)
        if translated:
            question_text = translated
    question_text = _sanitize_html_for_weasyprint(question_text)

    # Validation marker
    validation = q.get("Validation") or {}
    settings = validation.get("Settings") or {} if isinstance(validation, dict) else {}
    force = settings.get("ForceResponse") if isinstance(settings, dict) else None
    validation_marker = ""
    if force is not None:
        s = str(force).strip()
        if s.upper() == "ON":
            validation_marker = " *"
        elif s == "RequestResponse":
            validation_marker = " +"

    # Check for QuestionJS (inline or external file, consistent with Word export)
    has_js = bool((q.get("QuestionJS") or "").strip()) or bool(
        content.qid_to_js.get(qid)
    )

    indent = f'style="margin-left:{(depth + 1) * 20}px"'

    html = f'<div class="question" {indent}>\n'
    html += '<h3 class="question-header">'
    html += f'<span class="qid">[{_escape_html(qid)}]</span>'
    html += f'<span class="qid">[{_escape_html(qt_abbrev)}]</span>'
    if has_js:
        html += '<span class="qid">[JS]</span>'
    html += f" {_escape_html(export_tag)}{_escape_html(validation_marker)}"
    html += "</h3>\n"

    # Display logic
    logic = q.get("DisplayLogic")
    if logic:
        logic_str = _format_logic_blob(
            logic, questions=questions, translation_ctx=content.translation_ctx
        )
        if logic_str:
            styled_logic = _html_logic_with_highlights(f"DISPLAY IF: {logic_str}")
            html += f'<div class="logic">{styled_logic}</div>\n'

    # Question text
    html += f'<div class="question-text">{question_text}</div>\n'

    # Choices/Answers with proper ordering
    choices = q.get("Choices", {}) or {}
    answers = q.get("Answers", {}) or {}
    slider_like = qtype in {"Slider", "CS"}
    is_sbs_matrix = _is_sbs_matrix_question(q)

    def _ordered_ids(mapping: dict, *, order_key: str) -> list[str]:
        order = q.get(order_key) or []
        ids = [str(x) for x in order if str(x) in mapping]
        for k in mapping.keys():
            sk = str(k)
            if sk not in ids:
                ids.append(sk)
        return ids

    def _renderable_html(s: str) -> bool:
        return _has_renderable_text(_trim_html_edges(s))

    if is_sbs_matrix:
        ordered_choice_ids = _ordered_ids(choices, order_key="ChoiceOrder")

        items = []
        for choice_id in ordered_choice_ids:
            choice_data = choices.get(choice_id)
            if not isinstance(choice_data, dict):
                continue
            display = _coerce_display_text(choice_data.get("Display")).strip()
            if content.translation_ctx:
                key = content.translation_ctx.key_for_choice(qid, choice_id)
                translated = content.translation_ctx.target_map.get(key)
                if translated:
                    display = translated
            display = _sanitize_html_for_weasyprint(display)
            if not _renderable_html(display):
                continue
            items.append(
                f"<li><strong>[{_escape_html(choice_id)}]</strong> {display}</li>"
            )
        if items:
            html += '<div class="choice-label"><strong>Statements:</strong></div>\n'
            html += f'<ul class="choices">{"".join(items)}</ul>\n'

        additional = q.get("AdditionalQuestions") or {}
        if isinstance(additional, dict):
            for column_id in _ordered_numeric_string_ids(additional):
                column = additional.get(column_id)
                if not isinstance(column, dict):
                    continue
                col_text = _coerce_display_text(column.get("QuestionText"))
                if content.translation_ctx:
                    col_key = content.translation_ctx.key_for_sbs_column_question_text(
                        qid, str(column_id)
                    )
                    translated = content.translation_ctx.target_map.get(col_key)
                    if translated:
                        col_text = translated
                col_text = _sanitize_html_for_weasyprint(str(col_text))

                answer_map = column.get("Answers") or {}
                if not isinstance(answer_map, dict) or not answer_map:
                    answer_map = q.get("Answers") or {}
                    if not isinstance(answer_map, dict) or not answer_map:
                        continue
                    answer_order = q.get("AnswerOrder")
                else:
                    answer_order = column.get("AnswerOrder")
                ordered_answer_ids = _ordered_numeric_string_ids(
                    answer_map, answer_order
                )

                col_items = []
                for ans_id in ordered_answer_ids:
                    ans_data = answer_map.get(ans_id) if isinstance(answer_map, dict) else None
                    if not isinstance(ans_data, dict):
                        continue
                    ans_display = _coerce_display_text(ans_data.get("Display")).strip()
                    if content.translation_ctx:
                        ans_key = content.translation_ctx.key_for_sbs_column_answer(
                            qid, str(column_id), str(ans_id)
                        )
                        translated = content.translation_ctx.target_map.get(ans_key)
                        if translated:
                            ans_display = translated
                    ans_display = _sanitize_html_for_weasyprint(ans_display)
                    if not _renderable_html(ans_display):
                        continue
                    col_items.append(
                        f"<li><strong>[{_escape_html(str(column_id))}.{_escape_html(ans_id)}]</strong> {ans_display}</li>"
                    )
                if col_items:
                    html += (
                        f'<div class="choice-label"><strong>Column '
                        f'{_escape_html(str(column_id))}:</strong></div>\n'
                    )
                    html += f'<div class="question-text">{col_text}</div>\n'
                    html += f'<ul class="answers">{"".join(col_items)}</ul>\n'

    elif isinstance(choices, dict) and choices:
        ordered_choice_ids = _ordered_ids(choices, order_key="ChoiceOrder")

        items = []
        for choice_id in ordered_choice_ids:
            choice_data = choices.get(choice_id)
            if not isinstance(choice_data, dict):
                continue
            display = _coerce_display_text(choice_data.get("Display")).strip()
            if content.translation_ctx:
                key = content.translation_ctx.key_for_choice(qid, choice_id)
                translated = content.translation_ctx.target_map.get(key)
                if translated:
                    display = translated
            display = _sanitize_html_for_weasyprint(display)
            if (qtype == "Matrix" or slider_like) and not _renderable_html(display):
                continue
            items.append(
                f"<li><strong>[{_escape_html(choice_id)}]</strong> {display}</li>"
            )
        if items:
            # Slider and CS(+R) behave like matrix-style rows.
            label = "Statements" if qtype == "Matrix" or slider_like else "Choices"
            html += f'<div class="choice-label"><strong>{label}:</strong></div>\n'
            html += f'<ul class="choices">{"".join(items)}</ul>\n'

    labels = q.get("Labels") or {}
    if slider_like and isinstance(labels, dict) and labels:
        ordered_label_ids = sorted(
            [str(k) for k in labels.keys()],
            key=lambda x: int(x) if str(x).isdigit() else str(x),
        )
        items = []
        for lid in ordered_label_ids:
            lab = labels.get(lid)
            if not isinstance(lab, dict):
                continue
            display = _coerce_display_text(lab.get("Display")).strip()
            if content.translation_ctx:
                key = content.translation_ctx.key_for_label(qid, lid)
                translated = content.translation_ctx.target_map.get(key)
                if translated:
                    display = translated
            display = _sanitize_html_for_weasyprint(display)
            if not _renderable_html(display):
                continue
            items.append(f"<li><strong>[{_escape_html(lid)}]</strong> {display}</li>")
        if items:
            html += '<div class="choice-label"><strong>Labels:</strong></div>\n'
            html += f'<ul class="choices">{"".join(items)}</ul>\n'

    if isinstance(answers, dict) and answers:
        ordered_answer_ids = _ordered_ids(answers, order_key="AnswerOrder")

        items = []
        for ans_id in ordered_answer_ids:
            ans_data = answers.get(ans_id)
            if not isinstance(ans_data, dict):
                continue
            display = _coerce_display_text(ans_data.get("Display")).strip()
            if content.translation_ctx:
                key = content.translation_ctx.key_for_answer(qid, ans_id)
                translated = content.translation_ctx.target_map.get(key)
                if translated:
                    display = translated
            display = _sanitize_html_for_weasyprint(display)
            items.append(
                f"<li><strong>[{_escape_html(ans_id)}]</strong> {display}</li>"
            )
        if items:
            # Sliders and Constant Sum sliders should present a scale (not generic "Options").
            label = "Scale" if qtype == "Matrix" or slider_like else "Options"
            html += f'<div class="choice-label"><strong>{label}:</strong></div>\n'
            html += f'<ul class="answers">{"".join(items)}</ul>\n'
    elif slider_like:
        cfg = q.get("Configuration") or {}
        if isinstance(cfg, dict):
            lines: list[str] = []
            min_v = cfg.get("CSSliderMin")
            max_v = cfg.get("CSSliderMax")
            grid = cfg.get("GridLines")
            dec = cfg.get("NumDecimals")
            snap = cfg.get("SnapToGrid")
            show_val = cfg.get("ShowValue")
            if min_v is not None or max_v is not None:
                lo = "" if min_v is None else str(min_v)
                hi = "" if max_v is None else str(max_v)
                if lo and hi:
                    lines.append(f"Range: {lo}–{hi}")
                else:
                    lines.append(f"Range: {lo or '—'}–{hi or '—'}")
            if grid is not None:
                lines.append(f"GridLines: {grid}")
            if dec is not None:
                lines.append(f"Decimals: {dec}")
            if snap is not None:
                lines.append(f"SnapToGrid: {snap}")
            if show_val is not None:
                lines.append(f"ShowValue: {show_val}")
            if lines:
                items = "".join(f"<li>{_escape_html(line)}</li>" for line in lines)
                html += '<div class="choice-label"><strong>Scale:</strong></div>\n'
                html += f'<ul class="answers">{items}</ul>\n'

    # Validation details
    if isinstance(settings, dict) and settings:
        val_items = []
        for k in sorted(settings.keys()):
            if k == "ForceResponse":
                continue
            v = settings.get(k)
            if v is not None:
                val_items.append(f"{k}={v}")
        if val_items:
            html += f'<div class="annotation">Validation: {_escape_html("; ".join(val_items))}</div>\n'

    # JavaScript User-Visible Strings
    if has_js and content.include_js_strings:
        js_code = str(q.get("QuestionJS") or "")
        if js_code:
            js_lang = content.render_language or content.base_language
            js_strings = _extract_js_strings(js_code, target_language=js_lang)
            if js_strings:
                html += '<div class="js-strings"><strong>JavaScript User-Visible Strings:</strong></div>\n'
                html += '<ul class="js-strings-list">\n'
                for s in js_strings:
                    html += f"  <li>{_escape_html(s)}</li>\n"
                html += "</ul>\n"

    html += "</div>\n"
    return html


def _render_system_note_html(text: str, depth: int) -> str:
    """Render a system note (for flow annotations)."""
    indent = f'style="margin-left:{depth * 20}px"'
    return f'<div class="system-note" {indent}>{_escape_html(text)}</div>\n'


def _render_logic_line_html(text: str, depth: int) -> str:
    """Render a logic line (for branch conditions) with styled markers."""
    indent = f'style="margin-left:{depth * 20}px"'
    styled_text = _html_logic_with_highlights(text)
    return f'<div class="logic" {indent}>{styled_text}</div>\n'


def _html_logic_with_highlights(text: str) -> str:
    """Convert logic text with [[Q]], [[A]], [[OP]], [[EDF]] markers to styled HTML."""
    s = str(text or "")
    if not s:
        return ""

    result = []
    pos = 0
    for m in _LOGIC_MARK_RE.finditer(s):
        # Add text before marker
        if m.start() > pos:
            result.append(_escape_html(s[pos : m.start()]))

        kind = m.group(1)
        val = m.group(2) or ""

        # Style based on marker type (matching DOCX colors)
        if kind == "Q":
            # Question - black, not bold
            result.append(f'<span style="color:#000000">{_escape_html(val)}</span>')
        elif kind == "A":
            # Answer - blue, bold
            result.append(
                f'<span style="color:#0000C8;font-weight:bold">{_escape_html(val)}</span>'
            )
        elif kind == "EDF":
            # Embedded Data - green, bold
            result.append(
                f'<span style="color:#008000;font-weight:bold">{_escape_html(val)}</span>'
            )
        elif kind == "OP":
            # Operator - red, bold
            result.append(
                f'<span style="color:#C80000;font-weight:bold">{_escape_html(val)}</span>'
            )
        else:
            result.append(_escape_html(val))

        pos = m.end()

    # Add remaining text after last marker
    if pos < len(s):
        result.append(_escape_html(s[pos:]))

    return "".join(result)


def _render_embedded_data_html(
    node: dict, edf_overrides: dict[str, str] | None, depth: int
) -> str:
    """Render an EmbeddedData flow node."""
    fields = node.get("EmbeddedData") or []
    if not isinstance(fields, list):
        return ""

    items = []
    for field in fields:
        if not isinstance(field, dict):
            continue
        f_name = str(field.get("Field") or "").strip()
        f_val = str(field.get("Value") or "").strip()
        if not f_name:
            continue

        # Highlight if overridden by EDF
        if edf_overrides and f_name in edf_overrides:
            items.append(
                f"  <strong class='edf-token'>{_escape_html(f_name)}</strong> = {_escape_html(f_val)} (overridden by --edf)"
            )
        else:
            items.append(f"  {_escape_html(f_name)} = {_escape_html(f_val)}")

    if items:
        indent = f'style="margin-left:{depth * 20}px"'
        return f'<div class="system-note" {indent}>EMBEDDED DATA:<br>{"<br>".join(items)}</div>\n'
    return ""


def _render_web_service_html(node: dict, depth: int) -> str:
    """Render a WebService flow node with details."""
    flow_id = str(node.get("FlowID") or "").strip()
    method = str(node.get("Method") or "").strip()
    url = str(node.get("URL") or "").strip()
    content_type = str(node.get("ContentType") or "").strip()

    indent = f'style="margin-left:{depth * 20}px"'
    html = f'<div class="system-note" {indent}>'
    html += "<strong>WEB SERVICE</strong>"
    if method or url:
        html += f": <strong>{_escape_html(method)}</strong> {_escape_html(url)}"
    if flow_id:
        html += f" (FlowID={_escape_html(flow_id)})"
    if content_type:
        html += f"<br>Content-Type: {_escape_html(content_type)}"
    html += "</div>\n"
    return html


def _render_randomizer_html(node: dict, depth: int) -> str:
    """Render a BlockRandomizer flow node."""
    subset = node.get("SubSet")
    even = node.get("EvenPresentation")

    text = "BLOCK RANDOMIZER"
    if subset:
        text = f"{text}: subset={subset}"
    if even:
        text = f"{text}, even presentation"

    return _render_system_note_html(text, depth)


def _render_end_survey_html(node: dict, content: ExportContent, depth: int) -> str:
    """Render an EndSurvey flow node with message embedding."""
    opts = node.get("Options") or {}
    term = str(opts.get("SurveyTermination") or "").strip()
    flow_id = str(node.get("FlowID") or "").strip()
    lib_id = str(opts.get("EOSMessageLibrary") or "").strip()
    msg_id = str(opts.get("EOSMessage") or "").strip()

    label = f"END SURVEY: {term}" if term else "END SURVEY"
    details = []
    if flow_id:
        details.append(f"FlowID={flow_id}")
    if lib_id:
        details.append(f"EOSMessageLibrary={lib_id}")
    if msg_id:
        details.append(f"EOSMessage={msg_id}")
    if details:
        label = f"{label} ({', '.join(details)})"

    html = _render_system_note_html(label, depth)

    # Embed message content if available
    if term == "DisplayMessage" and lib_id and msg_id:
        msg_data = _read_eos_message_from_disk(lib_id, msg_id)
        if not msg_data:
            return html

        msg_desc = str(msg_data.get("description") or "").strip()
        messages = msg_data.get("messages") or {}
        if not isinstance(messages, dict) or not messages:
            return html

        base_key = _normalize_lang_code(content.base_language).lower()
        target_key = (
            _normalize_lang_code(content.render_language).lower()
            if content.render_language
            else ""
        )

        selected_keys = sorted(messages.keys(), key=str)
        if content.compare_to_base:
            preferred = [k for k in [base_key, target_key] if k and k in messages]
            if preferred:
                selected_keys = preferred
            elif base_key in messages:
                selected_keys = [base_key]
            elif target_key in messages:
                selected_keys = [target_key]
        else:
            if target_key and target_key in messages:
                selected_keys = [target_key]
            elif base_key in messages:
                selected_keys = [base_key]

        if not selected_keys:
            return html

        indent = f'style="margin-left:{(depth + 1) * 20}px"'
        html += f'<div class="system-note" {indent}>'
        if msg_desc:
            html += f"<div><strong>Message:</strong> {_escape_html(msg_desc)}</div>"

        # Render message variants in a simple table (monolingual: 1 col, bilingual: 2 cols).
        headers = "".join(
            f"<th>{_escape_html(_normalize_lang_code(k))}</th>" for k in selected_keys
        )
        cells = "".join(
            f"<td>{_sanitize_html_for_weasyprint(str(messages.get(k) or ''))}</td>"
            for k in selected_keys
        )
        html += f"""
        <table class="eos-message">
            <tr>{headers}</tr>
            <tr>{cells}</tr>
        </table>
        """
        html += "</div>\n"

    return html


def _render_question_html(qid: str, content: ExportContent) -> str:
    """Render a single question to HTML."""
    result = content.survey_payload.get("result", {}) or {}
    questions = result.get("Questions", {}) or {}
    q = questions.get(qid)

    if not isinstance(q, dict):
        return ""

    qtype = str(q.get("QuestionType") or "").strip()
    export_tag = str(q.get("DataExportTag") or "").strip() or qid

    # Get question text (with translation if applicable)
    question_text = str(q.get("QuestionText") or "").strip()
    if content.translation_ctx:
        key = content.translation_ctx.key_for_question_text(qid)
        translated = content.translation_ctx.target_map.get(key)
        if translated:
            question_text = translated

    # Sanitize HTML to work around WeasyPrint bugs
    question_text = _sanitize_html_for_weasyprint(question_text)

    # Question text rendering
    text_html = f'<div class="question-text">{question_text}</div>'

    # Choices/Answers rendering (simplified for MVP)
    choices_html = ""
    choices = q.get("Choices", {}) or {}
    if isinstance(choices, dict) and choices:
        items = []
        for choice_id, choice_data in sorted(choices.items()):
            if not isinstance(choice_data, dict):
                continue
            display = str(choice_data.get("Display") or "").strip()
            if content.translation_ctx:
                key = content.translation_ctx.key_for_choice(qid, choice_id)
                translated = content.translation_ctx.target_map.get(key)
                if translated:
                    display = translated
            # Sanitize HTML to work around WeasyPrint bugs
            display = _sanitize_html_for_weasyprint(display)
            items.append(f"<li>{display}</li>")

        if items:
            choices_html = f'<ul class="choices">{"".join(items)}</ul>'

    return f"""
    <div class="question" id="question-{_escape_html(qid)}">
        <h3 class="question-header"><span class="qid">{_escape_html(qid)}</span> / {_escape_html(export_tag)} [{_escape_html(qtype)}]</h3>
        {text_html}
        {choices_html}
    </div>
    """


def _render_survey_content_html(content: ExportContent) -> str:
    """Render main survey content section HTML with full flow traversal."""
    result = content.survey_payload.get("result", {}) or {}
    blocks = result.get("Blocks", {}) or {}
    questions = result.get("Questions", {}) or {}
    flow = result.get("SurveyFlow") or {}
    flow_list = flow.get("Flow") or []

    if not isinstance(flow_list, list):
        return """
        <section>
            <h1>SURVEY CONTENT (Flow Order)</h1>
            <p class="system-note">(SurveyFlow missing or malformed.)</p>
        </section>
        """

    if not content.active_qids:
        return """
        <section>
            <h1>SURVEY CONTENT (Flow Order)</h1>
            <p class="system-note">(No active questions detected in SurveyFlow non-Trash blocks.)</p>
        </section>
        """

    asked_qids: set[str] | None = set() if content.edf_overrides else None

    content_html = _traverse_flow_html(
        flow_list=flow_list,
        blocks=blocks,
        questions=questions,
        content=content,
        asked_qids=asked_qids,
        depth=0,
        flow_trace=content.flow_trace,
    )

    return f"""
    <section>
        <h1>SURVEY CONTENT (Flow Order)</h1>
        {content_html}
    </section>
    """


def _render_external_surfaces_html(content: ExportContent) -> str:
    """Render external translation surfaces section HTML."""
    # Questions with JS
    js_questions = [
        qid
        for qid, js_file in content.qid_to_js.items()
        if js_file and qid in content.active_qids
    ]

    js_html = ""
    if js_questions:
        items = "".join(
            f'<li><span class="qid">{_escape_html(qid)}</span> → {_escape_html(content.qid_to_js[qid])}</li>'
            for qid in sorted(js_questions)[:10]
        )
        if len(js_questions) > 10:
            items += f"<li><em>... and {len(js_questions) - 10} more</em></li>"
        js_html = f"<ul>{items}</ul>"
    else:
        js_html = "<p>(No questions with QuestionJS found.)</p>"

    return f"""
    <section class="external-surface">
        <h1>EXTERNAL TRANSLATION SURFACES</h1>
        <h2>Questions with QuestionJS</h2>
        <p><em>These questions use JavaScript that may contain user-visible strings. Review survey_js/ files.</em></p>
        {js_html}
    </section>
    """


def _build_pdf_html_template(content: ExportContent) -> str:
    """Build complete HTML template for PDF export.

    Args:
        content: Prepared export content

    Returns:
        Complete HTML document string
    """
    lang_code = content.render_language or "en"

    header_html = _render_header_html(content)
    translation_summary_html = _render_translation_summary_html(content)
    coverage_html = _render_coverage_summary_html(content)
    legend_html = _render_type_legend_html(content)
    content_html = _render_survey_content_html(content)
    external_html = _render_external_surfaces_html(content)

    # Mermaid diagram (if exists)
    mermaid_html = ""
    if content.mermaid_image_path and content.mermaid_image_path.exists():
        mermaid_html = f"""
        <section class="flow-diagram">
            <h1>FLOW DIAGRAM</h1>
            <img src="{content.mermaid_image_path}" alt="Survey Flow Diagram" />
        </section>
        """

    return f"""<!DOCTYPE html>
<html lang="{_escape_html(lang_code)}">
<head>
    <meta charset="UTF-8">
    <title>Translation Export: {_escape_html(content.survey_id)}</title>
</head>
<body>
    {header_html}
    {translation_summary_html}
    {coverage_html}
    {legend_html}
    {mermaid_html}
    {content_html}
    {external_html}
</body>
</html>
"""


def _render_to_pdf(content: ExportContent) -> Path:
    """Render prepared export content to a PDF file.

    Args:
        content: Prepared export content

    Returns:
        Path to the saved PDF file
    """
    try:
        from weasyprint import CSS, HTML
    except ImportError as e:
        raise ModuleNotFoundError(
            "PDF export requires WeasyPrint. Install it with: pip install 'qsync[pdf]'\n"
            "Note: WeasyPrint requires system libraries (cairo, pango, gdk-pixbuf, libffi).\n"
            "On macOS: brew install cairo pango gdk-pixbuf libffi\n"
            "On Ubuntu/Debian: apt-get install libcairo2 libpango-1.0-0 libgdk-pixbuf2.0-0 libffi-dev\n"
            "See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation"
        ) from e
    except OSError as e:
        # WeasyPrint installed but system libraries missing
        if "libgobject" in str(e) or "libcairo" in str(e) or "libpango" in str(e):
            raise RuntimeError(
                "PDF export: WeasyPrint is installed but cannot load required system libraries.\n"
                "On macOS: brew install cairo pango gdk-pixbuf libffi\n"
                "On Ubuntu/Debian: apt-get install libcairo2 libpango-1.0-0 libgdk-pixbuf2.0-0\n"
                "See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation"
            ) from e
        raise

    content.output_path.parent.mkdir(parents=True, exist_ok=True)

    # Generate HTML and CSS
    html_content = _build_pdf_html_template(content)
    css_content = _build_pdf_css()

    # Render to PDF
    HTML(string=html_content).write_pdf(
        content.output_path, stylesheets=[CSS(string=css_content)]
    )

    return content.output_path


def _env_flag_disabled(name: str) -> bool:
    v = (os.environ.get(name) or "").strip().lower()
    return v in {"0", "false", "no", "off"}


def _render_mermaid_to_png(code: str, out_path: Path) -> None:
    """Render Mermaid code to a PNG image using a lightweight remote renderer."""

    try:
        import requests
    except Exception as e:  # pragma: no cover
        raise ModuleNotFoundError(
            "Mermaid rendering requires the 'requests' package."
        ) from e

    # mermaid.ink expects a urlsafe base64 payload without padding.
    b64 = (
        base64.urlsafe_b64encode((code or "").encode("utf-8"))
        .decode("ascii")
        .rstrip("=")
    )
    url = f"https://mermaid.ink/img/{b64}"
    try:
        resp = requests.get(url, timeout=30)
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Failed to render Mermaid diagram (network error). "
            "This export embeds a rendered Mermaid chart and requires network access."
        ) from e

    if resp.status_code != 200 or not resp.content:
        raise RuntimeError(
            f"Failed to render Mermaid diagram (HTTP {resp.status_code})."
        )
    out_path.write_bytes(resp.content)


def _format_question_validation_line(question: dict) -> str:
    validation = question.get("Validation") or {}
    if not isinstance(validation, dict):
        return ""
    settings = validation.get("Settings") or {}
    if not isinstance(settings, dict) or not settings:
        return ""

    items: List[str] = []
    force = settings.get("ForceResponse")
    if force is not None:
        s = str(force).strip()
        required = s.upper() not in {"OFF", "FALSE", "0", "NONE", ""}
        items.append(f"Required: {'YES' if required else 'NO'} (ForceResponse={s})")

    for k in sorted(settings.keys(), key=str):
        if k == "ForceResponse":
            continue
        v = settings.get(k)
        if v is None:
            continue
        items.append(f"{k}={v}")

    if not items:
        return ""
    if len(items) == 1:
        return f"Validation: {items[0]}"
    return "Validation: " + "; ".join(items)


def _looks_like_routing_field(field: str) -> bool:
    f = (field or "").strip().lower()
    if not f:
        return False
    return (
        "lang" in f
        or "language" in f
        or "country" in f
        or "cntry" in f
        or f.endswith("_label")
    )


def _question_validation_marker(question: dict) -> str:
    """Return '*' for force response, '+' for request response, '' otherwise."""

    validation = question.get("Validation") or {}
    if not isinstance(validation, dict):
        return ""
    settings = validation.get("Settings") or {}
    if not isinstance(settings, dict):
        return ""
    force = settings.get("ForceResponse")
    if force is None:
        return ""
    s = str(force).strip()
    if s.upper() == "ON":
        return "*"
    if s == "RequestResponse":
        return "+"
    return ""


def _add_question_metadata_in_cell(
    cell,
    *,
    qid: str,
    qt_abbrev: str,
    has_js: bool,
    export_tag: str,
    validation_marker: str,
) -> None:
    """Render the compact question metadata line inside a table cell.

    Format:
      [QID][QT][JS] ExportTag *
    """
    p = _container_add_paragraph(cell)

    def add_token(token: str) -> None:
        r = p.add_run(token)
        r.bold = True
        _style_qid_run(r)

    add_token(f"[{str(qid or '').strip()}]")
    add_token(f"[{str(qt_abbrev or '').strip()}]")
    if has_js:
        add_token("[JS]")

    # Export tag and marker: from here on, explicitly Arial.
    tag_part = str(export_tag or "").strip()
    r_tag = p.add_run((" " + tag_part) if tag_part else " ")
    r_tag.bold = True
    try:
        r_tag.font.name = "Arial"
    except Exception:
        pass

    if validation_marker:
        r2 = p.add_run(f" {validation_marker}")
        r2.bold = True
        try:
            r2.font.name = "Arial"
        except Exception:
            pass


def _add_logic_line(container, text: str, *, depth: int) -> None:
    """Render branch/display logic in monospace, small, with highlighted segments."""

    p = _container_add_paragraph(container)
    _set_indent(p, depth=depth)
    _add_logic_runs_with_highlights(p, text)
    _shade_logic_paragraph(p)
    _style_logic_paragraph_spacing(p, container=container)


def _style_logic_paragraph_spacing(paragraph, *, container) -> None:
    """Ensure logic lines have a visible gap after them in doc-level flow output.

    Avoid adding extra spacing inside table cells, where the table structure already
    provides separation.
    """

    try:
        from docx.shared import Pt
    except Exception:
        return

    is_cell = container.__class__.__name__ == "_Cell"
    if is_cell:
        return
    try:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
    except Exception:
        return


def _style_logic_run(run) -> None:
    try:
        from docx.shared import Pt, RGBColor

        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(200, 0, 0)
    except Exception:
        return


def _style_logic_token_run(
    run, *, rgb: tuple[int, int, int], bold: bool = False
) -> None:
    try:
        from docx.shared import Pt, RGBColor

        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(int(rgb[0]), int(rgb[1]), int(rgb[2]))
        run.bold = bool(bold)
    except Exception:
        return


def _add_logic_runs_with_highlights(paragraph, text: str) -> None:
    """Render logic lines with Q/A/EDF segments highlighted via markers."""

    s = str(text or "")
    if not s:
        return

    pos = 0
    for m in _LOGIC_MARK_RE.finditer(s):
        if m.start() > pos:
            run = paragraph.add_run(s[pos : m.start()])
            _style_logic_token_run(run, rgb=(200, 0, 0))

        kind = m.group(1)
        val = m.group(2) or ""
        rgb = (200, 0, 0)
        bold = False
        if kind == "Q":
            rgb = (0, 0, 0)
        elif kind == "A":
            rgb = (0, 0, 200)
            bold = True
        elif kind == "EDF":
            rgb = (0, 128, 0)
            bold = True
        elif kind == "OP":
            rgb = (200, 0, 0)
            bold = True

        run2 = paragraph.add_run(val)
        _style_logic_token_run(run2, rgb=rgb, bold=bold)
        pos = m.end()

    if pos < len(s):
        run = paragraph.add_run(s[pos:])
        _style_logic_token_run(run, rgb=(200, 0, 0))


def _shade_logic_paragraph(paragraph) -> None:
    """Light gray background for logic lines (best effort)."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    try:
        ppr = paragraph._p.get_or_add_pPr()  # type: ignore[attr-defined]
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        ppr.append(shd)
    except Exception:
        return


def _style_table_label_paragraph(paragraph) -> None:
    """Spacing rules for labels that introduce a table (attach label to table below)."""

    try:
        from docx.shared import Pt

        # Do not add space-before here; rely on the preceding content's space-after
        # (docDefaults), and add explicit spacing after tables instead.
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
    except Exception:
        return


def _shade_block_header_paragraph(paragraph) -> None:
    """Gray background + spacing for block start lines (best effort)."""

    try:
        from docx.shared import Pt

        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.keep_with_next = True
    except Exception:
        pass

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    try:
        ppr = paragraph._p.get_or_add_pPr()  # type: ignore[attr-defined]
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        ppr.append(shd)
    except Exception:
        return


def _shade_panel_header_paragraph(paragraph) -> None:
    """Gray background header styling for table-based panels (best effort)."""

    try:
        from docx.shared import Pt

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
    except Exception:
        pass

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    try:
        ppr = paragraph._p.get_or_add_pPr()  # type: ignore[attr-defined]
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        ppr.append(shd)
    except Exception:
        return


def _style_edf_run(run) -> None:
    try:
        from docx.shared import RGBColor

        run.font.name = "Courier New"
        run.font.color.rgb = RGBColor(0, 128, 0)
    except Exception:
        return


def _add_doc_spacer_paragraph(doc, *, depth: int) -> None:
    """Add a small spacer paragraph after doc-level tables for consistent readability."""

    p = doc.add_paragraph("")
    _set_indent(p, depth=depth)
    try:
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Pt

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        # Avoid a full blank-line height; keep the spacer compact.
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(1)
    except Exception:
        return


def _add_block_header_leading_spacer(doc, *, depth: int) -> None:
    """Add a small spacer before block headers for consistent visual separation."""

    p = doc.add_paragraph("")
    _set_indent(p, depth=depth)
    try:
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Pt

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(1)
    except Exception:
        return


def _style_url_run(run) -> None:
    try:
        from docx.shared import RGBColor

        run.font.name = "Courier New"
        run.font.color.rgb = RGBColor(0, 0, 200)
    except Exception:
        return


def _add_text_with_edf_styling(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    rgb: tuple[int, int, int] | None = None,
) -> None:
    """Add text to a paragraph, styling ${e://Field/...} tokens as monospace green."""

    text = str(text or "")
    if not text:
        return

    pos = 0
    for m in _EDF_TOKEN_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if rgb is not None:
                _apply_run_rgb(run, rgb)
        token = m.group(0)
        run_t = paragraph.add_run(token)
        run_t.bold = bold
        run_t.italic = italic
        run_t.underline = underline
        _style_edf_run(run_t)
        pos = m.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold
        run.italic = italic
        run.underline = underline
        if rgb is not None:
            _apply_run_rgb(run, rgb)


def _apply_run_rgb(run, rgb: tuple[int, int, int]) -> None:
    try:
        from docx.shared import RGBColor

        r, g, b = rgb
        run.font.color.rgb = RGBColor(int(r), int(g), int(b))
    except Exception:
        return


def _extract_color_rgb(attrs) -> tuple[int, int, int] | None:
    """Parse color from HTML attributes (very small subset)."""

    if not attrs:
        return None
    color_val: str | None = None
    for k, v in attrs:
        if not k:
            continue
        lk = str(k).lower()
        if lk == "color" and v:
            color_val = str(v)
            break
        if lk == "style" and v:
            m = re.search(r"color\\s*:\\s*([^;]+)", str(v), flags=re.IGNORECASE)
            if m:
                color_val = m.group(1).strip()
                break
    if not color_val:
        return None
    s = color_val.strip().lower()
    if s.startswith("#") and len(s) in {4, 7}:
        if len(s) == 4:
            r = int(s[1] * 2, 16)
            g = int(s[2] * 2, 16)
            b = int(s[3] * 2, 16)
            return (r, g, b)
        r = int(s[1:3], 16)
        g = int(s[3:5], 16)
        b = int(s[5:7], 16)
        return (r, g, b)
    m = re.match(r"rgb\\(\\s*(\\d+)\\s*,\\s*(\\d+)\\s*,\\s*(\\d+)\\s*\\)", s)
    if m:
        return (int(m.group(1)), int(m.group(2)), int(m.group(3)))
    return None


def _extract_href(attrs) -> str | None:
    for k, v in attrs or []:
        if str(k).lower() == "href" and v:
            return str(v).strip()
    return None


_ALLOWED_INLINE_CSS_KEYS = {
    "color",
    "font-weight",
    "font-style",
    "text-decoration",
    "font-family",
    "font-size",
}


def _parse_inline_style(style_str: str | None) -> dict[str, str]:
    s = str(style_str or "").strip()
    if not s:
        return {}
    out: dict[str, str] = {}
    for part in s.split(";"):
        if ":" not in part:
            continue
        k, v = part.split(":", 1)
        k = k.strip().lower()
        v = v.strip()
        if not k or not v:
            continue
        out[k] = v
    return out


def _css_font_weight_is_bold(value: str) -> bool:
    v = str(value or "").strip().lower()
    if not v:
        return False
    if v == "bold":
        return True
    try:
        n = int(re.sub(r"[^0-9]", "", v) or "0")
        return n >= 600
    except Exception:
        return False


def _css_font_style_is_italic(value: str) -> bool:
    return str(value or "").strip().lower() == "italic"


def _css_text_decoration_is_underline(value: str) -> bool:
    v = str(value or "").strip().lower()
    return "underline" in v


def _css_font_family_is_monospace(value: str) -> bool:
    v = str(value or "").strip().lower()
    return "monospace" in v or "courier" in v


def _css_font_size_pt(value: str) -> int | None:
    v = str(value or "").strip().lower()
    if not v:
        return None
    m = re.match(r"^([0-9]+(?:\\.[0-9]+)?)\\s*(px|pt)?$", v)
    if not m:
        return None
    num = float(m.group(1))
    unit = m.group(2) or "px"
    if unit == "px":
        pt = int(round(num * 0.75))
    else:
        pt = int(round(num))
    if pt <= 0:
        return None
    # Clamp to avoid layout blowups.
    return max(6, min(28, pt))


def _href_normalize(href: str | None) -> str | None:
    h = str(href or "").strip()
    if not h:
        return None
    # Prefer web URLs over app-deep-links for translation docs.
    if h.lower().startswith("bsky:"):
        return "https://bsky.app/"
    return h


def _append_hyperlink(
    paragraph, *, url: str, text: str, bold: bool, italic: bool, size_pt: int | None
) -> None:
    """Append a clickable hyperlink to a paragraph (best-effort)."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except Exception:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = True
        try:
            from docx.shared import RGBColor

            run.font.color.rgb = RGBColor(0, 0, 200)
        except Exception:
            pass
        return

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000C8")
    r_pr.append(c)

    if size_pt is not None:
        sz = str(int(size_pt) * 2)
        e1 = OxmlElement("w:sz")
        e1.set(qn("w:val"), sz)
        r_pr.append(e1)
        e2 = OxmlElement("w:szCs")
        e2.set(qn("w:val"), sz)
        r_pr.append(e2)

    new_run.append(r_pr)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)  # type: ignore[attr-defined]


def _add_horizontal_rule(container, *, depth: int) -> None:
    """Add a horizontal divider line (best-effort Word-native border)."""

    p = _container_add_paragraph(container)
    _set_indent(p, depth=depth)
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        # Fallback: visible line of box-drawing chars
        p.add_run("────────────────────────────────────────")
        return

    # Always include a visible line so it renders consistently across Word viewers.
    line_run = p.add_run("────────────────────────────────────────")
    try:
        from docx.shared import RGBColor

        line_run.font.color.rgb = RGBColor(191, 191, 191)
    except Exception:
        pass

    try:
        ppr = p._p.get_or_add_pPr()  # type: ignore[attr-defined]
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "BFBFBF")
        pBdr.append(bottom)
        ppr.append(pBdr)
    except Exception:
        return


def _can_render_html_to_docx(html_str: str) -> bool:
    """Return True if we can safely render this HTML to Word (instead of RAW HTML).

    Conservative rules:
    - Reject scripts/styles/forms/iframes/svg/canvas and unknown tags.
    - Reject event handler attributes (on*).
    - Allow a small whitelist with limited attributes:
      - span/font: allow color + small deterministic CSS subset
      - a: allow href/rel/target/title
    """

    html_str = str(html_str or "")
    if not html_str.strip():
        return True

    from html.parser import HTMLParser

    allowed_tags = {
        "br",
        "p",
        "div",
        "details",
        "summary",
        "strong",
        "b",
        "em",
        "i",
        "u",
        "ul",
        "ol",
        "li",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "a",
        "sup",
        "sub",
        "code",
        "span",
        "font",
    }
    hard_reject = {
        "script",
        "style",
        "form",
        "input",
        "textarea",
        "select",
        "option",
        "button",
        "iframe",
        "svg",
        "canvas",
        "img",
        "video",
        "audio",
        "table",
    }

    ok = True

    def attrs_ok(tag: str, attrs) -> bool:
        for k, v in attrs or []:
            if not k:
                continue
            lk = str(k).lower()
            # Data/aria attrs are safe to ignore for Word rendering.
            if lk.startswith("data-") or lk.startswith("aria-"):
                continue
            if lk.startswith("on"):
                return False
            if tag in {"span", "font"}:
                if lk == "color":
                    if v and _extract_color_rgb([(k, v)]) is None:
                        return False
                    continue
                if lk == "style":
                    css = _parse_inline_style(str(v or ""))
                    # Only allow a deterministic subset of CSS.
                    for css_k in css.keys():
                        if css_k not in _ALLOWED_INLINE_CSS_KEYS:
                            return False
                    if (
                        "color" in css
                        and _extract_color_rgb([("style", f"color:{css['color']}")])
                        is None
                    ):
                        return False
                    if (
                        "font-size" in css
                        and _css_font_size_pt(css.get("font-size", "")) is None
                    ):
                        return False
                    # Other supported keys are applied best-effort.
                    continue
                if lk in {"class", "id"}:
                    continue
                return False
            if tag == "a":
                if lk == "href":
                    continue
                if lk in {"class", "id", "title", "target", "rel"}:
                    continue
                return False
            if tag == "details":
                if lk in {"open", "class", "id"}:
                    continue
                return False
            if tag == "summary":
                if lk in {"class", "id"}:
                    continue
                return False
            # default: ignore class/id, reject everything else
            if lk in {"class", "id"}:
                continue
            return False
        return True

    class Parser(HTMLParser):
        def handle_starttag(self, tag, attrs):
            nonlocal ok
            t = str(tag or "").lower()
            if t in hard_reject or t not in allowed_tags:
                ok = False
                return
            if not attrs_ok(t, attrs):
                ok = False

        def handle_startendtag(self, tag, attrs):
            self.handle_starttag(tag, attrs)

    try:
        Parser().feed(html_str)
    except Exception:
        return False
    return ok


def _style_edf_field_cell(cell) -> None:
    for p in getattr(cell, "paragraphs", []) or []:
        for run in p.runs:
            _style_edf_run(run)


def _set_table_column_widths_in(table, widths_in: list[float]) -> None:
    """Set fixed column widths for a python-docx table (best effort)."""

    try:
        from docx.shared import Inches
    except Exception:
        return

    try:
        table.autofit = False
    except Exception:
        pass

    for idx, w in enumerate(widths_in):
        try:
            col = table.columns[idx]
        except Exception:
            continue
        width = Inches(float(w))
        try:
            col.width = width
        except Exception:
            pass
        # Also set each cell width (helps some Word renderers honor the grid).
        try:
            for cell in col.cells:
                try:
                    cell.width = width
                except Exception:
                    continue
        except Exception:
            continue


def _shrink_table_font(table, *, size_pt: int) -> None:
    try:
        from docx.shared import Pt
    except Exception:
        return
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(size_pt)


def _eval_question_display_logic_visibility(
    question: dict,
    *,
    questions: dict,
    edf_overrides: dict[str, str],
    asked_qids: set[str],
) -> bool | None:
    """Return True/False when we can decide if the question is visible, else None."""

    display_logic = question.get("DisplayLogic")
    if not display_logic:
        return True
    if not isinstance(display_logic, dict):
        return None
    return _eval_boolean_expression_with_unasked_selected_false(
        display_logic, edf_overrides, asked_qids
    )


def _collect_branchlogic_embedded_field_keys(flow_obj: dict) -> set[str]:
    """Collect EmbeddedField LeftOperand keys used in SurveyFlow BranchLogic."""

    keys: set[str] = set()

    def walk(node: object) -> None:
        if node is None:
            return
        if isinstance(node, list):
            for child in node:
                walk(child)
            return
        if not isinstance(node, dict):
            return

        if str(node.get("Type") or "").strip() == "Branch":
            logic = node.get("BranchLogic")
            if isinstance(logic, dict):
                if_block = logic.get("0")
                if isinstance(if_block, dict):
                    for k, v in if_block.items():
                        if not str(k).isdigit() or not isinstance(v, dict):
                            continue
                        if (v.get("Type") or "") != "Expression":
                            continue
                        if (v.get("LogicType") or "").strip() != "EmbeddedField":
                            continue
                        left = (v.get("LeftOperand") or "").strip()
                        if left:
                            keys.add(left)

        for key in ("Flow", "Then", "Else", "ElseFlow"):
            if key in node:
                walk(node.get(key))

    if isinstance(flow_obj, dict):
        walk(flow_obj.get("Flow"))
    return keys


def _warn_about_unused_edf_overrides(
    doc,
    *,
    flow_obj: dict,
    edf_overrides: dict[str, str],
) -> None:
    """Warn when the user passes --edf keys that are not used in BranchLogic."""

    used = _collect_branchlogic_embedded_field_keys(flow_obj)
    unused = sorted([k for k in edf_overrides.keys() if k not in used])

    # Heuristic warnings about inconsistent EDF key spellings in the survey flow.
    # This is for messaging only (not for pruning): we want to surface bugs like
    # SVERSION vs S_VERSION early.
    def warn_norm(s: str) -> str:
        return re.sub(r"[_\\s-]+", "", str(s or "").strip().lower())

    groups: dict[str, set[str]] = {}
    for k in used:
        groups.setdefault(warn_norm(k), set()).add(k)

    related_variant_warnings: list[str] = []
    override_keys = set(edf_overrides.keys())
    for nk, variants in sorted(groups.items(), key=lambda kv: kv[0]):
        if len(variants) <= 1:
            continue
        # If user provided one variant, but not all, warn that some branches won't prune.
        if variants & override_keys and not variants.issubset(override_keys):
            missing = sorted(list(variants - override_keys))
            provided = sorted(list(variants & override_keys))
            related_variant_warnings.append(
                "Survey uses multiple EDF key spellings: "
                f"{', '.join(sorted(variants))}. "
                f"You provided: {', '.join(provided)}; missing: {', '.join(missing)}."
            )

    if not unused and not related_variant_warnings:
        return

    try:
        import sys

        if unused:
            print(
                f"[qsync:export-translation] WARNING: --edf keys not used in BranchLogic: {', '.join(unused)}. "
                "Next: verify the EDF key spelling in SurveyFlow BranchLogic (or remove unused --edf overrides).",
                file=sys.stderr,
            )
        for w in related_variant_warnings:
            print(
                f"[qsync:export-translation] WARNING: {w} "
                "Next: standardize EDF key spelling in SurveyFlow or provide the correct --edf key.",
                file=sys.stderr,
            )
        if used and (unused or related_variant_warnings):
            print(
                f"[qsync:export-translation] NOTE: BranchLogic EDF keys seen: {', '.join(sorted(used))}",
                file=sys.stderr,
            )
    except Exception:
        pass

    # Also surface this in the document so it remains visible when sharing the .docx.
    if unused:
        _add_system_note(
            doc,
            "WARNING: Some --edf keys are not used in any SurveyFlow BranchLogic. "
            "This usually indicates a survey EDF key mismatch (e.g., SVERSION vs S_VERSION). "
            f"Unused: {', '.join(unused)}. "
            "Next: fix the EDF key spelling in SurveyFlow or re-run export with the correct key.",
            depth=0,
        )
    for w in related_variant_warnings:
        _add_system_note(
            doc,
            f"WARNING: {w} Next: standardize EDF key spelling in SurveyFlow or provide the correct --edf key.",
            depth=0,
        )


def _container_add_paragraph(container, style: str | None = None):
    """Add a paragraph, reusing the placeholder paragraph in table cells when possible."""

    is_cell = container.__class__.__name__ == "_Cell"
    if is_cell:
        try:
            paras = getattr(container, "paragraphs", []) or []
            if len(paras) == 1 and not paras[0].text and not paras[0].runs:
                p = paras[0]
                if style:
                    try:
                        p.style = style
                    except Exception:
                        pass
                return p
        except Exception:
            pass

    if style is not None:
        try:
            return container.add_paragraph(style=style)
        except TypeError:
            return container.add_paragraph()
    return container.add_paragraph()


def _has_renderable_text(html_str: str) -> bool:
    s = str(html_str or "")
    if s.strip():
        # If it's mostly tags, fall back to stripped text.
        return bool(_strip_html(s).strip()) or should_treat_as_html(s)
    return False


def _format_randomizer(node: dict) -> str:
    flow_id = str(node.get("FlowID") or "").strip()
    subset = node.get("SubSet")
    even = node.get("EvenPresentation")
    parts = ["RANDOMIZER"]
    if flow_id:
        parts.append(f"FlowID={flow_id}")
    if subset is not None:
        parts.append(f"SubSet={subset}")
    if even is not None:
        parts.append(f"EvenPresentation={even}")
    if len(parts) == 1:
        return parts[0]
    return f"{parts[0]} ({', '.join(parts[1:])})"


def _render_web_service_node(doc, *, node: dict, depth: int) -> None:
    """Render a Qualtrics WebService SurveyFlow node as a single readable table."""

    flow_id = str(node.get("FlowID") or "").strip()
    method = str(node.get("Method") or "").strip()
    url = str(node.get("URL") or "").strip()
    content_type = str(node.get("ContentType") or "").strip()
    fire_and_forget = node.get("FireAndForget")
    stringify_values = node.get("StringifyValues")
    schema_version = node.get("SchemaVersion")

    # Build a single “card” table similar in spirit to a question block.
    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _set_table_column_widths_in(table, [1.6, 4.4])

    # Header row: merge cells and shade.
    header = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    p_hdr = _container_add_paragraph(header)
    _shade_panel_header_paragraph(p_hdr)
    r0 = p_hdr.add_run("WEB SERVICE")
    r0.bold = True
    if method or url:
        p_hdr.add_run(": ")
        r_m = p_hdr.add_run(method)
        r_m.bold = True
        p_hdr.add_run(" ")
        r_u = p_hdr.add_run(url)
        _style_url_run(r_u)
    if flow_id:
        p_hdr.add_run(" (FlowID=")
        r_id = p_hdr.add_run(flow_id)
        _style_id_run(r_id)
        p_hdr.add_run(")")

    def add_row(label: str):
        row = table.add_row().cells
        p0 = _container_add_paragraph(row[0])
        r = p0.add_run(label)
        r.bold = True
        return row[1]

    # Core fields
    if method:
        c = add_row("Method")
        _container_add_paragraph(c).add_run(method)
    if url:
        c = add_row("URL")
        p = _container_add_paragraph(c)
        r = p.add_run(url)
        _style_url_run(r)
    if content_type:
        c = add_row("ContentType")
        _container_add_paragraph(c).add_run(content_type)

    # Credential
    cred = node.get("Credential") or {}
    if isinstance(cred, dict) and cred:
        cid = str(cred.get("ID") or "").strip()
        fmt = str(cred.get("ParamFormat") or "").strip()
        name = str(cred.get("ParamName") or "").strip()
        templ = str(cred.get("ParamTemplate") or "").strip()
        c = add_row("Credential")
        p = _container_add_paragraph(c)
        parts = [x for x in [fmt, name] if x]
        p.add_run(" ".join(parts) if parts else "(present)")
        extras = []
        if cid:
            extras.append(f"ID={cid}")
        if templ:
            extras.append(f"Template={templ}")
        if extras:
            p.add_run(" (" + ", ".join(extras) + ")")

    # Request params
    req_params = node.get("RequestParams") or []
    if isinstance(req_params, list) and req_params:
        items: list[tuple[str, str]] = []
        for e in req_params:
            if not isinstance(e, dict):
                continue
            k = str(e.get("key") or "").strip()
            v = str(e.get("value") or "").strip()
            if k:
                items.append((k, v))
        if items:
            c = add_row("RequestParams")
            for i, (k, v) in enumerate(items):
                p = _container_add_paragraph(c) if i else _container_add_paragraph(c)
                rk = p.add_run(f"{k}: ")
                rk.bold = True
                _add_text_with_edf_styling(p, v)

    # Headers
    headers = node.get("Headers") or []
    if isinstance(headers, list) and headers:
        items = []
        for e in headers:
            if not isinstance(e, dict):
                continue
            k = str(e.get("key") or "").strip()
            v = str(e.get("value") or "").strip()
            if k:
                items.append((k, v))
        if items:
            c = add_row("Headers")
            for i, (k, v) in enumerate(items):
                p = _container_add_paragraph(c) if i else _container_add_paragraph(c)
                rk = p.add_run(f"{k}: ")
                rk.bold = True
                _add_text_with_edf_styling(p, v)

    # Body params
    body = node.get("Body") or []
    if isinstance(body, list) and body:
        items = []
        for e in body:
            if not isinstance(e, dict):
                continue
            k = str(e.get("key") or "").strip()
            v = str(e.get("value") or "").strip()
            if k:
                items.append((k, v))
        if items:
            c = add_row("Body")
            for i, (k, v) in enumerate(items):
                p = _container_add_paragraph(c) if i else _container_add_paragraph(c)
                rk = p.add_run(f"{k}: ")
                rk.bold = True
                _add_text_with_edf_styling(p, v)

    # Response map: response path -> EDF field
    resp_map = node.get("ResponseMap") or []
    if isinstance(resp_map, list) and resp_map:
        items = []
        for e in resp_map:
            if not isinstance(e, dict):
                continue
            k = str(e.get("key") or "").strip()
            v = str(e.get("value") or "").strip()
            if k or v:
                items.append((k, v))
        if items:
            c = add_row("ResponseMap")
            for i, (k, v) in enumerate(items):
                p = _container_add_paragraph(c) if i else _container_add_paragraph(c)
                p.add_run(f"{k} → ")
                r_edf = p.add_run(v)
                _style_edf_run(r_edf)

    # Flags
    flags: list[str] = []
    if fire_and_forget is not None:
        flags.append(f"FireAndForget={fire_and_forget}")
    if stringify_values is not None:
        flags.append(f"StringifyValues={stringify_values}")
    if schema_version is not None:
        flags.append(f"SchemaVersion={schema_version}")
    if flags:
        c = add_row("Flags")
        _container_add_paragraph(c).add_run(", ".join(flags))

    _set_indent_table(table, depth=depth)
    _shrink_table_font(table, size_pt=9)
    _add_doc_spacer_paragraph(doc, depth=depth)


def _render_eos_message_content(
    doc,
    *,
    library_id: str,
    message_id: str,
    flow_id: str,
    depth: int,
    include_html_source: bool,
    layout_heuristics: bool,
    base_language: str,
    render_language: str | None,
    compare_to_base: bool,
) -> None:
    payload = _read_eos_message_from_disk(library_id, message_id)
    if payload is None:
        _add_system_note(
            doc,
            f"(EOS message content not pulled to disk for {library_id}/{message_id}; run `qsync eos pull`.)",
            depth=depth,
        )
        return

    messages = payload.get("messages") or {}
    if not isinstance(messages, dict) or not messages:
        _add_system_note(
            doc,
            f"(EOS message {library_id}/{message_id} has no message variants on disk.)",
            depth=depth,
        )
        return

    p = doc.add_paragraph()
    _set_indent(p, depth=depth)
    _style_table_label_paragraph(p)
    if flow_id:
        p.add_run(
            f"EOS MESSAGE CONTENT (FlowID={flow_id}; {library_id}/{message_id})"
        ).bold = True
    else:
        p.add_run(f"EOS MESSAGE CONTENT ({library_id}/{message_id})").bold = True

    # Align EOS message presentation with question translations:
    # - monolingual exports: one column (target language if specified, else base)
    # - bilingual exports (--compare-to-base): base + target side-by-side columns
    base_key = _normalize_lang_code(base_language).lower()
    target_key = (
        _normalize_lang_code(render_language).lower() if render_language else ""
    )

    selected_keys = sorted(messages.keys(), key=str)
    if compare_to_base:
        preferred = [k for k in [base_key, target_key] if k and k in messages]
        if preferred:
            selected_keys = preferred
        elif base_key in messages:
            selected_keys = [base_key]
        elif target_key in messages:
            selected_keys = [target_key]
    else:
        if target_key and target_key in messages:
            selected_keys = [target_key]
        elif base_key in messages:
            # Baseline exports should not show all language variants.
            selected_keys = [base_key]

    cols = max(len(selected_keys), 1)
    table = doc.add_table(rows=2, cols=cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    if cols == 1:
        _set_table_column_widths_in(table, [6.0])
    else:
        per = 6.0 / cols
        _set_table_column_widths_in(table, [per for _ in range(cols)])

    # Header row: language codes
    hdr = table.rows[0].cells
    for idx in range(cols):
        key = selected_keys[idx] if idx < len(selected_keys) else ""
        hdr[idx].text = _normalize_lang_code(key) if key else ""
    try:
        for cell in hdr:
            for pp in getattr(cell, "paragraphs", []) or []:
                for rr in getattr(pp, "runs", []) or []:
                    rr.bold = True
    except Exception:
        pass

    # Content row: render message variants
    row = table.rows[1].cells
    for idx in range(cols):
        key = selected_keys[idx] if idx < len(selected_keys) else ""
        value = str(messages.get(key) or "") if key else ""
        _add_rich_text_block(
            row[idx],
            value,
            depth=0,
            include_html_source=include_html_source,
            layout_heuristics=layout_heuristics,
        )

    _set_indent_table(table, depth=depth)
    _add_doc_spacer_paragraph(doc, depth=depth)


def _render_embedded_data_node(
    doc,
    *,
    node: dict,
    depth: int,
    edf_overrides: dict[str, str] | None = None,
) -> None:
    """Render a SurveyFlow EmbeddedData node inline as a mini table."""

    flow_id = str(node.get("FlowID") or "").strip()
    entries = node.get("EmbeddedData") or []
    if not isinstance(entries, list) or not entries:
        p = doc.add_paragraph()
        _set_indent(p, depth=depth)
        if flow_id:
            p.add_run("EMBEDDED DATA (FlowID=")
            r = p.add_run(flow_id)
            _style_id_run(r)
            p.add_run("): (no assignments)")
        else:
            p.add_run("EMBEDDED DATA: (no assignments)")
        return

    p = doc.add_paragraph()
    _set_indent(p, depth=depth)
    _style_table_label_paragraph(p)
    if flow_id:
        r1 = p.add_run("EMBEDDED DATA WRITES (FlowID=")
        r1.bold = True
        r2 = p.add_run(flow_id)
        r2.bold = True
        _style_id_run(r2)
        r3 = p.add_run(")")
        r3.bold = True
    else:
        p.add_run("EMBEDDED DATA WRITES").bold = True

    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    # 8.5in page width with 1.25in L/R margins -> ~6.0in usable width.
    # Make Field/Value wider; Type/FlowID narrower.
    _set_table_column_widths_in(table, [2.0, 2.5, 0.75, 0.75])

    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    hdr[2].text = "Type"
    hdr[3].text = "FlowID"
    # Bold header row (translator scanability).
    try:
        for cell in hdr:
            for p in getattr(cell, "paragraphs", []) or []:
                for r in getattr(p, "runs", []) or []:
                    r.bold = True
    except Exception:
        pass

    seen_fields: set[str] = getattr(doc, "_qsync_translation_edf_seen_fields", set())
    overridden_in_this_node: set[str] = set()

    for e in entries:
        if not isinstance(e, dict):
            continue
        field = str(e.get("Field") or "").strip()
        if not field:
            continue
        value = e.get("Value")
        if (
            edf_overrides
            and field in edf_overrides
            and (field not in seen_fields or field in overridden_in_this_node)
        ):
            value_str = str(edf_overrides[field])
            overridden_in_this_node.add(field)
        else:
            value_str = str(value) if value is not None else EMBEDDED_EMPTY_VALUE
        ed_type = str(e.get("Type") or "").strip() or "Custom"

        row = table.add_row().cells
        row[0].text = field
        _style_edf_field_cell(row[0])

        # Value: allow EDF token highlighting within the cell.
        try:
            pval = _container_add_paragraph(row[1])
            _add_text_with_edf_styling(pval, value_str)
        except Exception:
            row[1].text = value_str

        row[2].text = ed_type
        try:
            pflow = _container_add_paragraph(row[3])
            r = pflow.add_run(flow_id)
            _style_id_run(r)
        except Exception:
            row[3].text = flow_id

        seen_fields.add(field)

    setattr(doc, "_qsync_translation_edf_seen_fields", seen_fields)

    _shrink_table_font(table, size_pt=9)
    _add_doc_spacer_paragraph(doc, depth=depth)


def _describe_question_type(
    *, question_type: str, selector: str, subselector: str
) -> str:
    qt = (question_type or "").strip()
    sel = (selector or "").strip()
    sub = (subselector or "").strip()
    code = "/".join([p for p in [qt, sel, sub] if p])

    key = (qt, sel, sub)
    # Common Qualtrics combos we see in this repo.
    pretty = {
        ("MC", "SAVR", ""): "Multiple Choice (single answer)",
        ("MC", "MAVR", ""): "Multiple Choice (multiple answer)",
        ("TE", "SL", ""): "Text Entry (single line)",
        ("TE", "ML", ""): "Text Entry (multi line)",
        ("DB", "", ""): "Descriptive Text",
        ("Matrix", "", ""): "Matrix",
        ("Timing", "PageTimer", ""): "Timing (page timer)",
        ("Timing", "", ""): "Timing",
    }.get(key)
    if pretty:
        return pretty
    # Partial matching on (QuestionType, Selector)
    pretty2 = {
        ("MC", "SAVR"): "Multiple Choice (single answer)",
        ("MC", "MAVR"): "Multiple Choice (multiple answer)",
        ("TE", "SL"): "Text Entry (single line)",
        ("TE", "ML"): "Text Entry (multi line)",
        ("Timing", "PageTimer"): "Timing (page timer)",
    }.get((qt, sel))
    if pretty2:
        return pretty2
    return code or qt or "Question"


def _question_type_abbrev_and_label(
    *, question_type: str, selector: str, subselector: str
) -> tuple[str, str]:
    """Return (abbrev, label) for the QID metadata line + legend."""

    qt = (question_type or "").strip()
    sel = (selector or "").strip()
    _sub = (subselector or "").strip()

    # Abbrev should be short and consistent; label should be human-readable.
    base: dict[str, tuple[str, str]] = {
        "SBS": ("SBS", "SBS Matrix"),
        "MC": ("MC", "Multiple Choice"),
        "TE": ("TE", "Text Entry"),
        "DB": ("DB", "Descriptive Text"),
        "Matrix": ("MAT", "Matrix"),
        "Timing": ("TIM", "Timing"),
        "Slider": ("SLD", "Slider"),
        "Captcha": ("CAP", "Captcha"),
        "Meta": ("META", "Meta"),
        "CS": ("CS", "Constant Sum"),
    }
    if qt in base:
        abbrev, label = base[qt]
    elif qt:
        # Fallback: keep short codes as-is; otherwise abbreviate.
        if len(qt) <= 5 and qt.replace("_", "").isalnum():
            abbrev = qt.upper()
        else:
            abbrev = re.sub(r"[^A-Za-z0-9]", "", qt.upper())[:5] or "QT"
        label = qt
    else:
        return ("QT", "Question")

    # Add a little specificity for legend clarity when we can.
    if qt == "Meta" and sel:
        label = f"Meta ({sel})"
    if qt == "MC" and sel in {"SAVR", "MAVR", "SAHR"}:
        # Keep code stable (MC), but clarify meaning in the legend.
        label = "Multiple Choice"
    if qt == "TE" and sel in {"SL", "ML", "FORM"}:
        label = "Text Entry"

    return (abbrev, label)


def _question_type_abbrev(q: dict) -> str:
    """Return short abbreviation for question type (for HTML rendering)."""
    qt = str(q.get("QuestionType") or "").strip()
    sel = str(q.get("Selector") or "").strip()
    sub = str(q.get("SubSelector") or "").strip()
    abbrev, _ = _question_type_abbrev_and_label(
        question_type=qt, selector=sel, subselector=sub
    )
    return abbrev


@dataclass(frozen=True)
class EndSurveyMessageRef:
    flow_id: str | None
    library_id: str
    message_id: str


def _extract_end_survey_message_refs(survey_flow: dict) -> List[EndSurveyMessageRef]:
    refs: List[EndSurveyMessageRef] = []

    def walk(node: object) -> None:
        if node is None:
            return
        if isinstance(node, list):
            for child in node:
                walk(child)
            return
        if not isinstance(node, dict):
            return

        if str(node.get("Type") or "") == "EndSurvey":
            opts = node.get("Options") or {}
            if (opts.get("SurveyTermination") or "") == "DisplayMessage":
                lib_id = (opts.get("EOSMessageLibrary") or "").strip()
                msg_id = (opts.get("EOSMessage") or "").strip()
                if lib_id and msg_id:
                    refs.append(
                        EndSurveyMessageRef(
                            flow_id=(node.get("FlowID") or None),
                            library_id=lib_id,
                            message_id=msg_id,
                        )
                    )

        for key in ("Flow", "Then", "Else", "ElseFlow"):
            if key in node:
                walk(node.get(key))

    if isinstance(survey_flow, dict):
        walk(survey_flow.get("Flow"))

    # Deduplicate while preserving order
    seen = set()
    out: List[EndSurveyMessageRef] = []
    for r in refs:
        k = (r.library_id, r.message_id, r.flow_id)
        if k in seen:
            continue
        seen.add(k)
        out.append(r)
    return out
